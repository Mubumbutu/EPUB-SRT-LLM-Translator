# file_processors.py
import copy
import ebooklib
import html as html_module
import logging
import os
import re
import shutil
import statistics
import uuid
from abc import ABC, abstractmethod
from ebooklib import epub
from epub_utils import NAMESPACES, read_epub
from lxml import etree
from PyQt6.QtCore import pyqtSignal, QThread
from typing import Dict, List, Optional, Set, Tuple

logger = logging.getLogger(__name__)

NAMESPACES = {
    'x': 'http://www.w3.org/1999/xhtml',
    'opf': 'http://www.idpf.org/2007/opf',
    'dc': 'http://purl.org/dc/elements/1.1/',
    'ncx': 'http://www.daisy.org/z3986/2005/ncx/',
    'epub': 'http://www.idpf.org/2007/ops',
}

class FileProcessor(ABC):
    @abstractmethod
    def load(self, path: str) -> Tuple[List[Dict], Optional[any]]:
        pass

    @abstractmethod
    def get_file_type(self) -> str:
        pass

class EPUBProcessor(FileProcessor):
    def __init__(self, app_settings: dict):
        self.app_settings = app_settings
        self.book = None
        self.paragraphs = []
        self.image_items = []
        self.css_class_styles = {}
 
        self.skip_inline_tags = app_settings.get('skip_inline_tags', {})
 
        self.PRIORITY_TAGS = ['p', 'pre', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']
 
        self.NON_INLINE_ELEMENTS = {
            'address', 'blockquote', 'dialog', 'div', 'figure', 'figcaption',
            'footer', 'header', 'legend', 'main', 'p', 'pre', 'search', 'article',
            'aside', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'hgroup', 'nav',
            'section', 'dd', 'dl', 'dt', 'menu', 'ol', 'ul', 'table', 'caption',
            'colgroup', 'col', 'thead', 'tbody', 'tfoot', 'tr', 'td', 'th', 'li'
        }
 
        self.NOISE_TAGS = ['rt', 'rp']
 
        self.RESERVE_TAGS = [
            'img', 'code', 'br', 'hr', 'sub', 'sup', 'kbd',
            'abbr', 'wbr', 'var', 'canvas', 'svg', 'script',
            'style', 'math'
        ]
 
        logger.info(f"EPUBProcessor initialized with {len(self.RESERVE_TAGS)} structural reserve tags, "
                    f"skip_inline_tags: {self.skip_inline_tags}")
 
    def _parse_css_files(self):
        css_class_styles = {}
        KEEP_CSS_PROPS = {
            'text-align',
            'font-style',
            'font-weight',
            'font-size',
            'color',
            'text-indent',
            'margin-left',
            'margin-right',
            'line-height',
            'background-color',
        }

        for manifest_item in self.book.manifest_items:
            if not isinstance(manifest_item, dict):
                continue
            if manifest_item.get('media_type') != 'text/css':
                continue

            href = manifest_item.get('href', '')
            if not href:
                continue

            css_path = os.path.join(self.book.content_dir, href)
            if not os.path.exists(css_path):
                continue

            try:
                with open(css_path, 'r', encoding='utf-8', errors='ignore') as f:
                    css_content = f.read()

                for match in re.finditer(r'\.([\w\-]+)\s*\{([^}]*)\}', css_content, re.DOTALL):
                    class_name = match.group(1)
                    declarations = match.group(2)
                    found_styles = {}

                    for decl in declarations.split(';'):
                        decl = decl.strip()
                        if ':' not in decl:
                            continue
                        prop, _, val = decl.partition(':')
                        prop = prop.strip().lower()
                        val = val.strip()

                        if val.endswith('!important'):
                            val = val[:-len('!important')].strip()

                        if prop in KEEP_CSS_PROPS:
                            if prop == 'text-align' and val not in ('center', 'left', 'right', 'justify'):
                                continue
                            if prop == 'font-style' and val not in ('italic', 'oblique', 'normal'):
                                continue
                            if prop == 'font-weight' and val not in ('bold', 'bolder', 'normal'):
                                continue
                            found_styles[prop] = val

                    if found_styles:
                        css_class_styles[class_name] = found_styles

            except Exception as e:
                logger.debug(f"CSS parse error for {css_path}: {e}")

        logger.info(f"EPUBProcessor: parsed {len(css_class_styles)} CSS class style(s)")
        return css_class_styles
 
    def get_file_type(self) -> str:
        return "epub"
 
    def load(self, path: str) -> Tuple[List[Dict], any]:
        try:
            self.book = read_epub(path)
            self.paragraphs = []
            self.image_items = []
            self.css_class_styles = self._parse_css_files()
 
            seen = set()
 
            spine_order = {}
            nav_hrefs = set()
            try:
                opf_path = os.path.join(self.book.temp_dir, self.book.opf_path)
 
                with open(opf_path, 'rb') as f:
                    opf_tree = etree.parse(f)
 
                spine_elem = opf_tree.find('.//opf:spine', namespaces=NAMESPACES)
 
                if spine_elem is not None:
                    itemrefs = spine_elem.findall('.//opf:itemref', namespaces=NAMESPACES)
                    for idx, itemref in enumerate(itemrefs):
                        idref = itemref.get('idref')
                        if idref:
                            spine_order[idref] = idx
                    logger.info(f"📚 Spine order built from OPF: {len(spine_order)} items")
                else:
                    logger.warning("⚠️ Warning: No <spine> found in OPF. Using alphabetical order.")
 
                manifest_elem = opf_tree.find('.//opf:manifest', namespaces=NAMESPACES)
                if manifest_elem is not None:
                    for mitem in manifest_elem.findall('opf:item', namespaces=NAMESPACES):
                        props = mitem.get('properties', '')
                        if 'nav' in props.split():
                            href = mitem.get('href', '')
                            if href:
                                nav_hrefs.add(href)
                                nav_hrefs.add(os.path.basename(href))
                if nav_hrefs:
                    logger.info(f"📚 Nav documents to skip: {nav_hrefs}")
 
            except Exception as e:
                logger.warning(f"⚠️ Warning: Could not build spine order ({e}). Using alphabetical order.")
                spine_order = {}
 
            doc_items = list(self.book.get_items_of_type('DOCUMENT'))
 
            def get_sort_key(item):
                try:
                    item_id = None
 
                    for manifest_item in self.book.manifest_items:
                        if isinstance(manifest_item, dict):
                            if manifest_item.get('href') == item.href:
                                item_id = manifest_item.get('id')
                                break
                        elif hasattr(manifest_item, 'href') and hasattr(manifest_item, 'id'):
                            if manifest_item.href == item.href:
                                item_id = manifest_item.id
                                break
 
                    if not item_id:
                        item_id = os.path.splitext(os.path.basename(item.href))[0]
 
                    if item_id in spine_order:
                        return (0, spine_order[item_id])
                    else:
                        return (1, item.href)
 
                except Exception as e:
                    logger.error(f"⚠️ Sort key error for {item.href}: {e}")
                    return (2, item.href)
 
            doc_items.sort(key=get_sort_key)
 
            logger.info(f"📖 Processing {len(doc_items)} documents in spine order...\n")
 
            total_items = 0
            items_with_body = 0
            total_priority_found = 0
            total_processed = 0
 
            for item in doc_items:
                total_items += 1
 
                item_basename = os.path.basename(item.href)
                if item.href in nav_hrefs or item_basename in nav_hrefs:
                    logger.debug(f"⏭ Skipping nav document (manifest): {item.href}")
                    continue
 
                if item.data is None:
                    logger.debug(f"❌ Item {item.href}: No data")
                    continue
 
                body = item.data.find('.//x:body', namespaces=NAMESPACES)
                if body is None:
                    logger.debug(f"❌ Item {item.href}: No body")
                    continue
 
                nav_elems = item.data.xpath('.//x:nav', namespaces=NAMESPACES)
                is_nav_doc = False
                for nav_el in nav_elems:
                    epub_type = nav_el.get(f'{{{NAMESPACES["epub"]}}}type', '')
                    if epub_type:
                        is_nav_doc = True
                        break
                if is_nav_doc:
                    logger.debug(f"⏭ Skipping navigation document (epub:type): {item.href}")
                    continue
 
                items_with_body += 1
 
                priority_xpath = ' | '.join([f'.//x:{tag}' for tag in self.PRIORITY_TAGS])
                priority_elements = body.xpath(priority_xpath, namespaces=NAMESPACES)
 
                li_elements = body.xpath('.//x:li', namespaces=NAMESPACES)
 
                logger.debug(f"📄 {item.href}")
                logger.debug(f"   Priority elements: {len(priority_elements)}")
                logger.debug(f"   <li> elements: {len(li_elements)}")
 
                if priority_elements:
                    total_priority_found += len(priority_elements)
                    for idx, elem in enumerate(priority_elements[:3]):
                        tag = etree.QName(elem).localname
                        text = etree.tostring(elem, encoding='unicode', method='text').strip()[:60]
                        logger.debug(f"   Priority[{idx+1}] <{tag}>: {text}...")
 
                if li_elements:
                    for idx, elem in enumerate(li_elements[:3]):
                        text = etree.tostring(elem, encoding='unicode', method='text').strip()[:60]
                        logger.debug(f"   <li>[{idx+1}]: {text}...")
 
                before_count = len(self.paragraphs)
 
                logger.debug(f"   🔍 Starting extraction...")
                self._extract_elements_lxml(body, item.href, seen)
 
                after_count = len(self.paragraphs)
                added = after_count - before_count
                total_processed += added
 
                logger.debug(f"   ✓ Added to paragraphs: {added}\n")
 
            logger.info("=" * 60)
            logger.info("SUMMARY:")
            logger.info(f"  Total items: {total_items}")
            logger.info(f"  Items with body: {items_with_body}")
            logger.info(f"  Priority elements found: {total_priority_found}")
            logger.info(f"  Fragments added: {total_processed}")
            logger.info("=" * 60 + "\n")
 
            try:
                for item in self.book.get_items_of_type('IMAGE'):
                    content = getattr(item, 'content', b'') or b''
                    if not content:
                        continue
                    item_id = getattr(item, 'id', '') or getattr(item, 'uid', '')
                    file_name = getattr(item, 'file_name', '') or getattr(item, 'href', '') or item_id
                    media_type = getattr(item, 'media_type', 'image/jpeg') or 'image/jpeg'
                    self.image_items.append({
                        'id': item_id,
                        'file_name': file_name,
                        'media_type': media_type,
                        'content': content,
                    })
                logger.info(f"EPUBProcessor: collected {len(self.image_items)} image items")
            except Exception as e:
                logger.warning(f"EPUBProcessor: could not collect image items: {e}")
 
            return self.paragraphs, self.book
 
        except Exception as e:
            logger.error(f"EPUB load error: {e}", exc_info=True)
            raise
 
    def _is_non_translatable_content(self, text: str) -> bool:
        text_no_placeholders = re.sub(r'<id_\d{2}>', '', text)
        text_no_placeholders = re.sub(r'</?p_\d{2}>', '', text_no_placeholders)
 
        text_stripped = text_no_placeholders.strip()
 
        if not text_stripped:
            logger.debug("Non-translatable: empty after removing placeholders")
            return True
 
        if re.match(r'^[\s\d\.,!?:;…\-–—\'\"\u201e\u201d\u201a\u2019]+$', text_stripped):
            logger.debug(f"Non-translatable: only special chars: {repr(text_stripped)}")
            return True
 
        if re.match(r'^[\s\*•–—]+$', text_stripped):
            logger.debug(f"Non-translatable: chapter separator: {repr(text_stripped)}")
            return True
 
        if re.match(r'^([\*•–—])\s*(\1\s*)+$', text_stripped):
            logger.debug(f"Non-translatable: repeated special chars: {repr(text_stripped)}")
            return True
 
        return False
 
    def _extract_elements_lxml(self, root, item_href, seen):
        CONTAINER_TAGS = {
            'ul', 'ol', 'dl',
            'table', 'tbody', 'thead', 'tfoot', 'tr',
            'div', 'section', 'article', 'aside', 'nav', 'main',
            'header', 'footer', 'figure', 'body',
        }
        XLINK_NS = 'http://www.w3.org/1999/xlink'
 
        def _resolve_src(src, item_href):
            item_dir = os.path.dirname(item_href) if '/' in item_href else ''
            if item_dir:
                return os.path.normpath(os.path.join(item_dir, src)).replace('\\', '/')
            return src.lstrip('./')
 
        def _collect_img_paragraphs(element, item_href):
            imgs = (
                element.xpath('.//x:img', namespaces=NAMESPACES)
                or element.xpath('.//img')
            )
            result = []
            for img in imgs:
                src = img.get('src', '')
                if src and not src.startswith('data:'):
                    result.append({
                        'id': f'img_{uuid.uuid4().hex[:8]}',
                        'original_text': '',
                        'translated_text': '',
                        'is_translated': False,
                        'item_href': item_href,
                        'element_type': 'image',
                        'image_href': _resolve_src(src, item_href),
                        'has_mismatch': False,
                    })
            svg_elems = (
                element.xpath('.//x:svg', namespaces=NAMESPACES)
                or element.xpath('.//svg')
            )
            for svg in svg_elems:
                for svg_child in svg.iter():
                    href = (
                        svg_child.get(f'{{{XLINK_NS}}}href', '')
                        or svg_child.get('href', '')
                    )
                    if href and not href.startswith('data:') and not href.startswith('#'):
                        result.append({
                            'id': f'img_{uuid.uuid4().hex[:8]}',
                            'original_text': '',
                            'translated_text': '',
                            'is_translated': False,
                            'item_href': item_href,
                            'element_type': 'image',
                            'image_href': _resolve_src(href, item_href),
                            'has_mismatch': False,
                        })
                        break
            return result
 
        for child in root:
            if not isinstance(child, etree._Element):
                logger.debug(f"   ⏭ Skipping non-element node")
                continue
 
            tag_name = etree.QName(child).localname
 
            if tag_name == 'img':
                src = child.get('src', '')
                if src and not src.startswith('data:'):
                    self.paragraphs.append({
                        'id': f'img_{uuid.uuid4().hex[:8]}',
                        'original_text': '',
                        'translated_text': '',
                        'is_translated': False,
                        'item_href': item_href,
                        'element_type': 'image',
                        'image_href': _resolve_src(src, item_href),
                        'has_mismatch': False,
                    })
                continue
 
            if tag_name == 'svg':
                found_href = ''
                for svg_child in child.iter():
                    href = (
                        svg_child.get(f'{{{XLINK_NS}}}href', '')
                        or svg_child.get('href', '')
                    )
                    if href and not href.startswith('data:') and not href.startswith('#'):
                        found_href = href
                        break
                if found_href:
                    self.paragraphs.append({
                        'id': f'img_{uuid.uuid4().hex[:8]}',
                        'original_text': '',
                        'translated_text': '',
                        'is_translated': False,
                        'item_href': item_href,
                        'element_type': 'image',
                        'image_href': _resolve_src(found_href, item_href),
                        'has_mismatch': False,
                    })
                continue
 
            if tag_name in self.PRIORITY_TAGS:
                text_content = (etree.tostring(child, encoding='unicode', method='text') or '').strip()
                img_children = (
                    child.xpath('.//x:img', namespaces=NAMESPACES)
                    or child.xpath('.//img')
                )
                svg_children = (
                    child.xpath('.//x:svg', namespaces=NAMESPACES)
                    or child.xpath('.//svg')
                )
                has_visual = bool(img_children or svg_children)
                if not text_content and has_visual:
                    self.paragraphs.extend(_collect_img_paragraphs(child, item_href))
                    continue
                logger.debug(f"   ✓ Priority tag <{tag_name}> - processing directly")
                self._process_element_lxml(child, item_href, seen)
                continue
 
            if tag_name in CONTAINER_TAGS:
                child_elements = [c for c in child if isinstance(c, etree._Element)]
 
                if not child_elements and self._has_any_text(child):
                    logger.debug(f"   📝 Leaf container <{tag_name}> with text - processing directly")
                    self._process_element_lxml(child, item_href, seen)
                else:
                    logger.debug(f"   📦 Container <{tag_name}> - recursing into children")
                    self._extract_elements_lxml(child, item_href, seen)
                continue
 
            element_has_content = False
 
            if self._is_inline_only_lxml(child):
                element_has_content = True
                logger.debug(f"   📝 <{tag_name}> is inline-only - will process")
            elif self._has_any_text(child):
                element_has_content = True
                text_preview = self._get_element_text(child)[:40]
                logger.debug(f"   📝 <{tag_name}> has text content - will process: {text_preview}...")
 
            if element_has_content:
                text_in_child = (etree.tostring(child, encoding='unicode', method='text') or '').strip()
                img_in_child = (
                    child.xpath('.//x:img', namespaces=NAMESPACES)
                    or child.xpath('.//img')
                )
                svg_in_child = (
                    child.xpath('.//x:svg', namespaces=NAMESPACES)
                    or child.xpath('.//svg')
                )
                if not text_in_child and (img_in_child or svg_in_child):
                    logger.debug(f"   📷 <{tag_name}> inline-only but visual-only - collecting as image(s)")
                    self.paragraphs.extend(_collect_img_paragraphs(child, item_href))
                else:
                    logger.debug(f"   ✓ Processing element: <{tag_name}>")
                    self._process_element_lxml(child, item_href, seen)
            else:
                logger.debug(f"   🔄 <{tag_name}> has no direct content - recursing to find children")
                self._extract_elements_lxml(child, item_href, seen)
 
    def _is_inline_only_lxml(self, element):
        non_inline_xpath = ' | '.join([f'.//x:{tag}' for tag in self.NON_INLINE_ELEMENTS])
        found = element.xpath(non_inline_xpath, namespaces=NAMESPACES)
        return len(found) == 0
 
    def _has_any_text(self, element):
        if element.text and element.text.strip():
            return True
 
        for descendant in element.iter():
            if descendant is element:
                continue
            if descendant.text and descendant.text.strip():
                return True
            if descendant.tail and descendant.tail.strip():
                return True
 
        return False
 
    def _get_element_text(self, element):
        return etree.tostring(element, encoding='unicode', method='text').strip()
 
    def _process_element_lxml(self, element, item_href, seen):
        tag_name = etree.QName(element).localname
 
        CONTAINER_TAGS = {
            'ul', 'ol', 'dl', 'table', 'tbody', 'thead', 'tfoot', 'tr',
            'div', 'section', 'article', 'aside', 'nav', 'main',
            'header', 'footer', 'figure', 'body'
        }
 
        if tag_name in CONTAINER_TAGS:
            child_elements = [c for c in element if isinstance(c, etree._Element)]
            if child_elements:
                logger.warning(f"⚠️ Container <{tag_name}> with children reached _process_element_lxml - SKIPPING")
                return
 
        use_inline = self.app_settings.get('use_inline_formatting', True)
 
        if use_inline:
            self._process_element_new(element, item_href, seen)
        else:
            self._process_element_legacy(element, item_href, seen)
 
    def _process_element_legacy(self, element, item_href, seen):
        tag_name = etree.QName(element).localname
 
        if element.get('id') is None:
            element.set('id', f"trans_{uuid.uuid4()}")
 
        element_id = element.get('id')
        element_copy = copy.deepcopy(element)
 
        for noise_tag in self.NOISE_TAGS:
            for noise_elem in element_copy.xpath(f'.//x:{noise_tag}', namespaces=NAMESPACES):
                parent = noise_elem.getparent()
                if parent is not None:
                    parent.remove(noise_elem)
 
        reserve_elements = []
        placeholder_pattern = '<id_{:02d}>'
        reserve_counter = 0
 
        for reserve_tag in self.RESERVE_TAGS:
            for reserve_elem in element_copy.xpath(f'.//x:{reserve_tag}', namespaces=NAMESPACES):
                reserve_html = etree.tostring(
                    reserve_elem,
                    encoding='unicode',
                    method='xml',
                    with_tail=False
                )
                reserve_html = re.sub(r'\s+xmlns(?::\w+)?="[^"]*"', '', reserve_html)
                reserve_elements.append(reserve_html)
 
                placeholder = placeholder_pattern.format(reserve_counter)
                reserve_counter += 1
 
                parent = reserve_elem.getparent()
                if parent is not None:
                    tail = reserve_elem.tail or ''
                    prev = reserve_elem.getprevious()
                    if prev is not None:
                        existing = prev.tail or ''
                        space_before = '' if (not existing or existing[-1].isspace()) else ' '
                        prev.tail = existing + space_before + placeholder + tail
                    else:
                        existing = parent.text or ''
                        space_before = '' if (not existing or existing[-1].isspace()) else ' '
                        parent.text = existing + space_before + placeholder + tail
                    parent.remove(reserve_elem)
 
        clean_text = self._get_element_text(element_copy)
 
        if not clean_text or not clean_text.strip():
            return
 
        prefix_tags, suffix_tags, clean_text = self._extract_boundary_reserve_tags(clean_text)
 
        is_non_translatable = self._is_non_translatable_content(clean_text)
 
        if is_non_translatable:
            logger.debug(f"Marking element {element_id} as non-translatable (will be preserved in EPUB)")
 
        key = (item_href, element_id)
        if key in seen:
            return
        seen.add(key)
 
        original_html = etree.tostring(
            element,
            encoding='unicode',
            method='xml',
            pretty_print=False
        )
 
        HEADING_LOCAL_NAMES = {'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}
        element_type = 'heading' if tag_name in HEADING_LOCAL_NAMES else tag_name
 
        para = {
            "id": element_id,
            "original_text": clean_text,
            "translated_text": "",
            "is_translated": False,
            "item_href": item_href,
            "element_type": element_type,
            "original_html": original_html,
            "has_mismatch": False,
            "reserve_elements": reserve_elements,
            "placeholder_pattern": placeholder_pattern,
            "processing_mode": "legacy",
            "prefix_reserve_tags": prefix_tags,
            "suffix_reserve_tags": suffix_tags,
            "is_non_translatable": is_non_translatable
        }
 
        style_parts = {}
        elem_align = (element.get('align') or '').strip().lower()
        if elem_align in ('center', 'left', 'right', 'justify'):
            style_parts['text-align'] = elem_align
        elem_style_raw = (element.get('style') or '').strip()
        for decl in elem_style_raw.split(';'):
            decl = decl.strip()
            if not decl or ':' not in decl:
                continue
            prop, _, val = decl.partition(':')
            prop = prop.strip().lower()
            if prop in {'text-align', 'color', 'font-size', 'font-weight', 'font-style',
                        'text-indent', 'margin-left', 'margin-right', 'background-color',
                        'letter-spacing', 'line-height'}:
                style_parts[prop] = val.strip()
        for cls in (element.get('class') or '').split():
            for prop, val in self.css_class_styles.get(cls, {}).items():
                if prop not in style_parts:
                    style_parts[prop] = val
        if style_parts:
            para['kindle_style'] = '; '.join(f'{k}:{v}' for k, v in style_parts.items())
 
        self.paragraphs.append(para)
 
        logger.debug(f"[LEGACY] Processed element {element_id}: {len(reserve_elements)} reserve, "
                     f"{len(prefix_tags)} prefix tags, {len(suffix_tags)} suffix tags, "
                     f"non-translatable={is_non_translatable}")
 
    def _process_element_new(self, element, item_href, seen):
        tag_name = etree.QName(element).localname
 
        if element.get('id') is None:
            element.set('id', f"trans_{uuid.uuid4()}")
 
        element_id = element.get('id')
        element_copy = copy.deepcopy(element)
 
        for noise_tag in self.NOISE_TAGS:
            for noise_elem in element_copy.xpath(f'.//x:{noise_tag}', namespaces=NAMESPACES):
                parent = noise_elem.getparent()
                if parent is not None:
                    parent.remove(noise_elem)
 
        reserve_elements = []
        placeholder_pattern = '<id_{:02d}>'
        reserve_counter = 0
 
        for reserve_tag in self.RESERVE_TAGS:
            for reserve_elem in element_copy.xpath(f'.//x:{reserve_tag}', namespaces=NAMESPACES):
                reserve_elem_copy = copy.deepcopy(reserve_elem)
                reserve_elem_copy.tail = None
                reserve_html = etree.tostring(reserve_elem_copy, encoding='unicode', method='xml')
                reserve_elements.append(reserve_html)
 
                placeholder = placeholder_pattern.format(reserve_counter)
                reserve_counter += 1
 
                parent = reserve_elem.getparent()
                if parent is not None:
                    tail = reserve_elem.tail or ''
                    prev = reserve_elem.getprevious()
                    if prev is not None:
                        prev.tail = (prev.tail or '') + placeholder + tail
                    else:
                        parent.text = (parent.text or '') + placeholder + tail
                    parent.remove(reserve_elem)
 
        self._remove_useless_spans(element_copy)
 
        inline_formatting_map = {}
        inline_counter = reserve_counter
 
        INLINE_FORMATTING_TAGS = [
            'a', 'i', 'b', 'em', 'strong',
            'u', 'sup', 'sub', 'small', 'span'
        ]
 
        clean_text, inline_formatting_map, final_counter = \
            self._replace_inline_formatting_with_placeholders(
                element_copy,
                INLINE_FORMATTING_TAGS,
                inline_counter,
                placeholder_pattern
            )
 
        clean_text = self._cleanup_empty_placeholders(clean_text)
 
        clean_text, inline_formatting_map = \
            self._flatten_placeholder_nesting(clean_text, inline_formatting_map)
 
        clean_text, non_translatable_map = \
            self._extract_non_translatable_placeholders(
                clean_text,
                inline_formatting_map
            )
 
        clean_text = self._cleanup_empty_placeholders(clean_text)
 
        if not clean_text or not clean_text.strip():
            return
 
        prefix_tags, suffix_tags, clean_text = \
            self._extract_boundary_reserve_tags(clean_text)
 
        is_non_translatable = self._is_non_translatable_content(clean_text)
 
        if is_non_translatable:
            logger.debug(f"Marking element {element_id} as non-translatable")
 
        auto_wrap_tags = self._detect_auto_wrap_tags(
            clean_text,
            inline_formatting_map
        )
 
        if auto_wrap_tags:
            clean_text = self._strip_outer_placeholders(
                clean_text,
                auto_wrap_tags
            )
            logger.debug(f"Auto-wrap detected: {len(auto_wrap_tags)} tag(s)")
 
        used_ids = set(int(m) for m in re.findall(r'<p_(\d{2})>', clean_text))
        non_translatable_ids = set(non_translatable_map.keys())
 
        auto_wrap_ids = set(tag_info['elem_id'] for tag_info in auto_wrap_tags) if auto_wrap_tags else set()
        keep_ids = used_ids | non_translatable_ids | auto_wrap_ids
 
        inline_formatting_map = {
            k: v for k, v in inline_formatting_map.items()
            if k in keep_ids
        }
 
        logger.debug(
            f"Pruned inline_formatting_map: {len(inline_formatting_map)} entries "
            f"(used: {sorted(used_ids)}, nt: {sorted(non_translatable_ids)}, "
            f"auto_wrap: {sorted(auto_wrap_ids)})"
        )
 
        key = (item_href, element_id)
        if key in seen:
            return
        seen.add(key)
 
        original_html = etree.tostring(
            element,
            encoding='unicode',
            method='xml',
            pretty_print=False
        )
 
        HEADING_LOCAL_NAMES = {'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}
        element_type = 'heading' if tag_name in HEADING_LOCAL_NAMES else tag_name
 
        para = {
            "id": element_id,
            "original_text": clean_text,
            "translated_text": "",
            "is_translated": False,
            "item_href": item_href,
            "element_type": element_type,
            "original_html": original_html,
            "has_mismatch": False,
            "reserve_elements": reserve_elements,
            "inline_formatting_map": inline_formatting_map,
            "non_translatable_placeholders": non_translatable_map,
            "placeholder_pattern": placeholder_pattern,
            "processing_mode": "inline",
            "prefix_reserve_tags": prefix_tags,
            "suffix_reserve_tags": suffix_tags,
            "is_non_translatable": is_non_translatable
        }
 
        style_parts = {}
        elem_align = (element.get('align') or '').strip().lower()
        if elem_align in ('center', 'left', 'right', 'justify'):
            style_parts['text-align'] = elem_align
        elem_style_raw = (element.get('style') or '').strip()
        for decl in elem_style_raw.split(';'):
            decl = decl.strip()
            if not decl or ':' not in decl:
                continue
            prop, _, val = decl.partition(':')
            prop = prop.strip().lower()
            if prop in {'text-align', 'color', 'font-size', 'font-weight', 'font-style',
                        'text-indent', 'margin-left', 'margin-right', 'background-color',
                        'letter-spacing', 'line-height'}:
                style_parts[prop] = val.strip()
        for cls in (element.get('class') or '').split():
            for prop, val in self.css_class_styles.get(cls, {}).items():
                if prop not in style_parts:
                    style_parts[prop] = val
        if style_parts:
            para['kindle_style'] = '; '.join(f'{k}:{v}' for k, v in style_parts.items())
 
        if auto_wrap_tags:
            para['auto_wrap_tags'] = auto_wrap_tags
 
        self.paragraphs.append(para)
 
    def _remove_useless_spans(self, element):
        spans_to_remove = []
 
        for span in element.xpath('.//x:span', namespaces=NAMESPACES):
            has_class = span.get('class') is not None
            has_style = span.get('style') is not None
            has_id = span.get('id') is not None
            has_lang = span.get('lang') is not None
            has_dir = span.get('dir') is not None
            has_epub_type = span.get(f'{{{NAMESPACES["epub"]}}}type') is not None
 
            if has_class or has_style or has_id or has_lang or has_dir or has_epub_type:
                continue
 
            text = (span.text or '')
            children = list(span)
 
            if not text.strip() and len(children) <= 1:
                spans_to_remove.append(span)
 
        removed_count = 0
        for span in reversed(spans_to_remove):
            parent = span.getparent()
            if parent is None:
                continue
 
            span_index = list(parent).index(span)
 
            if span.text:
                prev = span.getprevious()
                if prev is not None:
                    prev.tail = (prev.tail or '') + span.text
                else:
                    parent.text = (parent.text or '') + span.text
 
            for child in reversed(list(span)):
                parent.insert(span_index, child)
 
            if span.tail:
                new_prev = span.getprevious()
                if new_prev is not None:
                    new_prev.tail = (new_prev.tail or '') + span.tail
                else:
                    parent.text = (parent.text or '') + span.tail
 
            parent.remove(span)
            removed_count += 1
 
        if removed_count > 0:
            logger.info(f"✓ Removed {removed_count} useless <span> wrapper(s)")
 
    def _replace_inline_formatting_with_placeholders(self, element, inline_tags, start_counter, placeholder_pattern):
        formatting_map = {}
        counter = start_counter
 
        STRUCTURAL_SPAN_CLASSES = {
            'first-letter',
            'last-word',
            'item-number',
            'element-number',
        }
 
        replacements = []
 
        tags_to_process = []
        for tag in inline_tags:
            if self.skip_inline_tags.get(tag, False):
                logger.debug(f"Skipping <{tag}> tags (user preference)")
            else:
                tags_to_process.append(tag)
 
        if not tags_to_process:
            logger.warning("All inline formatting tags are skipped - no placeholders will be created")
 
        for tag in tags_to_process:
            for elem in element.xpath(f'.//x:{tag}', namespaces=NAMESPACES):
                if tag == 'span':
                    epub_type = elem.get(f'{{{NAMESPACES["epub"]}}}type')
                    if epub_type == 'pagebreak':
                        logger.debug(f"Skipping pagebreak span")
                        continue
 
                    class_attr = elem.get('class')
                    if class_attr in STRUCTURAL_SPAN_CLASSES:
                        logger.debug(f"Skipping structural span: class='{class_attr}'")
                        continue
 
                    if not elem.attrib:
                        text = (elem.text or '').strip()
                        if len(text) == 1 and text.isalpha():
                            parent = elem.getparent()
                            if parent is not None:
                                children = [c for c in parent if isinstance(c, etree._Element)]
                                if len(children) > 0 and children[0] == elem:
                                    logger.debug(f"Skipping drop cap span: '{text}'")
                                    continue
 
                elem_id = counter
                opening = "<p_{:02d}>".format(elem_id)
                closing = "</p_{:02d}>".format(elem_id)
 
                tag_name = etree.QName(elem).localname
                attributes = dict(elem.attrib)
 
                prev_sibling = elem.getprevious()
                if prev_sibling is not None:
                    preceding_text = prev_sibling.tail or ''
                else:
                    parent_elem = elem.getparent()
                    preceding_text = (parent_elem.text or '') if parent_elem is not None else ''
                has_leading_space = bool(preceding_text) and preceding_text[-1] in ' \t\n\r\u00a0'
 
                elem_tail = elem.tail or ''
                has_trailing_space = bool(elem_tail) and elem_tail[0] in ' \t\n\r\u00a0'
 
                logger.debug(
                    f"  elem_id={elem_id} <{tag_name}> "
                    f"has_leading_space={has_leading_space}, has_trailing_space={has_trailing_space}"
                )
 
                formatting_map[elem_id] = {
                    'tag': tag_name,
                    'attributes': attributes,
                    'opening_placeholder': opening,
                    'closing_placeholder': closing,
                    'has_leading_space': has_leading_space,
                    'has_trailing_space': has_trailing_space,
                }
 
                replacements.append((elem, opening, closing, elem_id))
                counter += 1
 
        result_text = self._serialize_element_with_placeholders(element, replacements)
 
        return result_text, formatting_map, counter
 
    def _serialize_element_with_placeholders(self, element, replacements):
        replace_map = {id(elem): (opening, closing) for elem, opening, closing, _ in replacements}
 
        def process_node(node, depth=0):
            parts = []
 
            node_id = id(node)
            if node_id in replace_map:
                opening, closing = replace_map[node_id]
 
                parts.append(opening)
 
                if node.text:
                    parts.append(node.text)
 
                for child in node:
                    parts.append(process_node(child, depth + 1))
                    if child.tail:
                        parts.append(child.tail)
 
                parts.append(closing)
 
                return ''.join(parts)
 
            else:
                if node.text:
                    parts.append(node.text)
 
                for child in node:
                    parts.append(process_node(child, depth + 1))
                    if child.tail:
                        parts.append(child.tail)
 
                return ''.join(parts)
 
        return process_node(element)
 
    def _cleanup_empty_placeholders(self, text):
        original_text = text
        max_iterations = 10
 
        for iteration in range(max_iterations):
            pattern = r'<p_(\d{2})></p_\1>'
            new_text = re.sub(pattern, '', text)
 
            if new_text == text:
                break
 
            text = new_text
 
        if text != original_text:
            original_count = original_text.count('<p_')
            new_count = text.count('<p_')
            removed = (original_count - new_count) // 2
            logger.info(f"✓ Removed {removed} empty placeholder pair(s)")
 
        return text
 
    def _flatten_placeholder_nesting(self, text, formatting_map):
        pattern = r'<p_(\d{2})>\s*(<p_\d{2}>.*?</p_\d{2}>)\s*</p_\1>'
 
        max_iterations = 5
        removed_ids = set()
 
        for iteration in range(max_iterations):
            matches = list(re.finditer(pattern, text, re.DOTALL))
 
            if not matches:
                break
 
            for match in reversed(matches):
                outer_id_str = match.group(1)
                inner_content = match.group(2)
 
                outer_id = int(outer_id_str)
 
                if outer_id not in formatting_map:
                    continue
 
                outer_info = formatting_map[outer_id]
                outer_tag = outer_info['tag']
                outer_attrs = outer_info['attributes']
 
                should_remove = False
 
                if outer_tag == 'span' and not outer_attrs:
                    should_remove = True
                    logger.debug(f"Flattening: <p_{outer_id_str}> is empty <span> → REMOVE")
 
                inner_match = re.match(r'<p_(\d{2})>', inner_content)
                if inner_match and not should_remove:
                    inner_id = int(inner_match.group(1))
 
                    if inner_id in formatting_map:
                        inner_info = formatting_map[inner_id]
 
                        if (outer_tag == inner_info['tag'] and
                            outer_attrs == inner_info['attributes']):
                            should_remove = True
                            logger.debug(f"Flattening: <p_{outer_id_str}> duplicates <p_{inner_id}> → REMOVE")
 
                if should_remove:
                    text = text[:match.start()] + inner_content + text[match.end():]
                    removed_ids.add(outer_id)
 
        for tag_id in removed_ids:
            if tag_id in formatting_map:
                del formatting_map[tag_id]
                logger.debug(f"✓ Removed <p_{tag_id:02d}> from formatting_map")
 
        if removed_ids:
            logger.info(f"✓ Flattened {len(removed_ids)} nested placeholder(s)")
 
        return text, formatting_map
 
    def _extract_non_translatable_placeholders(self, text, formatting_map):
        NON_TRANSLATABLE_PATTERN = r'^[\s\.,!?:;…]*$'
        placeholder_pattern = r'<p_(\d{2})>(.*?)</p_\1>'
 
        non_translatable_map = {}
 
        def replace_with_marker(match):
            tag_id = int(match.group(1))
            content = match.group(2)
 
            if re.match(NON_TRANSLATABLE_PATTERN, content):
                non_translatable_map[tag_id] = {
                    'full_match': match.group(0),
                    'content': content
                }
                logger.debug(f"Marked non-translatable: {repr(match.group(0))}")
                return f'<nt_{tag_id:02d}/>'
            else:
                return match.group(0)
 
        clean_text = re.sub(placeholder_pattern, replace_with_marker, text, flags=re.DOTALL)
 
        if non_translatable_map:
            logger.info(f"✓ Marked {len(non_translatable_map)} non-translatable placeholder(s) with markers")
            for tag_id, info in non_translatable_map.items():
                logger.debug(f"  p_{tag_id:02d} → <nt_{tag_id:02d}/>: {repr(info['content'])}")
 
        return clean_text, non_translatable_map
 
    def _detect_auto_wrap_tags(self, text, formatting_map):
        if not formatting_map:
            return None
 
        wrap_tags = []
        working_text = text.strip()
 
        while True:
            match = re.match(r'^<p_(\d{2})>(.*)</p_\1>$', working_text, re.DOTALL)
 
            if not match:
                break
 
            elem_id_str = match.group(1)
            elem_id = int(elem_id_str)
            inner_text = match.group(2)
 
            if elem_id not in formatting_map:
                break
 
            info = formatting_map[elem_id]
            wrap_tags.append({
                'elem_id': elem_id,
                'opening': info['opening_placeholder'],
                'closing': info['closing_placeholder'],
                'tag': info['tag'],
                'attributes': info['attributes']
            })
 
            working_text = inner_text.strip()
 
        if wrap_tags and not re.search(r'</?p_\d{2}>', working_text):
            logger.debug(f"✓ Auto-wrap detected: {len(wrap_tags)} tag(s)")
            for idx, tag_info in enumerate(wrap_tags):
                logger.debug(f"  [{idx}] <{tag_info['tag']}> (id={tag_info['elem_id']})")
            return wrap_tags
 
        return None
 
    def _strip_outer_placeholders(self, text, auto_wrap_tags):
        working_text = text.strip()
 
        for tag_info in auto_wrap_tags:
            opening = tag_info['opening']
            closing = tag_info['closing']
 
            if working_text.startswith(opening) and working_text.endswith(closing):
                working_text = working_text[len(opening):-len(closing)].strip()
                logger.debug(f"  Stripped {opening}...{closing}")
 
        return working_text
 
    def _extract_boundary_reserve_tags(self, text):
        prefix_tags = []
        suffix_tags = []
        clean_text = text
 
        tag_pattern = r'<id_\d{2}>'
 
        while True:
            clean_text = clean_text.lstrip()
 
            match = re.match(tag_pattern, clean_text)
            if match:
                tag = match.group(0)
                prefix_tags.append(tag)
                clean_text = clean_text[len(tag):]
            else:
                break
 
        while True:
            clean_text = clean_text.rstrip()
 
            match = re.search(tag_pattern + r'$', clean_text)
            if match:
                tag = match.group(0)
                suffix_tags.insert(0, tag)
                clean_text = clean_text[:-len(tag)]
            else:
                break
 
        clean_text = clean_text.strip()
 
        if prefix_tags or suffix_tags:
            logger.debug(f"Extracted boundary tags: prefix={prefix_tags}, suffix={suffix_tags}")
 
        return prefix_tags, suffix_tags, clean_text

class SRTProcessor(FileProcessor):
    def get_file_type(self) -> str:
        return "srt"

    def load(self, path: str) -> Tuple[List[Dict], None]:
        encodings_to_try = ['utf-8', 'utf-8-sig', 'windows-1250', 'iso-8859-2', 'cp1252', 'latin1']

        content = None
        used_encoding = None

        for encoding in encodings_to_try:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    content = f.read()
                used_encoding = encoding
                logger.info(f"Successfully loaded SRT file with encoding: {encoding}")
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                logger.warning(f"Error loading SRT with {encoding}: {e}")
                continue

        if content is None:
            raise ValueError(f"Failed to load SRT file. Tried encodings: {', '.join(encodings_to_try)}")

        try:
            blocks = [block.strip() for block in content.split('\n\n') if block.strip()]
            paragraphs = []

            for block in blocks:
                lines = block.split('\n')
                if len(lines) < 3:
                    continue

                block_number = lines[0].strip()
                timestamp = lines[1].strip()
                text_lines = lines[2:]

                original_lines_with_tags = text_lines

                srt_tags_by_line = []
                clean_lines = []

                for line_idx, line_text in enumerate(text_lines):
                    line_srt_tags = {}
                    clean_line = line_text

                    for match in re.finditer(r'<(i|b|u|font[^>]*)>', line_text):
                        tag = match.group(0)
                        pos = len(re.sub(r'<[^>]+>', '', line_text[:match.start()]))
                        if pos not in line_srt_tags:
                            line_srt_tags[pos] = []
                        line_srt_tags[pos].append(('open', tag))

                    for match in re.finditer(r'</(i|b|u|font)>', line_text):
                        tag = match.group(0)
                        pos = len(re.sub(r'<[^>]+>', '', line_text[:match.start()]))
                        if pos not in line_srt_tags:
                            line_srt_tags[pos] = []
                        line_srt_tags[pos].append(('close', tag))

                    clean_line = re.sub(r'<[^>]+>', '', line_text)
                    clean_lines.append(clean_line)
                    srt_tags_by_line.append(line_srt_tags)

                combined_text = '\n'.join(clean_lines)

                split_positions = []
                current_pos = 0
                for i, line in enumerate(clean_lines):
                    if i < len(clean_lines) - 1:
                        current_pos += len(line)
                        split_positions.append(current_pos)
                        current_pos += 1

                paragraphs.append({
                    'id': block_number,
                    'original_text': combined_text,
                    'translated_text': '',
                    'is_translated': False,
                    'item_href': path,
                    'element_type': 'subtitle_block',
                    'timestamp': timestamp,
                    'subtitle_block': block_number,
                    'has_mismatch': False,
                    'srt_tags_by_line': srt_tags_by_line,
                    'original_clean_lines': clean_lines,
                    'original_line_count': len(clean_lines),
                    'original_split_positions': split_positions,
                    'original_lines_with_tags': original_lines_with_tags
                })

            if not paragraphs:
                raise ValueError("No valid subtitle blocks found in SRT file")

            logger.info(f"Loaded {len(paragraphs)} subtitle blocks from SRT (Encoding: {used_encoding})")

            return paragraphs, None

        except Exception as e:
            logger.error(f"Failed to parse SRT file: {e}", exc_info=True)
            raise

    def _extract_srt_tags(self, line: str) -> Tuple[str, Dict[int, List[Tuple[str, str]]]]:
        tags_dict = {}
        clean_line = line

        tag_pattern = r'<(/?)([biu]|font[^>]*)>'

        matches = list(re.finditer(tag_pattern, line))

        if not matches:
            return line, {}

        offset = 0

        for match in matches:
            is_closing = match.group(1) == '/'
            tag_content = match.group(2)

            if tag_content == 'b':
                tag_type = 'bold'
                tag_value = '</b>' if is_closing else '<b>'
            elif tag_content == 'i':
                tag_type = 'italic'
                tag_value = '</i>' if is_closing else '<i>'
            elif tag_content == 'u':
                tag_type = 'underline'
                tag_value = '</u>' if is_closing else '<u>'
            elif tag_content.startswith('font'):
                tag_type = 'font'
                tag_value = f'</{tag_content}>' if is_closing else f'<{tag_content}>'
            else:
                continue

            pos = match.start() - offset

            if pos not in tags_dict:
                tags_dict[pos] = []
            tags_dict[pos].append((tag_type, tag_value))

            clean_line = clean_line[:match.start() - offset] + clean_line[match.end() - offset:]
            offset += len(match.group(0))

        return clean_line, tags_dict

class TXTProcessor(FileProcessor):
    MAX_CHARS_PER_FRAGMENT = 6000

    def get_file_type(self) -> str:
        return "txt"

    def load(self, path: str) -> Tuple[List[Dict], None]:
        encodings_to_try = ['utf-8', 'utf-8-sig', 'windows-1250', 'iso-8859-2', 'cp1252', 'latin1']

        content = None
        used_encoding = None

        for encoding in encodings_to_try:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    content = f.read()
                used_encoding = encoding
                logger.info(f"Successfully loaded TXT file with encoding: {encoding}")
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                logger.warning(f"Error loading TXT with {encoding}: {e}")
                continue

        if content is None:
            raise ValueError(f"Failed to load TXT file. Tried encodings: {', '.join(encodings_to_try)}")

        try:
            paragraphs = []

            if self._is_subtitle_txt(content):
                logger.info("TXT mode: SUBTITLE")
                paragraphs = self._load_subtitle_txt(content, path)
            else:
                fragment_id = 1

                paragraphs_raw = content.split('\n\n')
                paragraphs_raw = [p.strip() for p in paragraphs_raw if p.strip()]

                if len(paragraphs_raw) >= 2:
                    logger.info(f"TXT mode: PARAGRAPHS ({len(paragraphs_raw)} blocks)")

                    for para_index, text in enumerate(paragraphs_raw, start=1):
                        if len(text) <= self.MAX_CHARS_PER_FRAGMENT:
                            fragments = [text]
                        else:
                            fragments = self._split_long_text(text)

                        for part_index, part in enumerate(fragments, start=1):
                            paragraphs.append({
                                'id': str(fragment_id),
                                'original_text': part,
                                'translated_text': '',
                                'is_translated': False,
                                'item_href': path,
                                'element_type': 'paragraph_part',
                                'paragraph_number': para_index,
                                'has_mismatch': False,
                                'part_index': part_index,
                                'parts_total': len(fragments)
                            })
                            fragment_id += 1

                else:
                    lines = content.split('\n')
                    lines = [l.strip() for l in lines if l.strip()]

                    if len(lines) >= 2:
                        logger.info(f"TXT mode: LINES ({len(lines)} lines)")

                        for line_index, text in enumerate(lines, start=1):
                            if len(text) <= self.MAX_CHARS_PER_FRAGMENT:
                                fragments = [text]
                            else:
                                fragments = self._split_long_text(text)

                            for part_index, part in enumerate(fragments, start=1):
                                paragraphs.append({
                                    'id': str(fragment_id),
                                    'original_text': part,
                                    'translated_text': '',
                                    'is_translated': False,
                                    'item_href': path,
                                    'element_type': 'line_part',
                                    'paragraph_number': line_index,
                                    'has_mismatch': False,
                                    'part_index': part_index,
                                    'parts_total': len(fragments)
                                })
                                fragment_id += 1

                    else:
                        logger.info("TXT mode: SENTENCES")
                        sentences = self._split_into_sentences(content)

                        for sent_index, text in enumerate(sentences, start=1):
                            text = text.strip()
                            if not text:
                                continue

                            if len(text) <= self.MAX_CHARS_PER_FRAGMENT:
                                fragments = [text]
                            else:
                                fragments = self._split_long_text(text)

                            for part_index, part in enumerate(fragments, start=1):
                                paragraphs.append({
                                    'id': str(fragment_id),
                                    'original_text': part,
                                    'translated_text': '',
                                    'is_translated': False,
                                    'item_href': path,
                                    'element_type': 'sentence_part',
                                    'paragraph_number': sent_index,
                                    'has_mismatch': False,
                                    'part_index': part_index,
                                    'parts_total': len(fragments)
                                })
                                fragment_id += 1

            if not paragraphs:
                logger.warning("TXT mode: FALLBACK (single block)")
                paragraphs.append({
                    'id': '1',
                    'original_text': content.strip(),
                    'translated_text': '',
                    'is_translated': False,
                    'item_href': path,
                    'element_type': 'full_text',
                    'paragraph_number': 1,
                    'has_mismatch': False,
                    'part_index': 1,
                    'parts_total': 1
                })

            logger.info(f"Loaded {len(paragraphs)} fragments (Encoding: {used_encoding})")
            return paragraphs, None

        except Exception as e:
            logger.error(f"Failed to parse TXT file: {e}", exc_info=True)
            raise

    def _is_subtitle_txt(self, content: str) -> bool:
        subtitle_pattern = re.compile(r'^\[\d+\]\[\d+\]')
        lines = [l.strip() for l in content.split('\n') if l.strip()]
        if not lines:
            return False
        matched = sum(1 for l in lines if subtitle_pattern.match(l))
        return matched >= max(1, int(len(lines) * 0.8))

    def _load_subtitle_txt(self, content: str, path: str) -> List[Dict]:
        paragraphs = []
        fragment_id = 1
        block_id = 0
        subtitle_pattern = re.compile(r'^\[(\d+)\]\[(\d+)\](.*)')

        for raw_line in content.split('\n'):
            raw_line_stripped = raw_line.strip()
            if not raw_line_stripped:
                continue
            match = subtitle_pattern.match(raw_line_stripped)
            if not match:
                continue

            start = int(match.group(1))
            end = int(match.group(2))
            text_part = match.group(3)

            raw_parts = text_part.split('|') if '|' in text_part else [text_part]
            parts = [p.strip() for p in raw_parts if p.strip()]
            lines_total = len(parts)

            if not parts:
                block_id += 1
                continue

            for line_index, part in enumerate(parts):
                paragraphs.append({
                    'id': str(fragment_id),
                    'original_text': part,
                    'translated_text': '',
                    'is_translated': False,
                    'item_href': path,
                    'element_type': 'subtitle_line',
                    'has_mismatch': False,
                    'txt_subtitle_start': start,
                    'txt_subtitle_end': end,
                    'txt_subtitle_block_id': block_id,
                    'txt_subtitle_line_index': line_index,
                    'txt_subtitle_lines_total': lines_total,
                })
                fragment_id += 1

            block_id += 1

        return paragraphs

    def _split_into_sentences(self, text: str) -> List[str]:
        ABBREVIATIONS = {

            'np', 'itd', 'itp', 'tzw', 'ok', 'ul', 'al', 'pl', 'wg', 'vs',
            'nr', 'str', 'mgr', 'dr', 'prof', 'inż', 'lic', 'hab',
            'mln', 'mld', 'tys', 'godz', 'min', 'sek', 'zł', 'gr',
            'gen', 'ppłk', 'płk', 'mjr', 'kpt', 'por', 'szer',

            'mr', 'mrs', 'ms', 'dr', 'prof', 'sr', 'jr', 'rev', 'gen',
            'sgt', 'cpl', 'pvt', 'capt', 'col', 'lt', 'gov', 'pres',
            'dept', 'est', 'approx', 'avg', 'max', 'min', 'vol', 'no',
            'fig', 'eq', 'vs', 'etc', 'inc', 'corp', 'co', 'ltd',
            'jan', 'feb', 'mar', 'apr', 'jun', 'jul', 'aug', 'sep',
            'oct', 'nov', 'dec',

            'bzw', 'ca', 'dh', 'evtl', 'ggf', 'usw', 'vgl', 'zb', 'zbsp',
            'str', 'nr', 'hr', 'fr', 'dr', 'prof', 'dipl',

            'mme', 'mlle', 'mm', 'dr', 'pr', 'me', 'ste', 'st',
            'bd', 'av', 'pl', 'sq', 'env', 'nb', 'cf', 'ex',

            'sr', 'sra', 'srta', 'dr', 'dra', 'prof', 'lic',
            'num', 'pag', 'cap', 'art', 'fig', 'ed',

            'sig', 'dott', 'prof', 'ing', 'avv', 'geom', 'arch',

            'al', 'et', 'ibid', 'op', 'loc', 'cit', 'viz',
        }

        raw_parts = re.split(r'([.!?…]+)', text)

        sentences = []
        current = ''

        i = 0
        while i < len(raw_parts):
            part = raw_parts[i]

            if i % 2 == 0:
                current += part
                i += 1
            else:
                separator = part

                next_part = raw_parts[i + 1] if i + 1 < len(raw_parts) else ''

                if separator == '.' and next_part:
                    prev_word_match = re.search(r'(\w+)$', current, re.UNICODE)

                    if prev_word_match:
                        prev_word = prev_word_match.group(1).lower()

                        if prev_word in ABBREVIATIONS:
                            current += separator
                            i += 1
                            continue

                        if len(prev_word) == 1:
                            current += separator
                            i += 1
                            continue

                        if prev_word.isdigit():
                            next_stripped = next_part.lstrip()
                            first_char = next_stripped[0] if next_stripped else ''
                            if first_char and not first_char.isupper():
                                current += separator
                                i += 1
                                continue

                if next_part and not next_part.startswith((' ', '\n', '\t')):
                    current += separator
                    i += 1
                    continue

                current += separator
                sentence = current.strip()
                if sentence:
                    sentences.append(sentence)
                current = ''
                i += 1

        if current.strip():
            sentences.append(current.strip())

        if len(sentences) < 2:
            logger.warning("Sentence splitter: fallback to simple split")
            sentences = [s.strip() for s in re.split(r'(?<=[.!?…])\s+', text) if s.strip()]

        logger.debug(f"Split into {len(sentences)} sentences")
        return sentences

    def _split_long_text(self, text: str) -> List[str]:
        chunks = []
        current_chunk = ''

        sentences = self._split_into_sentences(text)

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            if len(sentence) <= self.MAX_CHARS_PER_FRAGMENT:
                separator = ' ' if current_chunk else ''
                if len(current_chunk) + len(separator) + len(sentence) <= self.MAX_CHARS_PER_FRAGMENT:
                    current_chunk += separator + sentence
                else:
                    if current_chunk:
                        chunks.append(current_chunk)
                    current_chunk = sentence
                continue

            if current_chunk:
                chunks.append(current_chunk)
                current_chunk = ''

            words = sentence.split()

            for word in words:
                if len(word) > self.MAX_CHARS_PER_FRAGMENT:
                    if current_chunk:
                        chunks.append(current_chunk)
                        current_chunk = ''
                    for i in range(0, len(word), self.MAX_CHARS_PER_FRAGMENT):
                        chunks.append(word[i:i + self.MAX_CHARS_PER_FRAGMENT])
                    continue

                separator = ' ' if current_chunk else ''
                if len(current_chunk) + len(separator) + len(word) <= self.MAX_CHARS_PER_FRAGMENT:
                    current_chunk += separator + word
                else:
                    if current_chunk:
                        chunks.append(current_chunk)
                    current_chunk = word

        if current_chunk:
            chunks.append(current_chunk)

        return chunks

class PDFProcessor(FileProcessor):
    def __init__(self, app_settings: dict):
        self.app_settings = app_settings
        self.original_pdf_path = None
        self.image_items = []

    def get_file_type(self) -> str:
        return "pdf"

    def load(self, path: str) -> Tuple[List[Dict], str]:
        try:
            import fitz

            self.original_pdf_path = path
            doc = fitz.open(path)
            paragraphs = []
            fragment_id = 0

            LIST_MARKER_RE = re.compile(r'^\s*(?:[-•–—*]|\d+[.):])\s+')

            for page_number, page in enumerate(doc):
                page_dict = page.get_text("dict", sort=True)
                page_height = page.rect.height

                valid_blocks = []
                for block in page_dict.get("blocks", []):
                    if block.get("type") != 0:
                        continue
                    lines = block.get("lines", [])
                    if not lines:
                        continue

                    merged_text = ""
                    first_span = None
                    for line in lines:
                        for span in line.get("spans", []):
                            if first_span is None:
                                first_span = span
                            merged_text += span.get("text", "").replace('\xad', '')
                        merged_text += " "

                    merged_text = merged_text.strip()

                    if not merged_text:
                        continue
                    if len(merged_text) < 15 and re.fullmatch(r'[\d\s.,\-]+', merged_text):
                        continue
                    if len(merged_text) <= 1:
                        continue
                    if first_span is None:
                        continue

                    font_size = first_span.get("size", 12.0)
                    flags = first_span.get("flags", 0)
                    color = first_span.get("color", 0)
                    bbox = list(block.get("bbox", [0, 0, 0, 0]))

                    valid_blocks.append({
                        'text': merged_text,
                        'bbox': bbox,
                        'font_size': font_size,
                        'bold': bool(flags & (1 << 4)),
                        'italic': bool(flags & (1 << 1)),
                        'color': color,
                    })

                if valid_blocks:
                    font_sizes = [b['font_size'] for b in valid_blocks]
                    try:
                        median_font_size = statistics.median(font_sizes)
                    except Exception:
                        median_font_size = 12.0
                else:
                    median_font_size = 12.0

                for para_number, block_data in enumerate(valid_blocks):
                    font_size = block_data['font_size']
                    bold = block_data['bold']
                    bbox = block_data['bbox']
                    y0 = bbox[1]
                    text = block_data['text']

                    is_heading = (
                        font_size > median_font_size * 1.35
                        or (bold and font_size >= median_font_size * 1.15)
                        or (text.rstrip().isupper() and 3 < len(text.strip()) < 60)
                    )

                    if is_heading:
                        element_type = "heading"
                    elif y0 < page_height * 0.12:
                        element_type = "header"
                    elif y0 > page_height * 0.88:
                        element_type = "footer"
                    elif LIST_MARKER_RE.match(text):
                        element_type = "list_item"
                    else:
                        element_type = "paragraph"

                    paragraphs.append({
                        'id': str(fragment_id),
                        'original_text': text,
                        'translated_text': '',
                        'is_translated': False,
                        'item_href': str(page_number),
                        'element_type': element_type,
                        'paragraph_number': para_number,
                        'has_mismatch': False,
                        'bbox': bbox,
                        'span_metadata': {
                            'font_size': font_size,
                            'bold': bold,
                            'italic': block_data['italic'],
                            'color': block_data['color'],
                        },
                    })
                    fragment_id += 1

            doc.close()
            paragraphs = self._merge_split_sentences(paragraphs)
            logger.info(f"PDFProcessor: loaded {len(paragraphs)} fragments from {path}")
            return paragraphs, path

        except Exception as e:
            logger.error(f"PDF load error: {e}", exc_info=True)
            raise

    @staticmethod
    def _ends_sentence(text: str) -> bool:
        stripped = text.rstrip()
        if not stripped:
            return True
        last_char = stripped[-1]
        closing_chars = ('"', "'", '\u201d', '\u2019', '\u00bb', '\u203a', ')', ']')
        if last_char in closing_chars:
            return len(stripped) >= 2 and stripped[-2] in '.!?\u2026'
        return last_char in '.!?\u2026'

    @staticmethod
    def _starts_continuation(text: str) -> bool:
        stripped = text.lstrip()
        if not stripped:
            return False
        return stripped[0].islower()

    def _merge_split_sentences(self, paragraphs: List[Dict]) -> List[Dict]:
        MAX_MERGED_LEN = 1500
        LIST_START_RE = re.compile(r'^\s*(?:[-•–—*]|\d+[.):])\s+')

        pages: Dict[str, List[Dict]] = {}
        page_order: List[str] = []
        for para in paragraphs:
            href = para.get('item_href', '')
            if href not in pages:
                pages[href] = []
                page_order.append(href)
            pages[href].append(para)

        result: List[Dict] = []
        total_merged = 0

        for href in page_order:
            page_paras = pages[href]
            merged_page: List[Dict] = []
            i = 0
            while i < len(page_paras):
                current = dict(page_paras[i])

                if (
                    current['original_text'].endswith('-')
                    and i + 1 < len(page_paras)
                    and current.get('element_type') == 'paragraph'
                    and page_paras[i + 1].get('element_type') == 'paragraph'
                ):
                    nxt = page_paras[i + 1]
                    nxt_text = nxt['original_text']
                    if nxt_text and nxt_text[0].islower():
                        curr_bbox = current.get('bbox', [0, 0, 0, 0])
                        nxt_bbox = nxt.get('bbox', [0, 0, 0, 0])
                        current['original_text'] = current['original_text'][:-1] + nxt_text
                        current['bbox'] = [
                            min(curr_bbox[0], nxt_bbox[0]),
                            min(curr_bbox[1], nxt_bbox[1]),
                            max(curr_bbox[2], nxt_bbox[2]),
                            max(curr_bbox[3], nxt_bbox[3]),
                        ]
                        i += 1
                        total_merged += 1
                        merged_page.append(current)
                        i += 1
                        continue

                while i + 1 < len(page_paras):
                    nxt = page_paras[i + 1]
                    if current.get('element_type') != 'paragraph':
                        break
                    if nxt.get('element_type') != 'paragraph':
                        break
                    curr_text = current['original_text']
                    nxt_text = nxt['original_text']
                    if self._ends_sentence(curr_text):
                        break
                    if curr_text.rstrip().endswith(':'):
                        break
                    if LIST_START_RE.match(nxt_text):
                        break
                    if not self._starts_continuation(nxt_text):
                        break
                    combined = curr_text + ' ' + nxt_text
                    if len(combined) > MAX_MERGED_LEN:
                        break
                    curr_bbox = current.get('bbox', [0, 0, 0, 0])
                    nxt_bbox = nxt.get('bbox', [0, 0, 0, 0])
                    font_size = current.get('span_metadata', {}).get('font_size', 12.0)
                    y_gap = nxt_bbox[1] - curr_bbox[3]
                    if y_gap > font_size * 2.5:
                        break
                    current['original_text'] = combined
                    current['bbox'] = [
                        min(curr_bbox[0], nxt_bbox[0]),
                        min(curr_bbox[1], nxt_bbox[1]),
                        max(curr_bbox[2], nxt_bbox[2]),
                        max(curr_bbox[3], nxt_bbox[3]),
                    ]
                    i += 1
                    total_merged += 1
                merged_page.append(current)
                i += 1
            result.extend(merged_page)

        if total_merged:
            logger.info(f"PDFProcessor: merged {total_merged} split-sentence fragment(s)")
        return result


class MobiProcessor(FileProcessor):
    def __init__(self, app_settings: dict):
        self.app_settings = app_settings
        self.original_kindle_path = None
        self._preview_tempdir = None
        self.image_items = []

    def get_file_type(self) -> str:
        return "mobi"

    def load(self, path: str) -> Tuple[List[Dict], str]:
        try:
            import mobi

            self.original_kindle_path = path
            self.image_items = []

            if self._preview_tempdir is not None:
                shutil.rmtree(self._preview_tempdir, ignore_errors=True)
                self._preview_tempdir = None

            tempdir, filepath = mobi.extract(path)
            self._preview_tempdir = tempdir

            try:
                paragraphs = self._parse_html(filepath)
            except Exception:
                raise

            logger.info(f"MobiProcessor: loaded {len(paragraphs)} fragments, "
                        f"{len(self.image_items)} image items from {path}")
            return paragraphs, path

        except Exception as e:
            logger.error(f"Mobi load error: {e}", exc_info=True)
            raise

    def cleanup_preview_tempdir(self):
        if self._preview_tempdir is not None:
            shutil.rmtree(self._preview_tempdir, ignore_errors=True)
            self._preview_tempdir = None

    def _parse_html(self, html_path: str) -> List[Dict]:
        import glob as _glob
        paragraphs = []
        fragment_id = 0
        media_type_map = {
            '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
            '.png': 'image/png', '.gif': 'image/gif',
            '.bmp': 'image/bmp', '.webp': 'image/webp',
        }
        use_inline = self.app_settings.get('use_inline_formatting', True)
        placeholder_pattern = '<p_{:02d}>'
        INLINE_TAGS = {'b', 'strong', 'i', 'em', 'u', 'span', 'small', 'a', 'font'}
        HEADING_TAGS = {'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}
        BLOCK_TAGS = {
            'p', 'div', 'section', 'article', 'blockquote', 'header', 'footer',
            'main', 'ul', 'ol', 'table', 'nav', 'aside', 'figure', 'figcaption',
            'address', 'pre', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
        }
        LEAF_BLOCK_TAGS = {'p', 'div', 'blockquote', 'header', 'footer', 'section', 'article', 'pre'}
        KEEP_CSS_PROPS = {
            'text-align', 'color', 'font-size', 'font-weight', 'font-style',
            'text-indent', 'margin-left', 'margin-right', 'letter-spacing',
            'line-height', 'float', 'vertical-align',
        }
        KEEP_BLOCK_PROPS = {
            'text-align', 'color', 'font-size', 'font-weight',
            'font-style', 'text-indent', 'margin-left', 'margin-right',
            'background-color', 'letter-spacing', 'line-height',
        }

        with open(html_path, 'rb') as f:
            content = f.read()

        try:
            tree = etree.fromstring(content, parser=etree.HTMLParser())
        except Exception:
            return paragraphs

        body = tree.find('.//body')
        if body is None:
            return paragraphs

        try:
            raw_html_bytes = content[:8000]
            logger.warning(f"MobiProcessor: RAW book.html (first 8000 bytes):\n{raw_html_bytes.decode('utf-8', errors='replace')}")
        except Exception:
            pass

        html_dir = os.path.dirname(os.path.abspath(html_path))

        css_class_styles = {}
        try:
            css_sources = []
            style_count = 0
            link_count = 0
            css_file_count = 0

            for se in tree.findall('.//style'):
                if se.text and se.text.strip():
                    css_sources.append(se.text)
                    style_count += 1

            for link_elem in tree.findall('.//link'):
                rel = (link_elem.get('rel') or '').lower()
                ltype = (link_elem.get('type') or '').lower()
                if 'stylesheet' in rel or ltype == 'text/css':
                    href = (link_elem.get('href') or '').strip()
                    link_count += 1
                    if href and not href.startswith(('http://', 'https://', 'data:')):
                        css_path = os.path.normpath(os.path.join(html_dir, href))
                        if os.path.isfile(css_path):
                            try:
                                with open(css_path, 'r', encoding='utf-8', errors='ignore') as _f:
                                    css_sources.append(_f.read())
                                css_file_count += 1
                            except Exception:
                                pass

            search_dirs = [html_dir]
            parent_dir = os.path.dirname(html_dir)
            if parent_dir and parent_dir != html_dir:
                search_dirs.append(parent_dir)
            for search_root in search_dirs:
                for css_path in _glob.glob(
                    os.path.join(search_root, '**', '*.css'), recursive=True
                ):
                    try:
                        with open(css_path, 'r', encoding='utf-8', errors='ignore') as _f:
                            css_sources.append(_f.read())
                        css_file_count += 1
                    except Exception:
                        pass

            logger.warning(f"MobiProcessor: html_dir={html_dir}")
            logger.warning(f"MobiProcessor: search_dirs={search_dirs}")
            logger.warning(f"MobiProcessor: html_dir contents={sorted(os.listdir(html_dir)) if os.path.exists(html_dir) else 'NOT EXISTS'}")
            logger.warning(f"MobiProcessor: parent_dir contents={sorted(os.listdir(parent_dir)) if os.path.exists(parent_dir) else 'NOT EXISTS'}")
            logger.warning(f"MobiProcessor: found <style> tags: {style_count} | <link> stylesheet: {link_count} | .css files loaded: {css_file_count}")
            logger.warning(f"MobiProcessor: total css_sources before OPF: {len(css_sources)}")

            opf_path = os.path.join(html_dir, 'content.opf')
            if os.path.isfile(opf_path):
                try:
                    opf_tree = etree.parse(opf_path)
                    css_from_opf = []
                    for item in opf_tree.findall('.//opf:item', namespaces={'opf': 'http://www.idpf.org/2007/opf'}):
                        mt = item.get('media-type', '')
                        href = item.get('href', '')
                        if mt == 'text/css' and href:
                            css_path = os.path.normpath(os.path.join(html_dir, href))
                            css_from_opf.append((href, css_path))
                            if os.path.isfile(css_path):
                                try:
                                    with open(css_path, 'r', encoding='utf-8', errors='ignore') as _f:
                                        css_sources.append(_f.read())
                                    css_file_count += 1
                                except Exception:
                                    pass
                    logger.warning(f"MobiProcessor: OPF found {len(css_from_opf)} CSS items: {[h for h, p in css_from_opf]}")
                except Exception as e:
                    logger.warning(f"MobiProcessor: OPF parse error: {e}")
            else:
                logger.warning("MobiProcessor: content.opf NOT FOUND")

            logger.warning(f"MobiProcessor: total css_sources after OPF: {len(css_sources)}")

            for css_text in css_sources:
                for m in re.finditer(
                    r'\.(-?[\w][\w-]*)\s*(?:,[^{]*)?\{([^}]*)\}', css_text, re.DOTALL
                ):
                    cls_name = m.group(1)
                    decls = m.group(2)
                    found = {}
                    for decl in decls.split(';'):
                        decl = decl.strip()
                        if ':' not in decl:
                            continue
                        prop, _, val = decl.partition(':')
                        prop = prop.strip().lower()
                        val = re.sub(r'\s*!important', '', val).strip()
                        if prop in KEEP_CSS_PROPS and val:
                            found[prop] = val
                    if found:
                        merged = css_class_styles.get(cls_name, {})
                        merged.update(found)
                        css_class_styles[cls_name] = merged

            logger.warning(f"MobiProcessor: FINAL classes found: {len(css_class_styles)} → {list(css_class_styles.keys())[:30]}")
        except Exception as e:
            logger.warning(f"MobiProcessor: CSS loading failed: {e}")

        cover_image_added = False
        try:
            opf_candidates = (
                _glob.glob(os.path.join(html_dir, '*.opf'))
                + _glob.glob(os.path.join(os.path.dirname(html_dir), '*.opf'))
            )
            if opf_candidates:
                opf_tree = etree.parse(opf_candidates[0])
                opf_root = opf_tree.getroot()
                opf_dir = os.path.dirname(opf_candidates[0])
                cover_id = None
                for el in opf_root.iter():
                    local = (el.tag if isinstance(el.tag, str) else '').split('}')[-1].lower()
                    if local == 'meta' and el.get('name', '').lower() == 'cover':
                        cover_id = el.get('content', '')
                        break
                if cover_id:
                    for el in opf_root.iter():
                        local = (el.tag if isinstance(el.tag, str) else '').split('}')[-1].lower()
                        if local == 'item' and el.get('id') == cover_id:
                            href = el.get('href', '')
                            if href:
                                cover_path = os.path.join(opf_dir, href)
                                if not os.path.isfile(cover_path):
                                    cover_path = os.path.join(html_dir, href)
                                if os.path.isfile(cover_path):
                                    ext = os.path.splitext(cover_path)[1].lower()
                                    mt = media_type_map.get(ext, 'image/jpeg')
                                    with open(cover_path, 'rb') as _f:
                                        cover_content = _f.read()
                                    basename = os.path.basename(cover_path)
                                    self.image_items.append({
                                        'id': basename,
                                        'file_name': basename,
                                        'media_type': mt,
                                        'content': cover_content,
                                        'is_cover': True,
                                    })
                                    cover_image_added = True
                            break
        except Exception:
            pass

        if not cover_image_added:
            for search_dir in (html_dir, os.path.dirname(html_dir)):
                for pattern in ('cover.jpg', 'cover.jpeg', 'cover.png', 'cover.gif', 'cover.webp'):
                    for sub in ('', 'Images', 'images'):
                        candidate = (
                            os.path.join(search_dir, sub, pattern)
                            if sub else
                            os.path.join(search_dir, pattern)
                        )
                        if os.path.isfile(candidate):
                            ext = os.path.splitext(candidate)[1].lower()
                            mt = media_type_map.get(ext, 'image/jpeg')
                            try:
                                with open(candidate, 'rb') as _f:
                                    cover_content = _f.read()
                                basename = os.path.basename(candidate)
                                self.image_items.append({
                                    'id': basename,
                                    'file_name': basename,
                                    'media_type': mt,
                                    'content': cover_content,
                                    'is_cover': True,
                                })
                                cover_image_added = True
                            except Exception:
                                pass
                            break
                if cover_image_added:
                    break

        def _is_leaf_block(element):
            for child in element:
                ctag = (child.tag if isinstance(child.tag, str) else '').lower().split('}')[-1]
                if ctag in BLOCK_TAGS:
                    return False
            return True

        def _extract_block_style(elem) -> str:
            style_parts = {}
            class_attr = (elem.get('class') or '').strip()
            for cls in class_attr.split():
                for prop, val in css_class_styles.get(cls, {}).items():
                    if prop in KEEP_BLOCK_PROPS and prop not in style_parts:
                        style_parts[prop] = val
            align_attr = (elem.get('align') or '').strip().lower()
            if not align_attr:
                parent_elem = elem.getparent()
                if parent_elem is not None:
                    align_attr = (parent_elem.get('align') or '').strip().lower()
            if align_attr in ('center', 'left', 'right', 'justify'):
                style_parts['text-align'] = align_attr
            width_attr = (elem.get('width') or '').strip()
            if not width_attr:
                parent_elem = elem.getparent()
                if parent_elem is not None:
                    width_attr = (parent_elem.get('width') or '').strip()
            if width_attr and width_attr not in ('0pt', '0px', '0%', '0', '0em'):
                if re.match(r'^\d+(\.\d+)?(em|px|pt|%)$', width_attr):
                    if 'text-indent' not in style_parts:
                        style_parts['text-indent'] = width_attr
            raw = (elem.get('style') or '').strip()
            for decl in raw.split(';'):
                decl = decl.strip()
                if not decl or ':' not in decl:
                    continue
                prop, _, val = decl.partition(':')
                prop = prop.strip().lower()
                if prop in KEEP_BLOCK_PROPS:
                    style_parts[prop] = val.strip()
            return '; '.join(f'{k}:{v}' for k, v in style_parts.items())

        pending_anchor_ids = []

        for elem in body.iter():
            tag = elem.tag if isinstance(elem.tag, str) else ''
            tag = tag.lower().split('}')[-1]

            if tag == 'a':
                parent = elem.getparent()
                parent_tag = (
                    parent.tag if parent is not None and isinstance(parent.tag, str) else ''
                ).lower().split('}')[-1]
                if parent_tag not in {'p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}:
                    for attr in ('id', 'name'):
                        val = (elem.get(attr) or '').strip()
                        if val and val not in pending_anchor_ids:
                            pending_anchor_ids.append(val)
                continue

            if tag == 'img':
                src = elem.get('src', '')
                if src and not src.startswith('data:'):
                    img_path = os.path.join(html_dir, src)
                    if not os.path.isfile(img_path):
                        alt_paths = [
                            os.path.join(html_dir, 'images', os.path.basename(src)),
                            os.path.join(html_dir, 'Images', os.path.basename(src)),
                        ]
                        for ap in alt_paths:
                            if os.path.isfile(ap):
                                img_path = ap
                                src = os.path.relpath(img_path, html_dir).replace('\\', '/')
                                break
                    if os.path.isfile(img_path):
                        ext = os.path.splitext(src)[1].lower()
                        mt = media_type_map.get(ext, 'image/jpeg')
                        try:
                            with open(img_path, 'rb') as f:
                                img_content = f.read()
                            basename = os.path.basename(src)
                            img_entry = {
                                'id': basename,
                                'file_name': basename,
                                'media_type': mt,
                                'content': img_content,
                            }
                            if not cover_image_added:
                                img_entry['is_cover'] = True
                                cover_image_added = True
                            self.image_items.append(img_entry)
                            img_para = {
                                'id': str(fragment_id),
                                'original_text': '',
                                'translated_text': '',
                                'is_translated': False,
                                'item_href': '0',
                                'element_type': 'image',
                                'image_href': basename,
                                'paragraph_number': fragment_id,
                                'has_mismatch': False,
                            }
                            if pending_anchor_ids:
                                img_para['extra_anchor_ids'] = list(pending_anchor_ids)
                                pending_anchor_ids = []
                            paragraphs.append(img_para)
                            fragment_id += 1
                        except Exception as e:
                            logger.warning(f"MobiProcessor: could not read image {img_path}: {e}")
                continue

            if tag in HEADING_TAGS:
                element_type = 'heading'
            elif tag == 'p':
                element_type = 'paragraph'
            elif tag in LEAF_BLOCK_TAGS and _is_leaf_block(elem):
                element_type = 'paragraph'
            else:
                continue

            elem_id_str = f'mobi_{fragment_id}'
            elem_id_attr = (elem.get('id') or '').strip()
            elem_style = _extract_block_style(elem)

            if use_inline:
                clean_text, inline_map = self._extract_html_inline(
                    elem, INLINE_TAGS, placeholder_pattern, css_class_styles
                )
                if not clean_text.strip():
                    if elem_id_attr and elem_id_attr not in pending_anchor_ids:
                        pending_anchor_ids.append(elem_id_attr)
                    for info in (inline_map or {}).values():
                        if info.get('tag') == 'a':
                            for attr_name in ('id', 'name'):
                                val = info.get('attributes', {}).get(attr_name, '').strip()
                                if val and val not in pending_anchor_ids:
                                    pending_anchor_ids.append(val)
                    continue
                actual_text = re.sub(r'</?p_\d{2}>', '', clean_text).strip()
                if not actual_text:
                    if elem_id_attr and elem_id_attr not in pending_anchor_ids:
                        pending_anchor_ids.append(elem_id_attr)
                    for info in (inline_map or {}).values():
                        if info.get('tag') == 'a':
                            for attr_name in ('id', 'name'):
                                val = info.get('attributes', {}).get(attr_name, '').strip()
                                if val and val not in pending_anchor_ids:
                                    pending_anchor_ids.append(val)
                    continue
                para = {
                    'id': str(fragment_id),
                    'original_text': clean_text,
                    'translated_text': '',
                    'is_translated': False,
                    'item_href': '0',
                    'element_type': element_type,
                    'paragraph_number': fragment_id,
                    'has_mismatch': False,
                    'processing_mode': 'inline',
                    'placeholder_pattern': placeholder_pattern,
                    'reserve_elements': [],
                    'prefix_reserve_tags': [],
                    'suffix_reserve_tags': [],
                    'is_non_translatable': False,
                }
                if element_type == 'heading':
                    para['heading_level'] = tag
                if inline_map:
                    para['inline_formatting_map'] = inline_map
            else:
                clean_text = (etree.tostring(elem, encoding='unicode', method='text') or '').strip()
                if not clean_text:
                    if elem_id_attr and elem_id_attr not in pending_anchor_ids:
                        pending_anchor_ids.append(elem_id_attr)
                    continue
                para = {
                    'id': str(fragment_id),
                    'original_text': clean_text,
                    'translated_text': '',
                    'is_translated': False,
                    'item_href': '0',
                    'element_type': element_type,
                    'paragraph_number': fragment_id,
                    'has_mismatch': False,
                    'processing_mode': 'legacy',
                    'original_html': self._build_kindle_original_html(elem, elem_id_str),
                    'reserve_elements': [],
                    'prefix_reserve_tags': [],
                    'suffix_reserve_tags': [],
                    'is_non_translatable': False,
                }
                if element_type == 'heading':
                    para['heading_level'] = tag

            if elem_id_attr:
                para['element_id'] = elem_id_attr

            if pending_anchor_ids:
                para['extra_anchor_ids'] = list(pending_anchor_ids)
                pending_anchor_ids = []

            if elem_style:
                para['kindle_style'] = elem_style

            paragraphs.append(para)
            fragment_id += 1

        return paragraphs

    def _extract_html_inline(self, element, inline_tags: set, placeholder_pattern: str, css_class_styles: dict = None):
        INLINE_TAG_MAP = {
            'b': 'b', 'strong': 'b', 'i': 'i', 'em': 'i',
            'u': 'u', 'span': 'span', 'small': 'small', 'a': 'a',
            'font': 'span',
        }
        FONT_SIZE_MAP = {
            '1': '0.6em', '2': '0.8em', '3': '1em', '4': '1.2em',
            '5': '1.5em', '6': '2em', '7': '2.5em',
            '+1': '1.1em', '+2': '1.3em', '+3': '1.6em', '+4': '2em',
            '-1': '0.85em', '-2': '0.7em',
        }
        ATTRS_TO_KEEP = {'href', 'style', 'class', 'id', 'target', 'title', 'name'}
        KEEP_INLINE_PROPS = {
            'font-size', 'font-weight', 'font-style', 'color',
            'text-decoration', 'vertical-align', 'letter-spacing',
            'float', 'line-height',
        }
        inline_formatting_map = {}
        counter = 0
        replacements = []

        for child in element.iter():
            child_tag = child.tag if isinstance(child.tag, str) else ''
            child_tag = child_tag.lower().split('}')[-1]
            if child_tag not in inline_tags:
                continue
            html_tag = INLINE_TAG_MAP.get(child_tag, child_tag)
            raw_attrs = dict(child.attrib) if child.attrib else {}
            filtered_attrs = {k: v for k, v in raw_attrs.items() if k in ATTRS_TO_KEEP}

            if child_tag == 'font':
                font_styles = {}
                size_val = raw_attrs.get('size', '').strip()
                if size_val and size_val in FONT_SIZE_MAP:
                    font_styles['font-size'] = FONT_SIZE_MAP[size_val]
                color_val = raw_attrs.get('color', '').strip()
                if color_val and color_val.lower() not in ('#000000', 'black', '#000'):
                    font_styles['color'] = color_val
                face_val = raw_attrs.get('face', '').strip()
                if face_val:
                    font_styles['font-family'] = face_val
                if not font_styles:
                    continue
                existing_style = filtered_attrs.get('style', '')
                existing_parts = {}
                for decl in existing_style.split(';'):
                    decl = decl.strip()
                    if ':' in decl:
                        p, _, v = decl.partition(':')
                        existing_parts[p.strip().lower()] = v.strip()
                font_styles.update(existing_parts)
                filtered_attrs['style'] = '; '.join(f'{k}:{v}' for k, v in font_styles.items())

            if css_class_styles and filtered_attrs.get('class'):
                class_styles = {}
                for cls in filtered_attrs['class'].split():
                    for prop, val in (css_class_styles.get(cls) or {}).items():
                        if prop in KEEP_INLINE_PROPS and prop not in class_styles:
                            class_styles[prop] = val
                if class_styles:
                    existing_style = filtered_attrs.get('style', '')
                    existing_parts = {}
                    for decl in existing_style.split(';'):
                        decl = decl.strip()
                        if ':' in decl:
                            p, _, v = decl.partition(':')
                            existing_parts[p.strip().lower()] = v.strip()
                    class_styles.update(existing_parts)
                    filtered_attrs['style'] = '; '.join(f'{k}:{v}' for k, v in class_styles.items())

            elem_id = counter
            opening = '<p_{:02d}>'.format(elem_id)
            closing = '</p_{:02d}>'.format(elem_id)

            inline_formatting_map[elem_id] = {
                'tag': html_tag,
                'attributes': filtered_attrs,
                'opening_placeholder': opening,
                'closing_placeholder': closing,
                'has_leading_space': False,
                'has_trailing_space': False,
            }
            replacements.append((child, opening, closing))
            counter += 1

        if not replacements:
            return (etree.tostring(element, encoding='unicode', method='text') or '').strip(), {}

        replace_map = {id(e): (op, cl) for e, op, cl in replacements}

        def serialize(node):
            parts = []
            nid = id(node)
            if nid in replace_map:
                op, cl = replace_map[nid]
                parts.append(op)
                if node.text:
                    parts.append(node.text)
                for child in node:
                    parts.append(serialize(child))
                    if child.tail:
                        parts.append(child.tail)
                parts.append(cl)
            else:
                if node.text:
                    parts.append(node.text)
                for child in node:
                    parts.append(serialize(child))
                    if child.tail:
                        parts.append(child.tail)
            return ''.join(parts)

        return serialize(element).strip(), inline_formatting_map

    def _build_kindle_original_html(self, element, element_id: str) -> str:
        XHTML_NS = 'http://www.w3.org/1999/xhtml'
        INLINE_MAP = {'b': 'b', 'strong': 'b', 'i': 'em', 'em': 'em', 'u': 'u'}

        def _serialize(node) -> str:
            tag = node.tag if isinstance(node.tag, str) else ''
            tag_local = tag.lower().split('}')[-1]
            xhtml_tag = INLINE_MAP.get(tag_local)
            inner = ''
            if node.text:
                inner += html_module.escape(node.text)
            for child in node:
                inner += _serialize(child)
                if child.tail:
                    inner += html_module.escape(child.tail)
            if xhtml_tag:
                return f'<{xhtml_tag}>{inner}</{xhtml_tag}>'
            return inner

        parts = []
        if element.text:
            parts.append(html_module.escape(element.text))
        for child in element:
            parts.append(_serialize(child))
            if child.tail:
                parts.append(html_module.escape(child.tail))
        return f'<p xmlns="{XHTML_NS}" id="{element_id}">{"".join(parts)}</p>'

class AZW3Processor:
    def __init__(self, app_settings: dict):
        self.app_settings = app_settings
        self.original_path = None
        self._preview_tempdir = None
        self.image_items: List[Dict] = []
        self.paragraphs: List[Dict] = []
        self.css_class_styles: Dict = {}
        self.book = None
        self._current_body_styles: Dict = {}
        self.skip_inline_tags = app_settings.get('skip_inline_tags', {})
        self.PRIORITY_TAGS = ['p', 'pre', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']
        self.NON_INLINE_ELEMENTS = {
            'address', 'blockquote', 'dialog', 'div', 'figure', 'figcaption',
            'footer', 'header', 'legend', 'main', 'p', 'pre', 'search', 'article',
            'aside', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'hgroup', 'nav',
            'section', 'dd', 'dl', 'dt', 'menu', 'ol', 'ul', 'table', 'caption',
            'colgroup', 'col', 'thead', 'tbody', 'tfoot', 'tr', 'td', 'th', 'li',
        }
        self.NOISE_TAGS = ['rt', 'rp']
        self.RESERVE_TAGS = [
            'img', 'code', 'br', 'hr', 'sub', 'sup', 'kbd',
            'abbr', 'wbr', 'var', 'canvas', 'svg', 'script',
            'style', 'math',
        ]
        self.HEADING_TAGS = {'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}
        self.BLOCK_TAGS = {
            'p', 'div', 'section', 'article', 'blockquote', 'header', 'footer',
            'main', 'ul', 'ol', 'table', 'nav', 'aside', 'figure', 'figcaption',
            'address', 'pre', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
        }
        self.INLINE_TAGS = {'b', 'strong', 'i', 'em', 'u', 'span', 'small', 'a'}
        self.CONTAINER_TAGS = {
            'ul', 'ol', 'dl',
            'table', 'tbody', 'thead', 'tfoot', 'tr',
            'div', 'section', 'article', 'aside', 'nav', 'main',
            'header', 'footer', 'figure', 'body',
        }
        self.INLINE_FORMATTING_TAGS = [
            'a', 'i', 'b', 'em', 'strong',
            'u', 'sup', 'sub', 'small', 'span',
        ]
        self.STRUCTURAL_SPAN_CLASSES = {
            'first-letter', 'last-word', 'item-number', 'element-number',
        }
        self.INLINE_TAG_MAP = {
            'b': 'b', 'strong': 'b', 'i': 'i', 'em': 'i',
            'u': 'u', 'span': 'span', 'small': 'small', 'a': 'a',
        }
        self.MEDIA_TYPE_MAP = {
            '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
            '.png': 'image/png', '.gif': 'image/gif',
            '.bmp': 'image/bmp', '.webp': 'image/webp',
        }
        self.KEEP_STYLE_PROPS = {
            'text-align', 'color', 'font-size', 'font-weight',
            'font-style', 'text-indent', 'margin-left', 'margin-right',
            'background-color', 'letter-spacing', 'line-height',
        }
        self.KEEP_CSS_PROPS = {'text-align'}
        self.PLACEHOLDER_PATTERN = '<p_{:02d}>'
        self.RESERVE_PLACEHOLDER_PATTERN = '<id_{:02d}>'

    def get_file_type(self) -> str:
        return "azw3"

    def load(self, path: str) -> Tuple[List[Dict], str]:
        try:
            import mobi
            self.original_path = path
            self.image_items = []
            self.paragraphs = []
            if self._preview_tempdir is not None:
                shutil.rmtree(self._preview_tempdir, ignore_errors=True)
                self._preview_tempdir = None
            tempdir, filepath = mobi.extract(path)
            self._preview_tempdir = tempdir
            ext = os.path.splitext(filepath)[1].lower()
            if ext == '.epub':
                self._load_from_epub(filepath)
            else:
                self.paragraphs = self._parse_html(filepath)
            logger.info(
                f"AZW3Processor: loaded {len(self.paragraphs)} fragments, "
                f"{len(self.image_items)} image items from {path}"
            )
            return self.paragraphs, path
        except Exception as e:
            logger.error(f"AZW3 load error: {e}", exc_info=True)
            raise

    def cleanup_preview_tempdir(self):
        if self._preview_tempdir is not None:
            shutil.rmtree(self._preview_tempdir, ignore_errors=True)
            self._preview_tempdir = None

    def _local_xpath(self, element, tag_name):
        if element is None:
            return []
        return element.xpath(f'.//*[local-name()="{tag_name}"]')

    def _local_find(self, element, tag_name):
        if element is None:
            return None
        results = self._local_xpath(element, tag_name)
        return results[0] if results else None

    def _normalize_namespaces(self, item):
        if item.data is None:
            return
        root = item.data
        if isinstance(root.tag, str) and root.tag.startswith('{'):
            return
        ns = 'http://www.w3.org/1999/xhtml'
        for elem in root.iter():
            if isinstance(elem.tag, str) and not elem.tag.startswith('{'):
                elem.tag = f'{{{ns}}}{elem.tag}'

    def _load_from_epub(self, epub_path: str):
        try:
            self.book = read_epub(epub_path)
            self.css_class_styles = self._parse_css_files()
            seen = set()
            spine_order = {}
            nav_hrefs = set()

            try:
                opf_path = os.path.join(self.book.temp_dir, self.book.opf_path)
                with open(opf_path, 'rb') as f:
                    opf_tree = etree.parse(f)
                spine_elem = opf_tree.find('.//opf:spine', namespaces=NAMESPACES)
                if spine_elem is not None:
                    for idx, itemref in enumerate(spine_elem.findall('.//opf:itemref', namespaces=NAMESPACES)):
                        idref = itemref.get('idref')
                        if idref:
                            spine_order[idref] = idx
                manifest_elem = opf_tree.find('.//opf:manifest', namespaces=NAMESPACES)
                if manifest_elem is not None:
                    for mitem in manifest_elem.findall('opf:item', namespaces=NAMESPACES):
                        if 'nav' in (mitem.get('properties') or '').split():
                            href = mitem.get('href', '')
                            if href:
                                nav_hrefs.add(href)
                                nav_hrefs.add(os.path.basename(href))
            except Exception:
                pass

            logger.debug(f"AZW3Processor._load_from_epub: epub_path={epub_path}")
            logger.debug(f"AZW3Processor._load_from_epub: spine_order={spine_order}")
            logger.debug(f"AZW3Processor._load_from_epub: nav_hrefs={nav_hrefs}")
            logger.debug(f"AZW3Processor._load_from_epub: css_class_styles count={len(self.css_class_styles)}")

            doc_items = list(self.book.get_items_of_type('DOCUMENT'))

            def get_sort_key(item):
                try:
                    item_id = next((m.get('id') for m in self.book.manifest_items if isinstance(m, dict) and m.get('href') == item.href), None)
                    if item_id in spine_order:
                        return (0, spine_order[item_id])
                    return (1, item.href)
                except Exception:
                    return (2, item.href)

            doc_items.sort(key=get_sort_key)

            logger.debug(f"AZW3Processor._load_from_epub: doc_items count={len(doc_items)}")
            logger.debug(f"AZW3Processor._load_from_epub: doc_items hrefs={[item.href for item in doc_items]}")

            for item in doc_items:
                item_basename = os.path.basename(item.href)
                if item.href in nav_hrefs or item_basename in nav_hrefs:
                    logger.debug(f"AZW3Processor._load_from_epub: SKIPPING nav item href={item.href}")
                    continue
                if item.data is None:
                    logger.debug(f"AZW3Processor._load_from_epub: SKIPPING item with no data href={item.href}")
                    continue

                self._normalize_namespaces(item)

                body = self._local_find(item.data, 'body')
                if body is None:
                    logger.debug(f"AZW3Processor._load_from_epub: SKIPPING item no body found href={item.href}")
                    continue

                nav_elems = self._local_xpath(item.data, 'nav')
                is_nav_doc = any(nav_el.get('{http://www.idpf.org/2007/ops}type') for nav_el in nav_elems)
                if is_nav_doc:
                    logger.debug(f"AZW3Processor._load_from_epub: SKIPPING nav doc href={item.href}")
                    continue

                body_class_styles = {}
                for cls in (body.get('class') or '').split():
                    for prop, val in self.css_class_styles.get(cls, {}).items():
                        if prop not in body_class_styles:
                            body_class_styles[prop] = val
                self._current_body_styles = body_class_styles

                paras_before = len(self.paragraphs)
                self._extract_elements_lxml(body, item.href, seen)
                paras_after = len(self.paragraphs)

                self._current_body_styles = {}

                logger.debug(
                    f"AZW3Processor._load_from_epub: item href={item.href} -> "
                    f"extracted {paras_after - paras_before} fragments (total={paras_after})"
                )

            for item in self.book.get_items_of_type('IMAGE'):
                content = getattr(item, 'content', b'') or b''
                if content:
                    self.image_items.append({
                        'id': getattr(item, 'id', '') or getattr(item, 'uid', ''),
                        'file_name': getattr(item, 'file_name', '') or getattr(item, 'href', ''),
                        'media_type': getattr(item, 'media_type', 'image/jpeg'),
                        'content': content,
                    })

            logger.debug(f"AZW3Processor._load_from_epub: image_items count={len(self.image_items)}")
            logger.debug(f"AZW3Processor._load_from_epub: image_items ids={[img.get('id') for img in self.image_items]}")
            logger.debug(f"AZW3Processor._load_from_epub: image_items file_names={[img.get('file_name') for img in self.image_items]}")
            logger.debug(f"AZW3Processor._load_from_epub: image_items media_types={[img.get('media_type') for img in self.image_items]}")

            if self.image_items and self.paragraphs:
                if not any(img.get('is_cover') for img in self.image_items):
                    for p in self.paragraphs:
                        if p.get('element_type') == 'image' and p.get('image_href'):
                            for img in self.image_items:
                                if os.path.basename(img.get('file_name', '')) == os.path.basename(p['image_href']):
                                    img['is_cover'] = True
                                    logger.debug(
                                        f"AZW3Processor._load_from_epub: cover assigned via para match "
                                        f"image_href={p['image_href']} -> file_name={img.get('file_name')}"
                                    )
                                    break
                            break
                    else:
                        self.image_items[0]['is_cover'] = True
                        logger.debug(
                            f"AZW3Processor._load_from_epub: cover assigned to first image_item "
                            f"id={self.image_items[0].get('id')} file_name={self.image_items[0].get('file_name')}"
                        )

            cover_img = next((img for img in self.image_items if img.get('is_cover')), None)
            logger.debug(
                f"AZW3Processor._load_from_epub: final cover_img="
                f"{'id=' + cover_img['id'] + ' file_name=' + cover_img.get('file_name', '') if cover_img else 'NOT FOUND'}"
            )

            image_paras = [p for p in self.paragraphs if p.get('element_type') == 'image']
            logger.debug(f"AZW3Processor._load_from_epub: image paragraphs count={len(image_paras)}")
            logger.debug(f"AZW3Processor._load_from_epub: image_hrefs in paragraphs={[p.get('image_href') for p in image_paras]}")

        except Exception as e:
            logger.error(f"AZW3Processor: EPUB load error: {e}", exc_info=True)
            raise

    def _extract_elements_lxml(self, root, item_href: str, seen: set, list_context=None):
        XLINK_NS = 'http://www.w3.org/1999/xlink'

        for child in root:
            if not isinstance(child, etree._Element):
                continue
            tag_name = etree.QName(child).localname.lower()

            if tag_name == 'img':
                src = child.get('src', '')
                if src and not src.startswith('data:'):
                    img_para = {
                        'id': f'img_{uuid.uuid4().hex[:8]}',
                        'original_text': '',
                        'translated_text': '',
                        'is_translated': False,
                        'item_href': item_href,
                        'element_type': 'image',
                        'image_href': src,
                        'has_mismatch': False,
                    }
                    if list_context and list_context.get('list_id'):
                        img_para['list_metadata'] = {
                            'list_id': list_context['list_id'],
                            'list_type': list_context.get('list_type', 'ul'),
                            'list_class': list_context.get('list_class', ''),
                            'li_index': list_context.get('li_index', 0),
                            'li_class': list_context.get('li_class', ''),
                            'parent_list_id': list_context.get('parent_list_id'),
                            'parent_li_index': list_context.get('parent_li_index'),
                        }
                    self.paragraphs.append(img_para)
                continue

            if tag_name == 'svg':
                for svg_child in child.iter():
                    href = svg_child.get('href') or svg_child.get(f'{{{XLINK_NS}}}href')
                    if href and not href.startswith(('data:', '#')):
                        img_para = {
                            'id': f'img_{uuid.uuid4().hex[:8]}',
                            'original_text': '',
                            'translated_text': '',
                            'is_translated': False,
                            'item_href': item_href,
                            'element_type': 'image',
                            'image_href': href,
                            'has_mismatch': False,
                        }
                        if list_context and list_context.get('list_id'):
                            img_para['list_metadata'] = {
                                'list_id': list_context['list_id'],
                                'list_type': list_context.get('list_type', 'ul'),
                                'list_class': list_context.get('list_class', ''),
                                'li_index': list_context.get('li_index', 0),
                                'li_class': list_context.get('li_class', ''),
                                'parent_list_id': list_context.get('parent_list_id'),
                                'parent_li_index': list_context.get('parent_li_index'),
                            }
                        self.paragraphs.append(img_para)
                        break
                continue

            if tag_name == 'li':
                child_list_context = None
                if list_context is not None and '_li_counter' in list_context:
                    li_idx = list_context['_li_counter']
                    list_context['_li_counter'] += 1
                    child_list_context = dict(list_context)
                    child_list_context['li_index'] = li_idx
                    child_list_context['li_class'] = (child.get('class') or '').strip()

                child_elem_tags = [
                    etree.QName(c).localname.lower()
                    for c in child
                    if isinstance(c, etree._Element)
                ]
                has_block = any(
                    t in self.PRIORITY_TAGS or t in self.CONTAINER_TAGS
                    for t in child_elem_tags
                )
                if has_block:
                    self._extract_elements_lxml(child, item_href, seen, child_list_context)
                else:
                    text_content = (etree.tostring(child, encoding='unicode', method='text') or '').strip()
                    img_children = self._local_xpath(child, 'img')
                    svg_children = self._local_xpath(child, 'svg')
                    if not text_content and (img_children or svg_children):
                        for img_elem in img_children:
                            src = img_elem.get('src', '')
                            if src and not src.startswith('data:'):
                                img_para = {
                                    'id': f'img_{uuid.uuid4().hex[:8]}',
                                    'original_text': '',
                                    'translated_text': '',
                                    'is_translated': False,
                                    'item_href': item_href,
                                    'element_type': 'image',
                                    'image_href': src,
                                    'has_mismatch': False,
                                }
                                if child_list_context and child_list_context.get('list_id'):
                                    img_para['list_metadata'] = {
                                        'list_id': child_list_context['list_id'],
                                        'list_type': child_list_context.get('list_type', 'ul'),
                                        'list_class': child_list_context.get('list_class', ''),
                                        'li_index': child_list_context.get('li_index', 0),
                                        'li_class': child_list_context.get('li_class', ''),
                                        'parent_list_id': child_list_context.get('parent_list_id'),
                                        'parent_li_index': child_list_context.get('parent_li_index'),
                                    }
                                self.paragraphs.append(img_para)
                        for svg_elem in svg_children:
                            for svg_inner in svg_elem.iter():
                                href = svg_inner.get('href') or svg_inner.get(f'{{{XLINK_NS}}}href')
                                if href and not href.startswith(('data:', '#')):
                                    img_para = {
                                        'id': f'img_{uuid.uuid4().hex[:8]}',
                                        'original_text': '',
                                        'translated_text': '',
                                        'is_translated': False,
                                        'item_href': item_href,
                                        'element_type': 'image',
                                        'image_href': href,
                                        'has_mismatch': False,
                                    }
                                    if child_list_context and child_list_context.get('list_id'):
                                        img_para['list_metadata'] = {
                                            'list_id': child_list_context['list_id'],
                                            'list_type': child_list_context.get('list_type', 'ul'),
                                            'list_class': child_list_context.get('list_class', ''),
                                            'li_index': child_list_context.get('li_index', 0),
                                            'li_class': child_list_context.get('li_class', ''),
                                            'parent_list_id': child_list_context.get('parent_list_id'),
                                            'parent_li_index': child_list_context.get('parent_li_index'),
                                        }
                                    self.paragraphs.append(img_para)
                                    break
                    elif text_content or img_children or svg_children:
                        self._process_element_lxml(child, item_href, seen, child_list_context)
                continue

            if tag_name in ('ol', 'ul'):
                new_list_context = {
                    'list_id': uuid.uuid4().hex[:8],
                    'list_type': tag_name,
                    'list_class': (child.get('class') or '').strip(),
                    'li_index': -1,
                    'li_class': '',
                    '_li_counter': 0,
                    'parent_list_id': list_context['list_id'] if list_context and list_context.get('list_id') else None,
                    'parent_li_index': list_context.get('li_index') if list_context else None,
                }
                self._extract_elements_lxml(child, item_href, seen, new_list_context)
                continue

            if tag_name in self.PRIORITY_TAGS:
                text_content = (etree.tostring(child, encoding='unicode', method='text') or '').strip()
                img_children = self._local_xpath(child, 'img')
                svg_children = self._local_xpath(child, 'svg')
                if not text_content and (img_children or svg_children):
                    for img_elem in img_children:
                        src = img_elem.get('src', '')
                        if src and not src.startswith('data:'):
                            img_para = {
                                'id': f'img_{uuid.uuid4().hex[:8]}',
                                'original_text': '',
                                'translated_text': '',
                                'is_translated': False,
                                'item_href': item_href,
                                'element_type': 'image',
                                'image_href': src,
                                'has_mismatch': False,
                            }
                            if list_context and list_context.get('list_id'):
                                img_para['list_metadata'] = {
                                    'list_id': list_context['list_id'],
                                    'list_type': list_context.get('list_type', 'ul'),
                                    'list_class': list_context.get('list_class', ''),
                                    'li_index': list_context.get('li_index', 0),
                                    'li_class': list_context.get('li_class', ''),
                                    'parent_list_id': list_context.get('parent_list_id'),
                                    'parent_li_index': list_context.get('parent_li_index'),
                                }
                            self.paragraphs.append(img_para)
                    for svg_elem in svg_children:
                        for svg_inner in svg_elem.iter():
                            href = svg_inner.get('href') or svg_inner.get(f'{{{XLINK_NS}}}href')
                            if href and not href.startswith(('data:', '#')):
                                img_para = {
                                    'id': f'img_{uuid.uuid4().hex[:8]}',
                                    'original_text': '',
                                    'translated_text': '',
                                    'is_translated': False,
                                    'item_href': item_href,
                                    'element_type': 'image',
                                    'image_href': href,
                                    'has_mismatch': False,
                                }
                                if list_context and list_context.get('list_id'):
                                    img_para['list_metadata'] = {
                                        'list_id': list_context['list_id'],
                                        'list_type': list_context.get('list_type', 'ul'),
                                        'list_class': list_context.get('list_class', ''),
                                        'li_index': list_context.get('li_index', 0),
                                        'li_class': list_context.get('li_class', ''),
                                        'parent_list_id': list_context.get('parent_list_id'),
                                        'parent_li_index': list_context.get('parent_li_index'),
                                    }
                                self.paragraphs.append(img_para)
                                break
                    continue
                self._process_element_lxml(child, item_href, seen, list_context)
                continue

            if tag_name in self.CONTAINER_TAGS:
                if not [c for c in child if isinstance(c, etree._Element)] and self._has_any_text(child):
                    self._process_element_lxml(child, item_href, seen, list_context)
                else:
                    self._extract_elements_lxml(child, item_href, seen, list_context)
                continue

            if self._is_inline_only_lxml(child) or self._has_any_text(child):
                self._process_element_lxml(child, item_href, seen, list_context)

    def _is_inline_only_lxml(self, element) -> bool:
        return len(self._local_xpath(element, 'p')) == 0 and \
               len(self._local_xpath(element, 'div')) == 0 and \
               len(self._local_xpath(element, 'h1')) == 0

    def _process_element_lxml(self, element, item_href: str, seen: set, list_context=None):
        if self.app_settings.get('use_inline_formatting', True):
            self._process_element_inline(element, item_href, seen, list_context)
        else:
            self._process_element_legacy(element, item_href, seen, list_context)

    def _parse_css_files(self) -> Dict:
        css_class_styles = {}
        KEEP_CSS_PROPS = {
            'text-align', 'font-size', 'font-weight', 'font-style',
            'color', 'text-indent', 'margin-left', 'margin-right',
            'margin-top', 'margin-bottom', 'margin',
            'padding-left', 'padding-right', 'padding-top', 'padding-bottom', 'padding',
            'line-height', 'letter-spacing', 'font-variant',
            'text-decoration', 'vertical-align',
            'float', 'display',
        }
        for manifest_item in self.book.manifest_items:
            if not isinstance(manifest_item, dict):
                continue
            if manifest_item.get('media_type') != 'text/css':
                continue
            href = manifest_item.get('href', '')
            if not href:
                continue
            css_path = os.path.join(self.book.content_dir, href)
            if not os.path.exists(css_path):
                continue
            try:
                with open(css_path, 'r', encoding='utf-8', errors='ignore') as f:
                    css_content = f.read()
                for match in re.finditer(r'([^{}]*)\{([^}]*)\}', css_content, re.DOTALL):
                    selector = match.group(1).strip()
                    declarations = match.group(2)
                    if not selector or selector.startswith('@'):
                        continue
                    class_names = re.findall(r'\.(\w+)', selector)
                    if not class_names:
                        continue
                    found_styles = {}
                    for decl in declarations.split(';'):
                        decl = decl.strip()
                        if ':' not in decl:
                            continue
                        prop, _, val = decl.partition(':')
                        prop = prop.strip().lower()
                        val = val.strip()
                        if val.endswith('!important'):
                            val = val[:-len('!important')].strip()
                        if prop in KEEP_CSS_PROPS:
                            found_styles[prop] = val
                    if found_styles:
                        for class_name in class_names:
                            if class_name in css_class_styles:
                                css_class_styles[class_name].update(found_styles)
                            else:
                                css_class_styles[class_name] = dict(found_styles)
            except Exception as e:
                logger.debug(f"CSS parse error for {css_path}: {e}")
        logger.info(f"AZW3Processor: parsed {len(css_class_styles)} CSS class style(s)")
        return css_class_styles

    def _parse_html(self, html_path: str) -> List[Dict]:
        paragraphs = []
        fragment_id = 0
        use_inline = self.app_settings.get('use_inline_formatting', True)
        placeholder_pattern = self.PLACEHOLDER_PATTERN
        cover_image_added = False

        logger.debug(f"AZW3Processor._parse_html: html_path={html_path}")

        with open(html_path, 'rb') as f:
            content = f.read()

        try:
            raw_html_bytes = content[:8000]
            logger.debug(f"AZW3Processor._parse_html: RAW html (first 8000 bytes):\n{raw_html_bytes.decode('utf-8', errors='replace')}")
        except Exception:
            pass

        try:
            tree = etree.fromstring(content, parser=etree.HTMLParser())
        except Exception:
            return paragraphs
        body = tree.find('.//body')
        if body is None:
            logger.warning(f"AZW3Processor._parse_html: no <body> found in {html_path}")
            return paragraphs

        html_dir = os.path.dirname(html_path)
        parent_dir = os.path.dirname(html_dir)
        logger.debug(f"AZW3Processor._parse_html: html_dir={html_dir}")
        logger.debug(f"AZW3Processor._parse_html: html_dir contents={sorted(os.listdir(html_dir)) if os.path.exists(html_dir) else 'NOT EXISTS'}")
        logger.debug(f"AZW3Processor._parse_html: parent_dir contents={sorted(os.listdir(parent_dir)) if os.path.exists(parent_dir) else 'NOT EXISTS'}")

        try:
            css_sources = []
            for style_elem in tree.findall('.//style'):
                if style_elem.text and style_elem.text.strip():
                    css_sources.append(style_elem.text)
            for link_elem in tree.findall('.//link'):
                rel = (link_elem.get('rel') or '').lower()
                if rel == 'stylesheet' or (link_elem.get('type') or '').lower() == 'text/css':
                    href = (link_elem.get('href') or '').strip()
                    if href and not href.startswith(('http://', 'https://', 'data:')):
                        css_path = os.path.normpath(os.path.join(html_dir, href))
                        if os.path.isfile(css_path):
                            try:
                                with open(css_path, 'r', encoding='utf-8', errors='ignore') as fcss:
                                    css_sources.append(fcss.read())
                            except Exception:
                                pass
            for root_dir in (html_dir, parent_dir):
                if not os.path.isdir(root_dir):
                    continue
                for sub_root, _, files in os.walk(root_dir):
                    for file in files:
                        if file.lower().endswith('.css'):
                            css_path = os.path.join(sub_root, file)
                            try:
                                with open(css_path, 'r', encoding='utf-8', errors='ignore') as fcss:
                                    css_sources.append(fcss.read())
                            except Exception:
                                pass
            KEEP_CSS_PROPS_HTML = {
                'text-align', 'font-size', 'font-weight', 'font-style',
                'color', 'text-indent', 'margin-left', 'margin-right',
                'margin-top', 'margin-bottom', 'margin',
                'padding-left', 'padding-right', 'padding-top', 'padding-bottom', 'padding',
                'line-height', 'letter-spacing', 'font-variant',
                'text-decoration', 'vertical-align',
                'float', 'display',
            }
            for css_text in css_sources:
                for match in re.finditer(r'\.(-?[\w][\w-]*)\s*\{([^}]*)\}', css_text, re.DOTALL):
                    class_name = match.group(1)
                    declarations = match.group(2)
                    found_styles = {}
                    for decl in declarations.split(';'):
                        decl = decl.strip()
                        if ':' not in decl:
                            continue
                        prop, _, val = decl.partition(':')
                        prop = prop.strip().lower()
                        val = val.strip()
                        if val.endswith('!important'):
                            val = val[:-len('!important')].strip()
                        if prop in KEEP_CSS_PROPS_HTML:
                            found_styles[prop] = val
                    if found_styles:
                        if class_name not in self.css_class_styles:
                            self.css_class_styles[class_name] = {}
                        self.css_class_styles[class_name].update(found_styles)
            logger.debug(f"AZW3Processor._parse_html: css_class_styles count={len(self.css_class_styles)}")
        except Exception as e:
            logger.warning(f"AZW3Processor._parse_html: CSS loading failed: {e}")

        body_class_styles = {}
        for cls in (body.get('class') or '').split():
            for prop, val in self.css_class_styles.get(cls, {}).items():
                if prop not in body_class_styles:
                    body_class_styles[prop] = val
        self._current_body_styles = body_class_styles

        def _is_leaf_block(element):
            for child in element:
                ctag = (child.tag if isinstance(child.tag, str) else '').lower().split('}')[-1]
                if ctag in self.BLOCK_TAGS:
                    return False
            return True

        def _extract_block_style(elem):
            style_parts = {}
            align_attr = (elem.get('align') or '').strip().lower()
            if align_attr in ('center', 'left', 'right', 'justify'):
                style_parts['text-align'] = align_attr
            raw = (elem.get('style') or '').strip()
            for decl in raw.split(';'):
                decl = decl.strip()
                if not decl or ':' not in decl:
                    continue
                prop, _, val = decl.partition(':')
                prop = prop.strip().lower()
                if prop in self.KEEP_STYLE_PROPS:
                    style_parts[prop] = val.strip()
            for cls in (elem.get('class') or '').split():
                if cls in self.css_class_styles:
                    for prop, val in self.css_class_styles[cls].items():
                        if prop not in style_parts:
                            style_parts[prop] = val
            for prop, val in self._current_body_styles.items():
                if prop not in style_parts:
                    style_parts[prop] = val
            return '; '.join(f'{k}:{v}' for k, v in style_parts.items())

        for elem in body.iter():
            tag = elem.tag if isinstance(elem.tag, str) else ''
            tag = tag.lower().split('}')[-1]
            if tag == 'img':
                src = elem.get('src', '')
                if src and not src.startswith('data:'):
                    img_path = os.path.join(html_dir, src)
                    if not os.path.isfile(img_path):
                        for alt_dir in ('images', 'Images'):
                            alt_path = os.path.join(html_dir, alt_dir, os.path.basename(src))
                            if os.path.isfile(alt_path):
                                img_path = alt_path
                                src = os.path.relpath(img_path, html_dir).replace('\\', '/')
                                break
                    if os.path.isfile(img_path):
                        ext = os.path.splitext(src)[1].lower()
                        mt = self.MEDIA_TYPE_MAP.get(ext, 'image/jpeg')
                        try:
                            with open(img_path, 'rb') as f:
                                img_content = f.read()
                            basename = os.path.basename(src)
                            img_entry = {
                                'id': basename,
                                'file_name': basename,
                                'media_type': mt,
                                'content': img_content,
                            }
                            if not cover_image_added:
                                img_entry['is_cover'] = True
                                cover_image_added = True
                            self.image_items.append(img_entry)
                            logger.debug(
                                f"AZW3Processor._parse_html: found image src={src} -> "
                                f"img_path={img_path} basename={basename} mt={mt} is_cover={img_entry.get('is_cover', False)}"
                            )
                            paragraphs.append({
                                'id': str(fragment_id),
                                'original_text': '',
                                'translated_text': '',
                                'is_translated': False,
                                'item_href': '0',
                                'element_type': 'image',
                                'image_href': basename,
                                'paragraph_number': fragment_id,
                                'has_mismatch': False,
                            })
                            fragment_id += 1
                        except Exception as e:
                            logger.warning(f"AZW3Processor: could not read image {img_path}: {e}")
                    else:
                        logger.warning(f"AZW3Processor._parse_html: image NOT FOUND src={src} tried path={img_path}")
                continue
            if tag in self.HEADING_TAGS:
                element_type = tag
            elif tag == 'p':
                element_type = 'paragraph'
            elif tag == 'div' and _is_leaf_block(elem):
                element_type = 'paragraph'
            else:
                continue
            elem_id_str = f'azw3_{fragment_id}'
            elem_id_attr = (elem.get('id') or '').strip()
            elem_style = _extract_block_style(elem)
            if use_inline:
                clean_text, inline_map = self._html_extract_inline(elem, placeholder_pattern)
                if not clean_text.strip():
                    continue
                actual_text = re.sub(r'</?p_\d{2}>', '', clean_text).strip()
                if not actual_text:
                    continue
                para = {
                    'id': str(fragment_id),
                    'original_text': clean_text,
                    'translated_text': '',
                    'is_translated': False,
                    'item_href': '0',
                    'element_type': element_type,
                    'paragraph_number': fragment_id,
                    'has_mismatch': False,
                    'processing_mode': 'inline',
                    'placeholder_pattern': placeholder_pattern,
                    'reserve_elements': [],
                    'prefix_reserve_tags': [],
                    'suffix_reserve_tags': [],
                    'is_non_translatable': False,
                }
                if inline_map:
                    para['inline_formatting_map'] = inline_map
            else:
                clean_text = (etree.tostring(elem, encoding='unicode', method='text') or '').strip()
                if not clean_text:
                    continue
                para = {
                    'id': str(fragment_id),
                    'original_text': clean_text,
                    'translated_text': '',
                    'is_translated': False,
                    'item_href': '0',
                    'element_type': element_type,
                    'paragraph_number': fragment_id,
                    'has_mismatch': False,
                    'processing_mode': 'legacy',
                    'original_html': self._build_html_original(elem, elem_id_str),
                    'reserve_elements': [],
                    'prefix_reserve_tags': [],
                    'suffix_reserve_tags': [],
                    'is_non_translatable': False,
                }
            if elem_id_attr:
                para['element_id'] = elem_id_attr
            if elem_style:
                para['kindle_style'] = elem_style
            paragraphs.append(para)
            fragment_id += 1

        element_type_dist = {}
        for p in paragraphs:
            et = p.get('element_type', 'unknown')
            element_type_dist[et] = element_type_dist.get(et, 0) + 1
        logger.debug(f"AZW3Processor._parse_html: total paragraphs={len(paragraphs)} element_type_dist={element_type_dist}")
        logger.debug(f"AZW3Processor._parse_html: self.image_items count={len(self.image_items)} ids={[img.get('id') for img in self.image_items]}")

        self._current_body_styles = {}
        return paragraphs

    def _html_extract_inline(self, element, placeholder_pattern: str) -> Tuple[str, Dict]:
        ATTRS_TO_KEEP = {'href', 'style', 'class', 'id', 'target', 'title', 'name'}
        VOID_TAGS = {'br', 'hr', 'wbr'}
        inline_formatting_map = {}
        counter = 0
        replacements = []
        for child in element.iter():
            child_tag = child.tag if isinstance(child.tag, str) else ''
            child_tag = child_tag.lower().split('}')[-1]
            if child_tag not in self.INLINE_TAGS:
                continue
            html_tag = self.INLINE_TAG_MAP.get(child_tag, child_tag)
            elem_id = counter
            opening = '<p_{:02d}>'.format(elem_id)
            closing = '</p_{:02d}>'.format(elem_id)
            filtered_attrs = {k: v for k, v in (child.attrib or {}).items() if k in ATTRS_TO_KEEP}
            inline_formatting_map[elem_id] = {
                'tag': html_tag,
                'attributes': filtered_attrs,
                'opening_placeholder': opening,
                'closing_placeholder': closing,
                'has_leading_space': False,
                'has_trailing_space': False,
            }
            replacements.append((child, opening, closing))
            counter += 1
        if not replacements:
            return html_module.escape(
                (etree.tostring(element, encoding='unicode', method='text') or '').strip()
            ), {}
        replace_map = {id(e): (op, cl) for e, op, cl in replacements}

        def serialize(node):
            parts = []
            nid = id(node)
            tag = node.tag if isinstance(node.tag, str) else ''
            tag_local = tag.lower().split('}')[-1]
            if nid in replace_map:
                op, cl = replace_map[nid]
                parts.append(op)
                if node.text:
                    parts.append(html_module.escape(node.text))
                for child in node:
                    parts.append(serialize(child))
                    if child.tail:
                        parts.append(html_module.escape(child.tail))
                parts.append(cl)
            elif tag_local in VOID_TAGS:
                parts.append(f'<{tag_local}/>')
            else:
                if node.text:
                    parts.append(html_module.escape(node.text))
                for child in node:
                    parts.append(serialize(child))
                    if child.tail:
                        parts.append(html_module.escape(child.tail))
            return ''.join(parts)

        return serialize(element).strip(), inline_formatting_map

    def _build_html_original(self, element, element_id: str) -> str:
        XHTML_NS = 'http://www.w3.org/1999/xhtml'
        INLINE_MAP = {'b': 'b', 'strong': 'b', 'i': 'em', 'em': 'em', 'u': 'u'}
        def _serialize(node) -> str:
            tag = node.tag if isinstance(node.tag, str) else ''
            tag_local = tag.lower().split('}')[-1]
            xhtml_tag = INLINE_MAP.get(tag_local)
            inner = ''
            if node.text:
                inner += html_module.escape(node.text)
            for child in node:
                inner += _serialize(child)
                if child.tail:
                    inner += html_module.escape(child.tail)
            if xhtml_tag:
                return f'<{xhtml_tag}>{inner}</{xhtml_tag}>'
            return inner
        parts = []
        if element.text:
            parts.append(html_module.escape(element.text))
        for child in element:
            parts.append(_serialize(child))
            if child.tail:
                parts.append(html_module.escape(child.tail))
        return f'<p xmlns="{XHTML_NS}" id="{element_id}">{"".join(parts)}</p>'

    def _has_any_text(self, element) -> bool:
        if element.text and element.text.strip():
            return True
        for descendant in element.iter():
            if descendant is element:
                continue
            if descendant.text and descendant.text.strip():
                return True
            if descendant.tail and descendant.tail.strip():
                return True
        return False

    def _process_element_inline(self, element, item_href: str, seen: set, list_context=None):
        tag_name = etree.QName(element).localname
        original_element_id = element.get('id')
        if original_element_id is None:
            element.set('id', f"trans_{uuid.uuid4()}")
        element_id = element.get('id')
        element_class = (element.get('class') or '').strip()
        element_copy = copy.deepcopy(element)
        for noise_tag in self.NOISE_TAGS:
            for noise_elem in self._local_xpath(element_copy, noise_tag):
                parent = noise_elem.getparent()
                if parent is not None:
                    parent.remove(noise_elem)
        reserve_elements = []
        reserve_counter = 0
        for reserve_tag in self.RESERVE_TAGS:
            for reserve_elem in self._local_xpath(element_copy, reserve_tag):
                reserve_elem_copy = copy.deepcopy(reserve_elem)
                reserve_elem_copy.tail = None
                reserve_html = etree.tostring(reserve_elem_copy, encoding='unicode', method='xml')
                reserve_elements.append(reserve_html)
                placeholder = self.RESERVE_PLACEHOLDER_PATTERN.format(reserve_counter)
                reserve_counter += 1
                parent = reserve_elem.getparent()
                if parent is not None:
                    tail = reserve_elem.tail or ''
                    prev = reserve_elem.getprevious()
                    if prev is not None:
                        prev.tail = (prev.tail or '') + placeholder + tail
                    else:
                        parent.text = (parent.text or '') + placeholder + tail
                    parent.remove(reserve_elem)
        self._remove_useless_spans(element_copy)
        inline_counter = reserve_counter
        clean_text, inline_formatting_map, _ = self._replace_inline_formatting_with_placeholders(
            element_copy, inline_counter
        )
        clean_text = self._cleanup_empty_placeholders(clean_text)
        clean_text, inline_formatting_map = self._flatten_placeholder_nesting(clean_text, inline_formatting_map)
        clean_text, non_translatable_map = self._extract_non_translatable_placeholders(
            clean_text, inline_formatting_map
        )
        clean_text = self._cleanup_empty_placeholders(clean_text)
        if not clean_text or not clean_text.strip():
            return
        prefix_tags, suffix_tags, clean_text = self._extract_boundary_reserve_tags(clean_text)
        is_non_translatable = self._is_non_translatable_content(clean_text)
        auto_wrap_tags = self._detect_auto_wrap_tags(clean_text, inline_formatting_map)
        if auto_wrap_tags:
            clean_text = self._strip_outer_placeholders(clean_text, auto_wrap_tags)
        used_ids = set(int(m) for m in re.findall(r'<p_(\d{2})>', clean_text))
        non_translatable_ids = set(non_translatable_map.keys())
        auto_wrap_ids = set(tag_info['elem_id'] for tag_info in auto_wrap_tags) if auto_wrap_tags else set()
        keep_ids = used_ids | non_translatable_ids | auto_wrap_ids
        inline_formatting_map = {k: v for k, v in inline_formatting_map.items() if k in keep_ids}
        key = (item_href, element_id)
        if key in seen:
            return
        seen.add(key)
        original_html = etree.tostring(element, encoding='unicode', method='xml', pretty_print=False)
        element_type = tag_name
        para = {
            'id': element_id,
            'original_text': clean_text,
            'translated_text': '',
            'is_translated': False,
            'item_href': item_href,
            'element_type': element_type,
            'original_html': original_html,
            'has_mismatch': False,
            'reserve_elements': reserve_elements,
            'inline_formatting_map': inline_formatting_map,
            'non_translatable_placeholders': non_translatable_map,
            'placeholder_pattern': self.PLACEHOLDER_PATTERN,
            'processing_mode': 'inline',
            'prefix_reserve_tags': prefix_tags,
            'suffix_reserve_tags': suffix_tags,
            'is_non_translatable': is_non_translatable,
        }
        if original_element_id:
            para['element_id'] = original_element_id
        if element_class:
            para['element_class'] = element_class
        style_str = self._extract_element_style(element)
        if style_str:
            para['kindle_style'] = style_str
        if auto_wrap_tags:
            para['auto_wrap_tags'] = auto_wrap_tags
        if list_context and list_context.get('list_id'):
            para['list_metadata'] = {
                'list_id': list_context['list_id'],
                'list_type': list_context.get('list_type', 'ul'),
                'list_class': list_context.get('list_class', ''),
                'li_index': list_context.get('li_index', 0),
                'li_class': list_context.get('li_class', ''),
                'parent_list_id': list_context.get('parent_list_id'),
                'parent_li_index': list_context.get('parent_li_index'),
            }
        self.paragraphs.append(para)

    def _process_element_legacy(self, element, item_href: str, seen: set, list_context=None):
        tag_name = etree.QName(element).localname
        original_element_id = element.get('id')
        if original_element_id is None:
            element.set('id', f"trans_{uuid.uuid4()}")
        element_id = element.get('id')
        element_class = (element.get('class') or '').strip()
        element_copy = copy.deepcopy(element)
        for noise_tag in self.NOISE_TAGS:
            for noise_elem in self._local_xpath(element_copy, noise_tag):
                parent = noise_elem.getparent()
                if parent is not None:
                    parent.remove(noise_elem)
        reserve_elements = []
        reserve_counter = 0
        for reserve_tag in self.RESERVE_TAGS:
            for reserve_elem in self._local_xpath(element_copy, reserve_tag):
                reserve_html = etree.tostring(reserve_elem, encoding='unicode', method='xml', with_tail=False)
                reserve_html = re.sub(r'\s+xmlns(?::\w+)?="[^"]*"', '', reserve_html)
                reserve_elements.append(reserve_html)
                placeholder = self.RESERVE_PLACEHOLDER_PATTERN.format(reserve_counter)
                reserve_counter += 1
                parent = reserve_elem.getparent()
                if parent is not None:
                    tail = reserve_elem.tail or ''
                    prev = reserve_elem.getprevious()
                    if prev is not None:
                        existing = prev.tail or ''
                        space_before = '' if (not existing or existing[-1].isspace()) else ' '
                        prev.tail = existing + space_before + placeholder + tail
                    else:
                        existing = parent.text or ''
                        space_before = '' if (not existing or existing[-1].isspace()) else ' '
                        parent.text = existing + space_before + placeholder + tail
                    parent.remove(reserve_elem)
        clean_text = (etree.tostring(element_copy, encoding='unicode', method='text') or '').strip()
        if not clean_text:
            return
        prefix_tags, suffix_tags, clean_text = self._extract_boundary_reserve_tags(clean_text)
        is_non_translatable = self._is_non_translatable_content(clean_text)
        key = (item_href, element_id)
        if key in seen:
            return
        seen.add(key)
        original_html = etree.tostring(element, encoding='unicode', method='xml', pretty_print=False)
        element_type = tag_name
        para = {
            'id': element_id,
            'original_text': clean_text,
            'translated_text': '',
            'is_translated': False,
            'item_href': item_href,
            'element_type': element_type,
            'original_html': original_html,
            'has_mismatch': False,
            'reserve_elements': reserve_elements,
            'placeholder_pattern': self.RESERVE_PLACEHOLDER_PATTERN,
            'processing_mode': 'legacy',
            'prefix_reserve_tags': prefix_tags,
            'suffix_reserve_tags': suffix_tags,
            'is_non_translatable': is_non_translatable,
        }
        if original_element_id:
            para['element_id'] = original_element_id
        if element_class:
            para['element_class'] = element_class
        style_str = self._extract_element_style(element)
        if style_str:
            para['kindle_style'] = style_str
        if list_context and list_context.get('list_id'):
            para['list_metadata'] = {
                'list_id': list_context['list_id'],
                'list_type': list_context.get('list_type', 'ul'),
                'list_class': list_context.get('list_class', ''),
                'li_index': list_context.get('li_index', 0),
                'li_class': list_context.get('li_class', ''),
                'parent_list_id': list_context.get('parent_list_id'),
                'parent_li_index': list_context.get('parent_li_index'),
            }
        self.paragraphs.append(para)

    def _extract_element_style(self, element) -> str:
        style_parts = {}
        elem_align = (element.get('align') or '').strip().lower()
        if elem_align in ('center', 'left', 'right', 'justify'):
            style_parts['text-align'] = elem_align
        for decl in (element.get('style') or '').split(';'):
            decl = decl.strip()
            if not decl or ':' not in decl:
                continue
            prop, _, val = decl.partition(':')
            prop = prop.strip().lower()
            if prop in self.KEEP_STYLE_PROPS:
                style_parts[prop] = val.strip()
        for cls in (element.get('class') or '').split():
            for prop, val in self.css_class_styles.get(cls, {}).items():
                if prop not in style_parts:
                    style_parts[prop] = val
        for prop, val in self._current_body_styles.items():
            if prop not in style_parts:
                style_parts[prop] = val
        return '; '.join(f'{k}:{v}' for k, v in style_parts.items())

    def _remove_useless_spans(self, element):
        spans_to_remove = []
        for span in element.xpath('.//*[local-name()="span"]'):
            has_meaningful_attr = any([
                span.get('class') is not None,
                span.get('style') is not None,
                span.get('id') is not None,
                span.get('lang') is not None,
                span.get('dir') is not None,
                span.get('{http://www.idpf.org/2007/ops}type') is not None,
            ])
            if has_meaningful_attr:
                continue
            text = (span.text or '')
            children = list(span)
            if not text.strip() and len(children) <= 1:
                spans_to_remove.append(span)
        for span in reversed(spans_to_remove):
            parent = span.getparent()
            if parent is None:
                continue
            span_index = list(parent).index(span)
            if span.text:
                prev = span.getprevious()
                if prev is not None:
                    prev.tail = (prev.tail or '') + span.text
                else:
                    parent.text = (parent.text or '') + span.text
            for child in reversed(list(span)):
                parent.insert(span_index, child)
            if span.tail:
                new_prev = span.getprevious()
                if new_prev is not None:
                    new_prev.tail = (new_prev.tail or '') + span.tail
                else:
                    parent.text = (parent.text or '') + span.tail
            parent.remove(span)

    def _replace_inline_formatting_with_placeholders(self, element, start_counter: int) -> Tuple[str, Dict, int]:
        _reserve_split = re.compile(r'(<id_\d{2}>)')

        def _esc(s: str) -> str:
            if not s:
                return s
            parts = _reserve_split.split(s)
            return ''.join(
                part if re.match(r'^<id_\d{2}>$', part) else html_module.escape(part)
                for part in parts
            )

        formatting_map = {}
        counter = start_counter
        replacements = []
        tags_to_process = [
            tag for tag in self.INLINE_FORMATTING_TAGS
            if not self.skip_inline_tags.get(tag, False)
        ]
        for tag in tags_to_process:
            for elem in self._local_xpath(element, tag):
                if tag == 'span':
                    epub_type = elem.get('{http://www.idpf.org/2007/ops}type')
                    if epub_type == 'pagebreak':
                        continue
                    if not elem.attrib:
                        text = (elem.text or '').strip()
                        if len(text) == 1 and text.isalpha():
                            parent = elem.getparent()
                            if parent is not None:
                                children = [c for c in parent if isinstance(c, etree._Element)]
                                if children and children[0] == elem:
                                    continue
                elem_id = counter
                opening = '<p_{:02d}>'.format(elem_id)
                closing = '</p_{:02d}>'.format(elem_id)
                tag_name = etree.QName(elem).localname
                attributes = {k: v for k, v in elem.attrib.items() if not k.startswith('{')}
                prev_sibling = elem.getprevious()
                if prev_sibling is not None:
                    preceding_text = prev_sibling.tail or ''
                else:
                    parent_elem = elem.getparent()
                    preceding_text = (parent_elem.text or '') if parent_elem is not None else ''
                has_leading_space = bool(preceding_text) and preceding_text[-1] in ' \t\n\r\u00a0'
                elem_tail = elem.tail or ''
                has_trailing_space = bool(elem_tail) and elem_tail[0] in ' \t\n\r\u00a0'
                formatting_map[elem_id] = {
                    'tag': tag_name,
                    'attributes': attributes,
                    'opening_placeholder': opening,
                    'closing_placeholder': closing,
                    'has_leading_space': has_leading_space,
                    'has_trailing_space': has_trailing_space,
                }
                replacements.append((elem, opening, closing, elem_id))
                counter += 1
        replace_map = {id(elem): (opening, closing) for elem, opening, closing, _ in replacements}

        def process_node(node):
            parts = []
            if id(node) in replace_map:
                opening, closing = replace_map[id(node)]
                parts.append(opening)
                if node.text:
                    parts.append(_esc(node.text))
                for child in node:
                    parts.append(process_node(child))
                    if child.tail:
                        parts.append(_esc(child.tail))
                parts.append(closing)
            else:
                if node.text:
                    parts.append(_esc(node.text))
                for child in node:
                    parts.append(process_node(child))
                    if child.tail:
                        parts.append(_esc(child.tail))
            return ''.join(parts)

        return process_node(element), formatting_map, counter

    def _cleanup_empty_placeholders(self, text: str) -> str:
        for _ in range(10):
            new_text = re.sub(r'<p_(\d{2})>\s*</p_\1>', '', text)
            if new_text == text:
                break
            text = new_text
        return text

    def _flatten_placeholder_nesting(self, text: str, formatting_map: Dict) -> Tuple[str, Dict]:
        pattern = r'<p_(\d{2})>\s*(<p_\d{2}>.*?</p_\d{2}>)\s*</p_\1>'
        removed_ids = set()
        for _ in range(5):
            matches = list(re.finditer(pattern, text, re.DOTALL))
            if not matches:
                break
            for match in reversed(matches):
                outer_id = int(match.group(1))
                inner_content = match.group(2)
                if outer_id not in formatting_map:
                    continue
                outer_info = formatting_map[outer_id]
                should_remove = False
                if outer_info['tag'] == 'span' and not outer_info['attributes']:
                    should_remove = True
                else:
                    inner_match = re.match(r'<p_(\d{2})>', inner_content)
                    if inner_match:
                        inner_id = int(inner_match.group(1))
                        if inner_id in formatting_map:
                            inner_info = formatting_map[inner_id]
                            if (outer_info['tag'] == inner_info['tag']
                                    and outer_info['attributes'] == inner_info['attributes']):
                                should_remove = True
                if should_remove:
                    text = text[:match.start()] + inner_content + text[match.end():]
                    removed_ids.add(outer_id)
        for tag_id in removed_ids:
            formatting_map.pop(tag_id, None)
        return text, formatting_map

    def _extract_non_translatable_placeholders(self, text: str, formatting_map: Dict) -> Tuple[str, Dict]:
        NON_TRANSLATABLE_PATTERN = r'^[\s\.,!?:;…]*$'
        non_translatable_map = {}
        def replace_with_marker(match):
            tag_id = int(match.group(1))
            content = match.group(2)
            if re.match(NON_TRANSLATABLE_PATTERN, content):
                non_translatable_map[tag_id] = {
                    'full_match': match.group(0),
                    'content': content,
                }
                return f'<nt_{tag_id:02d}/>'
            return match.group(0)
        clean_text = re.sub(r'<p_(\d{2})>(.*?)</p_\1>', replace_with_marker, text, flags=re.DOTALL)
        return clean_text, non_translatable_map

    def _detect_auto_wrap_tags(self, text: str, formatting_map: Dict) -> Optional[List[Dict]]:
        if not formatting_map:
            return None
        wrap_tags = []
        working_text = text.strip()
        while True:
            match = re.match(r'^<p_(\d{2})>(.*)</p_\1>$', working_text, re.DOTALL)
            if not match:
                break
            elem_id = int(match.group(1))
            inner_text = match.group(2)
            if elem_id not in formatting_map:
                break
            info = formatting_map[elem_id]
            wrap_tags.append({
                'elem_id': elem_id,
                'opening': info['opening_placeholder'],
                'closing': info['closing_placeholder'],
                'tag': info['tag'],
                'attributes': info['attributes'],
            })
            working_text = inner_text.strip()
        if wrap_tags and not re.search(r'</?p_\d{2}>', working_text):
            return wrap_tags
        return None

    def _strip_outer_placeholders(self, text: str, auto_wrap_tags: List[Dict]) -> str:
        working_text = text.strip()
        for tag_info in auto_wrap_tags:
            opening = tag_info['opening']
            closing = tag_info['closing']
            if working_text.startswith(opening) and working_text.endswith(closing):
                working_text = working_text[len(opening):-len(closing)].strip()
        return working_text

    def _extract_boundary_reserve_tags(self, text: str) -> Tuple[List[str], List[str], str]:
        prefix_tags = []
        suffix_tags = []
        clean_text = text
        tag_pattern = r'<id_\d{2}>'
        while True:
            clean_text = clean_text.lstrip()
            match = re.match(tag_pattern, clean_text)
            if match:
                prefix_tags.append(match.group(0))
                clean_text = clean_text[len(match.group(0)):]
            else:
                break
        while True:
            clean_text = clean_text.rstrip()
            match = re.search(tag_pattern + r'$', clean_text)
            if match:
                suffix_tags.insert(0, match.group(0))
                clean_text = clean_text[:-len(match.group(0))]
            else:
                break
        return prefix_tags, suffix_tags, clean_text.strip()

    def _is_non_translatable_content(self, text: str) -> bool:
        text_clean = re.sub(r'<id_\d{2}>', '', text)
        text_clean = re.sub(r'</?p_\d{2}>', '', text_clean).strip()
        if not text_clean:
            return True
        if re.match(r'^[\s\d\.,!?:;…\-–—\'\"\u201e\u201d\u201a\u2019]+$', text_clean):
            return True
        if re.match(r'^[\s\*•–—]+$', text_clean):
            return True
        if re.match(r'^([\*•–—])\s*(\1\s*)+$', text_clean):
            return True
        return False

class FB2Processor(FileProcessor):
    FB2_NS = 'http://www.gribuser.ru/xml/fictionbook/2.0'
 
    def __init__(self, app_settings: dict):
        self.app_settings = app_settings
        self.original_fb2_path = None
        self.image_items = []
        self.book_language = 'en'
 
    def get_file_type(self) -> str:
        return "fb2"
 
    def _build_fb2_original_html(self, element, element_id: str) -> str:
        XHTML_NS = 'http://www.w3.org/1999/xhtml'
        XLINK_NS = 'http://www.w3.org/1999/xlink'
        TAG_MAP = {
            f'{{{self.FB2_NS}}}emphasis': 'em',
            f'{{{self.FB2_NS}}}strong': 'b',
            f'{{{self.FB2_NS}}}strikethrough': 's',
            f'{{{self.FB2_NS}}}sub': 'sub',
            f'{{{self.FB2_NS}}}sup': 'sup',
            f'{{{self.FB2_NS}}}code': 'code',
            f'{{{self.FB2_NS}}}style': 'span',
        }
 
        def _serialize(node) -> str:
            local = node.tag.split('}')[-1] if '}' in node.tag else node.tag
            inner = ''
            if node.text:
                inner += html_module.escape(node.text)
            for child in node:
                inner += _serialize(child)
                if child.tail:
                    inner += html_module.escape(child.tail)
            xhtml_tag = TAG_MAP.get(node.tag)
            if xhtml_tag:
                return f'<{xhtml_tag}>{inner}</{xhtml_tag}>'
            if local == 'a':
                href = node.get(f'{{{XLINK_NS}}}href') or node.get('href', '')
                if href:
                    return f'<a href="{html_module.escape(href)}">{inner}</a>'
                return inner
            return inner
 
        parts = []
        if element.text:
            parts.append(html_module.escape(element.text))
        for child in element:
            parts.append(_serialize(child))
            if child.tail:
                parts.append(html_module.escape(child.tail))
        inner_html = ''.join(parts)
        return f'<p xmlns="{XHTML_NS}" id="{element_id}">{inner_html}</p>'
 
    def _extract_fb2_inline(self, element) -> tuple:
        XLINK_NS = 'http://www.w3.org/1999/xlink'
 
        FB2_TAG_TO_HTML = {
            f'{{{self.FB2_NS}}}emphasis': 'i',
            f'{{{self.FB2_NS}}}strong': 'b',
            f'{{{self.FB2_NS}}}strikethrough': 's',
            f'{{{self.FB2_NS}}}sub': 'sub',
            f'{{{self.FB2_NS}}}sup': 'sup',
            f'{{{self.FB2_NS}}}code': 'code',
            f'{{{self.FB2_NS}}}style': 'span',
        }
 
        inline_formatting_map = {}
        counter = 0
        replacements = []
 
        for child in element:
            local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
 
            if child.tag in FB2_TAG_TO_HTML:
                html_tag = FB2_TAG_TO_HTML[child.tag]
                attrs = {}
                if local == 'style':
                    name = child.get('name', '')
                    if name:
                        attrs['class'] = name
            elif local == 'a':
                html_tag = 'a'
                href = child.get(f'{{{XLINK_NS}}}href') or child.get('href', '')
                link_type = child.get(f'{{{XLINK_NS}}}type') or child.get('type', '')
                attrs = {}
                if href:
                    attrs['href'] = href
                if link_type:
                    attrs['data-type'] = link_type
            else:
                continue
 
            elem_id = counter
            opening = '<p_{:02d}>'.format(elem_id)
            closing = '</p_{:02d}>'.format(elem_id)
            prev = child.getprevious()
            preceding = (element.text or '') if prev is None else (prev.tail or '')
            inline_formatting_map[elem_id] = {
                'tag': html_tag,
                'attributes': attrs,
                'opening_placeholder': opening,
                'closing_placeholder': closing,
                'has_leading_space': bool(preceding) and preceding[-1] in ' \t',
                'has_trailing_space': bool(child.tail) and child.tail[0] in ' \t',
            }
            replacements.append((child, opening, closing))
            counter += 1
 
        if not replacements:
            return ''.join(element.itertext()).strip(), {}
 
        replace_map = {id(e): (op, cl) for e, op, cl in replacements}
 
        def serialize(node):
            parts = []
            nid = id(node)
            if nid in replace_map:
                op, cl = replace_map[nid]
                parts.append(op)
                if node.text:
                    parts.append(node.text)
                for child in node:
                    parts.append(serialize(child))
                    if child.tail:
                        parts.append(child.tail)
                parts.append(cl)
            else:
                if node.text:
                    parts.append(node.text)
                for child in node:
                    parts.append(serialize(child))
                    if child.tail:
                        parts.append(child.tail)
            return ''.join(parts)
 
        clean_text = serialize(element).strip()
        return clean_text, inline_formatting_map
 
    def load(self, path: str) -> Tuple[List[Dict], str]:
        try:
            import base64 as _base64

            self.original_fb2_path = path
            self.image_items = []
            self.book_language = 'en'
            paragraphs = []
            fragment_id = 0
            use_inline = self.app_settings.get('use_inline_formatting', True)
            placeholder_pattern = '<p_{:02d}>'

            with open(path, 'rb') as f:
                content = f.read()

            logger.debug(f"FB2Processor.load: path={path} content_size={len(content)}")

            if content[:2] == b'\x1f\x8b':
                import gzip
                content = gzip.decompress(content)
                logger.debug(f"FB2Processor.load: gzip decompressed, new content_size={len(content)}")

            tree = etree.fromstring(content)
            ns = {'fb': self.FB2_NS}
            XLINK_NS = 'http://www.w3.org/1999/xlink'

            logger.debug(f"FB2Processor.load: tree root tag={tree.tag}")

            IMAGE_CONTENT_TYPES = {
                'image/jpeg', 'image/png', 'image/gif',
                'image/bmp', 'image/webp', 'image/svg+xml',
            }

            lang_elem = tree.find(f'.//{{{self.FB2_NS}}}title-info/{{{self.FB2_NS}}}lang')
            if lang_elem is not None:
                lang_text = (lang_elem.text or '').strip()
                if lang_text:
                    self.book_language = lang_text

            logger.debug(f"FB2Processor.load: book_language={self.book_language}")

            cover_binary_id = None
            coverpage = tree.find(f'.//{{{self.FB2_NS}}}coverpage')
            if coverpage is not None:
                cover_img_elem = coverpage.find(f'{{{self.FB2_NS}}}image')
                if cover_img_elem is not None:
                    raw_href = (
                        cover_img_elem.get(f'{{{XLINK_NS}}}href')
                        or cover_img_elem.get('href', '')
                    )
                    if raw_href:
                        cover_binary_id = raw_href.lstrip('#')

            logger.debug(f"FB2Processor.load: cover_binary_id={cover_binary_id!r}")

            for binary in tree.findall(f'{{{self.FB2_NS}}}binary'):
                binary_id = binary.get('id', '')
                content_type = binary.get('content-type', 'image/jpeg')
                if content_type not in IMAGE_CONTENT_TYPES:
                    continue
                raw_data = (binary.text or '').strip()
                if not binary_id or not raw_data:
                    continue
                try:
                    img_bytes = _base64.b64decode(raw_data)
                    entry = {
                        'id': binary_id,
                        'file_name': binary_id,
                        'media_type': content_type,
                        'content': img_bytes,
                    }
                    if cover_binary_id and binary_id == cover_binary_id:
                        entry['is_cover'] = True
                    self.image_items.append(entry)
                except Exception as e:
                    logger.warning(f"FB2Processor: could not decode binary {binary_id}: {e}")

            logger.debug(f"FB2Processor.load: image_items count={len(self.image_items)}")
            logger.debug(f"FB2Processor.load: image_items ids={[img['id'] for img in self.image_items]}")
            logger.debug(f"FB2Processor.load: image_items is_cover={[img.get('is_cover') for img in self.image_items]}")

            title_info_elem = tree.find(f'.//{{{self.FB2_NS}}}title-info')
            if title_info_elem is not None:
                bt_elem = title_info_elem.find(f'{{{self.FB2_NS}}}book-title')
                if bt_elem is not None and (bt_elem.text or '').strip():
                    book_title_text = bt_elem.text.strip()
                    logger.debug(f"FB2Processor.load: book_title={book_title_text!r}")
                    paragraphs.append({
                        'id': str(fragment_id),
                        'original_text': book_title_text,
                        'translated_text': '',
                        'is_translated': False,
                        'item_href': 'metadata',
                        'element_type': 'book_title',
                        'paragraph_number': 0,
                        'has_mismatch': False,
                        'processing_mode': 'inline',
                        'placeholder_pattern': placeholder_pattern,
                        'reserve_elements': [],
                        'prefix_reserve_tags': [],
                        'suffix_reserve_tags': [],
                        'is_non_translatable': False,
                    })
                    fragment_id += 1
            else:
                logger.warning(f"FB2Processor.load: title-info NOT FOUND")

            def _base_para(fid, chapter_href, element_type, para_number):
                return {
                    'id': str(fid),
                    'original_text': '',
                    'translated_text': '',
                    'is_translated': False,
                    'item_href': chapter_href,
                    'element_type': element_type,
                    'paragraph_number': para_number,
                    'has_mismatch': False,
                    'processing_mode': 'inline',
                    'placeholder_pattern': placeholder_pattern,
                    'reserve_elements': [],
                    'prefix_reserve_tags': [],
                    'suffix_reserve_tags': [],
                    'is_non_translatable': False,
                }

            def _add_para(p_elem, chapter_href, element_type, para_number):
                nonlocal fragment_id
                elem_id_str = f'fb2_{fragment_id}'
                if use_inline:
                    clean_text, inline_map = self._extract_fb2_inline(p_elem)
                    if not clean_text:
                        return para_number
                    para = _base_para(fragment_id, chapter_href, element_type, para_number)
                    para['original_text'] = clean_text
                    if inline_map:
                        para['inline_formatting_map'] = inline_map
                else:
                    clean_text = ''.join(p_elem.itertext()).strip()
                    if not clean_text:
                        return para_number
                    para = _base_para(fragment_id, chapter_href, element_type, para_number)
                    para['original_text'] = clean_text
                    para['processing_mode'] = 'legacy'
                    para['original_html'] = self._build_fb2_original_html(p_elem, elem_id_str)
                paragraphs.append(para)
                fragment_id += 1
                return para_number + 1

            def _add_empty_line(chapter_href, para_number):
                nonlocal fragment_id
                para = _base_para(fragment_id, chapter_href, 'empty_line', para_number)
                para['is_non_translatable'] = True
                paragraphs.append(para)
                fragment_id += 1
                return para_number + 1

            def _add_image_frag(href, chapter_href, para_number):
                nonlocal fragment_id
                if not href:
                    return para_number
                paragraphs.append({
                    'id': str(fragment_id),
                    'original_text': '',
                    'translated_text': '',
                    'is_translated': False,
                    'item_href': chapter_href,
                    'element_type': 'image',
                    'image_href': href,
                    'paragraph_number': para_number,
                    'has_mismatch': False,
                })
                fragment_id += 1
                return para_number + 1

            def _add_table(table_elem, chapter_href, para_number):
                nonlocal fragment_id
                for row in table_elem.findall(f'{{{self.FB2_NS}}}tr'):
                    cells = [
                        c for c in row
                        if (c.tag.split('}')[-1] if '}' in c.tag else c.tag) in ('td', 'th')
                    ]
                    if not cells:
                        continue
                    row_parts = []
                    for cell in cells:
                        cell_text = ''.join(cell.itertext()).strip()
                        if cell_text:
                            row_parts.append(cell_text)
                    if not row_parts:
                        continue
                    para = _base_para(fragment_id, chapter_href, 'table_row', para_number)
                    para['original_text'] = ' | '.join(row_parts)
                    paragraphs.append(para)
                    fragment_id += 1
                    para_number += 1
                return para_number

            def _process_title(title_elem, chapter_href, para_number):
                p_children = list(title_elem.iter(f'{{{self.FB2_NS}}}p'))
                if p_children:
                    for p_elem in p_children:
                        para_number = _add_para(p_elem, chapter_href, 'heading', para_number)
                else:
                    text = ''.join(title_elem.itertext()).strip()
                    if text:
                        para_number = _add_para(title_elem, chapter_href, 'heading', para_number)
                return para_number

            def _process_epigraph(epi_elem, chapter_href, para_number):
                for p_elem in epi_elem.iter(f'{{{self.FB2_NS}}}p'):
                    para_number = _add_para(p_elem, chapter_href, 'blockquote', para_number)
                for ta_elem in epi_elem.findall(f'{{{self.FB2_NS}}}text-author'):
                    para_number = _add_para(ta_elem, chapter_href, 'blockquote', para_number)
                return para_number

            def _process_cite(cite_elem, chapter_href, para_number):
                for p_elem in cite_elem.iter(f'{{{self.FB2_NS}}}p'):
                    para_number = _add_para(p_elem, chapter_href, 'blockquote', para_number)
                for ta_elem in cite_elem.findall(f'{{{self.FB2_NS}}}text-author'):
                    para_number = _add_para(ta_elem, chapter_href, 'blockquote', para_number)
                return para_number

            def _process_poem(poem_elem, chapter_href, para_number):
                title_elem = poem_elem.find(f'{{{self.FB2_NS}}}title')
                if title_elem is not None:
                    para_number = _process_title(title_elem, chapter_href, para_number)
                for epi_elem in poem_elem.findall(f'{{{self.FB2_NS}}}epigraph'):
                    para_number = _process_epigraph(epi_elem, chapter_href, para_number)
                for stanza in poem_elem.findall(f'{{{self.FB2_NS}}}stanza'):
                    stanza_title = stanza.find(f'{{{self.FB2_NS}}}title')
                    if stanza_title is not None:
                        para_number = _process_title(stanza_title, chapter_href, para_number)
                    stanza_subtitle = stanza.find(f'{{{self.FB2_NS}}}subtitle')
                    if stanza_subtitle is not None:
                        para_number = _add_para(stanza_subtitle, chapter_href, 'heading', para_number)
                    for v_elem in stanza.findall(f'{{{self.FB2_NS}}}v'):
                        para_number = _add_para(v_elem, chapter_href, 'paragraph', para_number)
                for ta_elem in poem_elem.findall(f'{{{self.FB2_NS}}}text-author'):
                    para_number = _add_para(ta_elem, chapter_href, 'blockquote', para_number)
                return para_number

            def _process_child(child, chapter_href, para_number):
                nonlocal fragment_id
                local = child.tag.split('}')[-1] if '}' in child.tag else child.tag

                if local == 'image':
                    href = child.get(f'{{{XLINK_NS}}}href') or child.get('href', '')
                    return _add_image_frag(href, chapter_href, para_number)

                if local == 'title':
                    return _process_title(child, chapter_href, para_number)

                if local == 'subtitle':
                    return _add_para(child, chapter_href, 'heading', para_number)

                if local == 'empty-line':
                    return _add_empty_line(chapter_href, para_number)

                if local == 'epigraph':
                    return _process_epigraph(child, chapter_href, para_number)

                if local == 'cite':
                    return _process_cite(child, chapter_href, para_number)

                if local == 'poem':
                    return _process_poem(child, chapter_href, para_number)

                if local == 'table':
                    return _add_table(child, chapter_href, para_number)

                if local == 'section':
                    return para_number

                if local == 'p':
                    p_img_children = child.findall(f'{{{self.FB2_NS}}}image')
                    p_text = ''.join(child.itertext()).strip()
                    if p_img_children and not p_text:
                        for img_child in p_img_children:
                            img_href = img_child.get(f'{{{XLINK_NS}}}href') or img_child.get('href', '')
                            para_number = _add_image_frag(img_href, chapter_href, para_number)
                        return para_number
                    return _add_para(child, chapter_href, 'paragraph', para_number)

                return para_number

            def _process_section_children(section_elem, chapter_href, is_footnote_body):
                section_id = (section_elem.get('id') or '').strip()
                children_tags = [
                    (c.tag.split('}')[-1] if '}' in c.tag else c.tag)
                    for c in section_elem
                ]
                logger.debug(
                    f"FB2Processor._process_section_children: chapter_href={chapter_href!r} "
                    f"section_id={section_id!r} is_footnote={is_footnote_body} "
                    f"children_tags={children_tags}"
                )
                para_number = 0

                section_image = section_elem.find(f'{{{self.FB2_NS}}}image')
                if section_image is not None:
                    href = section_image.get(f'{{{XLINK_NS}}}href') or section_image.get('href', '')
                    para_number = _add_image_frag(href, chapter_href, para_number)

                for child in section_elem:
                    local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if local == 'section':
                        continue
                    if local == 'image' and child is section_image:
                        continue
                    if is_footnote_body and local == 'p':
                        p_text = ''.join(child.itertext()).strip()
                        p_img_children = child.findall(f'{{{self.FB2_NS}}}image')
                        if p_img_children and not p_text:
                            for img_child in p_img_children:
                                img_href = img_child.get(f'{{{XLINK_NS}}}href') or img_child.get('href', '')
                                para_number = _add_image_frag(img_href, chapter_href, para_number)
                        else:
                            para_number = _add_para(child, chapter_href, 'footnote', para_number)
                    else:
                        para_number = _process_child(child, chapter_href, para_number)

                logger.debug(
                    f"FB2Processor._process_section_children: chapter_href={chapter_href!r} "
                    f"finished para_number={para_number} total_paragraphs_so_far={len(paragraphs)}"
                )

            FOOTNOTE_BODY_NAMES = frozenset({
                'notes', 'footnotes', 'comments', 'note', 'footnote', 'comment',
            })
            BODY_CONTENT_LOCALS = frozenset({
                'title', 'p', 'image', 'epigraph', 'cite', 'poem',
                'subtitle', 'empty-line', 'table',
            })
            bodies = tree.findall('.//fb:body', ns)
            chapter_index = 0

            logger.debug(f"FB2Processor.load: bodies count={len(bodies)}")
            logger.debug(f"FB2Processor.load: body names={[(b.get('name') or '') for b in bodies]}")

            for body in bodies:
                body_name = (body.get('name') or '').lower()
                is_footnote_body = body_name in FOOTNOTE_BODY_NAMES

                has_body_content = any(
                    (child.tag.split('}')[-1] if '}' in child.tag else child.tag)
                    in BODY_CONTENT_LOCALS
                    for child in body
                )

                logger.debug(
                    f"FB2Processor.load: processing body name={body_name!r} "
                    f"is_footnote={is_footnote_body} has_body_content={has_body_content}"
                )

                if has_body_content:
                    chapter_href = f'chapter_{chapter_index}'
                    para_number = 0
                    fragments_before = len(paragraphs)
                    for child in body:
                        local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                        if local == 'section':
                            continue
                        para_number = _process_child(child, chapter_href, para_number)
                    if len(paragraphs) > fragments_before:
                        logger.debug(
                            f"FB2Processor.load: body {body_name!r} direct content -> "
                            f"chapter_href={chapter_href} extracted {len(paragraphs) - fragments_before} fragments"
                        )
                        chapter_index += 1

                section_list = list(body.iter(f'{{{self.FB2_NS}}}section'))
                logger.debug(
                    f"FB2Processor.load: body {body_name!r} sections count={len(section_list)} (including nested)"
                )

                for section in body.iter(f'{{{self.FB2_NS}}}section'):
                    chapter_href = f'chapter_{chapter_index}'
                    paras_before_section = len(paragraphs)
                    _process_section_children(section, chapter_href, is_footnote_body)
                    paras_after_section = len(paragraphs)
                    logger.debug(
                        f"FB2Processor.load: section chapter_href={chapter_href} -> "
                        f"extracted {paras_after_section - paras_before_section} fragments (total={paras_after_section})"
                    )
                    chapter_index += 1

            element_type_dist = {}
            for p in paragraphs:
                et = p.get('element_type', 'unknown')
                element_type_dist[et] = element_type_dist.get(et, 0) + 1
            logger.debug(
                f"FB2Processor.load: FINAL paragraphs={len(paragraphs)} "
                f"image_items={len(self.image_items)} element_type_dist={element_type_dist}"
            )
            image_hrefs_in_paras = [p.get('image_href') for p in paragraphs if p.get('element_type') == 'image']
            logger.debug(f"FB2Processor.load: image_hrefs in paragraphs={image_hrefs_in_paras}")

            logger.info(f"FB2Processor: loaded {len(paragraphs)} fragments, "
                        f"{len(self.image_items)} image items from {path}")
            return paragraphs, path

        except Exception as e:
            logger.error(f"FB2 load error: {e}", exc_info=True)
            raise


class DOCXProcessor(FileProcessor):
    def __init__(self, app_settings: dict):
        self.app_settings = app_settings
        self.original_docx_path = None
        self.image_items = []

    def get_file_type(self) -> str:
        return "docx"

    def _build_docx_original_html(self, para_obj, element_id: str) -> str:
        XHTML_NS = 'http://www.w3.org/1999/xhtml'
        runs = para_obj.runs
        if not runs:
            return f'<p xmlns="{XHTML_NS}" id="{element_id}">{html_module.escape(para_obj.text)}</p>'

        parts = []
        i = 0
        while i < len(runs):
            run = runs[i]
            is_bold = bool(run.bold)
            is_italic = bool(run.italic)
            is_underline = bool(run.underline)
            if not (is_bold or is_italic or is_underline):
                parts.append(html_module.escape(run.text))
                i += 1
                continue
            seg = run.text
            j = i + 1
            while j < len(runs):
                nxt = runs[j]
                if (bool(nxt.bold) == is_bold and
                        bool(nxt.italic) == is_italic and
                        bool(nxt.underline) == is_underline):
                    seg += nxt.text
                    j += 1
                else:
                    break
            escaped = html_module.escape(seg)
            if is_underline:
                escaped = f'<u>{escaped}</u>'
            if is_italic:
                escaped = f'<em>{escaped}</em>'
            if is_bold:
                escaped = f'<b>{escaped}</b>'
            parts.append(escaped)
            i = j

        return f'<p xmlns="{XHTML_NS}" id="{element_id}">{"".join(parts)}</p>'


    def _extract_docx_inline(self, para_obj) -> tuple:
        runs = para_obj.runs
        if not runs:
            return para_obj.text, {}

        if not any(r.bold or r.italic or r.underline for r in runs):
            return para_obj.text, {}

        inline_formatting_map = {}
        counter = 0
        parts = []
        i = 0

        while i < len(runs):
            run = runs[i]
            is_bold = bool(run.bold)
            is_italic = bool(run.italic)
            is_underline = bool(run.underline)

            if not (is_bold or is_italic or is_underline):
                parts.append(run.text)
                i += 1
                continue

            seg_text = run.text
            j = i + 1
            while j < len(runs):
                nxt = runs[j]
                if (bool(nxt.bold) == is_bold and
                        bool(nxt.italic) == is_italic and
                        bool(nxt.underline) == is_underline):
                    seg_text += nxt.text
                    j += 1
                else:
                    break

            if not seg_text:
                i = j
                continue

            inner = seg_text
            if is_underline:
                eid = counter
                op = '<p_{:02d}>'.format(eid)
                cl = '</p_{:02d}>'.format(eid)
                inline_formatting_map[eid] = {
                    'tag': 'u', 'attributes': {},
                    'opening_placeholder': op, 'closing_placeholder': cl,
                    'has_leading_space': False, 'has_trailing_space': False,
                }
                inner = op + inner + cl
                counter += 1
            if is_italic:
                eid = counter
                op = '<p_{:02d}>'.format(eid)
                cl = '</p_{:02d}>'.format(eid)
                inline_formatting_map[eid] = {
                    'tag': 'i', 'attributes': {},
                    'opening_placeholder': op, 'closing_placeholder': cl,
                    'has_leading_space': False, 'has_trailing_space': False,
                }
                inner = op + inner + cl
                counter += 1
            if is_bold:
                eid = counter
                op = '<p_{:02d}>'.format(eid)
                cl = '</p_{:02d}>'.format(eid)
                inline_formatting_map[eid] = {
                    'tag': 'b', 'attributes': {},
                    'opening_placeholder': op, 'closing_placeholder': cl,
                    'has_leading_space': False, 'has_trailing_space': False,
                }
                inner = op + inner + cl
                counter += 1

            parts.append(inner)
            i = j

        return ''.join(parts).strip(), inline_formatting_map

    def load(self, path: str) -> Tuple[List[Dict], str]:
        try:
            import zipfile
            from docx import Document
            from docx.oxml.ns import qn

            self.original_docx_path = path
            self.image_items = []
            paragraphs = []
            fragment_id = 0
            chapter_index = 0
            para_number = 0
            use_inline = self.app_settings.get('use_inline_formatting', True)
            placeholder_pattern = '<p_{:02d}>'

            media_type_map = {
                '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
                '.png': 'image/png', '.gif': 'image/gif',
                '.bmp': 'image/bmp', '.webp': 'image/webp',
                '.svg': 'image/svg+xml',
            }

            try:
                with zipfile.ZipFile(path, 'r') as z:
                    for name in z.namelist():
                        if name.startswith('word/media/'):
                            ext = os.path.splitext(name)[1].lower()
                            mt = media_type_map.get(ext, 'image/jpeg')
                            img_content = z.read(name)
                            basename = os.path.basename(name)
                            self.image_items.append({
                                'id': basename,
                                'file_name': basename,
                                'media_type': mt,
                                'content': img_content,
                            })
            except Exception as e:
                logger.warning(f"DOCXProcessor: could not extract images: {e}")

            doc = Document(path)

            for para in doc.paragraphs:
                has_image = False
                image_href = None
                try:
                    for run in para.runs:
                        drawing = run._element.find('.//' + qn('w:drawing'))
                        if drawing is not None:
                            blip = drawing.find('.//' + qn('a:blip'))
                            if blip is not None:
                                r_embed = blip.get(qn('r:embed'))
                                if r_embed and r_embed in doc.part.rels:
                                    rel = doc.part.rels[r_embed]
                                    image_href = os.path.basename(rel.target_ref)
                                    has_image = True
                                    break
                except Exception:
                    pass

                if has_image and image_href:
                    paragraphs.append({
                        'id': str(fragment_id),
                        'original_text': '',
                        'translated_text': '',
                        'is_translated': False,
                        'item_href': f'chapter_{chapter_index}',
                        'element_type': 'image',
                        'image_href': image_href,
                        'paragraph_number': para_number,
                        'has_mismatch': False,
                    })
                    fragment_id += 1
                    para_number += 1
                    continue

                style_name = para.style.name if para.style else ''
                if style_name.startswith('Heading'):
                    element_type = 'heading'
                    chapter_index += 1
                    para_number = 0
                else:
                    element_type = 'paragraph'

                elem_id_str = f'docx_{fragment_id}'

                if use_inline:
                    clean_text, inline_map = self._extract_docx_inline(para)
                    clean_text = clean_text.strip()
                    if not clean_text:
                        continue
                    entry = {
                        'id': str(fragment_id),
                        'original_text': clean_text,
                        'translated_text': '',
                        'is_translated': False,
                        'item_href': f'chapter_{chapter_index}',
                        'element_type': element_type,
                        'paragraph_number': para_number,
                        'has_mismatch': False,
                        'processing_mode': 'inline',
                        'placeholder_pattern': placeholder_pattern,
                        'reserve_elements': [],
                        'prefix_reserve_tags': [],
                        'suffix_reserve_tags': [],
                        'is_non_translatable': False,
                    }
                    if inline_map:
                        entry['inline_formatting_map'] = inline_map
                else:
                    clean_text = para.text.strip()
                    if not clean_text:
                        continue
                    entry = {
                        'id': str(fragment_id),
                        'original_text': clean_text,
                        'translated_text': '',
                        'is_translated': False,
                        'item_href': f'chapter_{chapter_index}',
                        'element_type': element_type,
                        'paragraph_number': para_number,
                        'has_mismatch': False,
                        'processing_mode': 'legacy',
                        'original_html': self._build_docx_original_html(para, elem_id_str),
                        'reserve_elements': [],
                        'prefix_reserve_tags': [],
                        'suffix_reserve_tags': [],
                        'is_non_translatable': False,
                    }

                paragraphs.append(entry)
                fragment_id += 1
                para_number += 1

            logger.info(f"DOCXProcessor: loaded {len(paragraphs)} fragments, "
                        f"{len(self.image_items)} image items from {path}")
            return paragraphs, path

        except Exception as e:
            logger.error(f"DOCX load error: {e}", exc_info=True)
            raise


# ─── Creators / Writers ──────────────────────────────────────────────────────

def _extract_text_align(kindle_style: str) -> str:
    for decl in (kindle_style or '').split(';'):
        decl = decl.strip()
        if ':' in decl:
            prop, _, val = decl.partition(':')
            if prop.strip() == 'text-align' and val.strip() in ('center', 'left', 'right', 'justify'):
                return val.strip()
    return ''

def _resolve_text(para: dict, keep_html: bool = False) -> str:
    def _apply_auto_wrap(result_text: str) -> str:
        if not keep_html or not result_text:
            return result_text
        auto_wrap = para.get('auto_wrap_tags') or []
        for tag_info in reversed(auto_wrap):
            tag = tag_info['tag']
            attrs = tag_info.get('attributes', {})
            if attrs:
                attrs_str = ' '.join(
                    f'{k}="{html_module.escape(str(v))}"' for k, v in attrs.items()
                )
                result_text = f'<{tag} {attrs_str}>{result_text}</{tag}>'
            else:
                result_text = f'<{tag}>{result_text}</{tag}>'
        return result_text

    if para.get('is_translated') and para.get('aligned_translated_html'):
        aligned_html = para['aligned_translated_html']
        if keep_html:
            try:
                elem = etree.fromstring(aligned_html.encode('utf-8'))
                inner = ''
                if elem.text:
                    inner += html_module.escape(elem.text)
                for child in elem:
                    child_str = etree.tostring(child, encoding='unicode', method='xml', with_tail=False)
                    child_str = re.sub(r'\s+xmlns(?::\w+)?="[^"]*"', '', child_str)
                    inner += child_str
                    if child.tail:
                        inner += html_module.escape(child.tail)
                return _apply_auto_wrap(inner)
            except Exception:
                pass
        clean = re.sub(r'<[^>]+>', '', aligned_html)
        return re.sub(r'\s+', ' ', clean).strip()

    text = (
        para.get('translated_text') if para.get('is_translated')
        else para.get('original_text', '')
    ) or ''

    text = re.sub(r'<nt_\d{2}/>', '', text)

    reserve_elements = para.get('reserve_elements') or []

    inline_map = para.get('inline_formatting_map')
    if inline_map and re.search(r'</?p_\d{2}>', text):
        if keep_html:
            tag_lookup = {}
            for elem_id, info in inline_map.items():
                tag = info['tag']
                attrs = info['attributes']
                if attrs:
                    attrs_str = ' '.join(
                        f'{k}="{html_module.escape(str(v))}"' for k, v in attrs.items()
                    )
                    tag_lookup[info['opening_placeholder']] = f'<{tag} {attrs_str}>'
                else:
                    tag_lookup[info['opening_placeholder']] = f'<{tag}>'
                tag_lookup[info['closing_placeholder']] = f'</{tag}>'

            def _expand_reserve(seg_text):
                seg_text = re.sub(r'</id_\d{2}>', '', seg_text)
                if not reserve_elements:
                    return html_module.escape(re.sub(r'<id_\d{2}>', '', seg_text))
                pieces = re.split(r'(<id_\d{2}>)', seg_text)
                out = []
                for piece in pieces:
                    m = re.match(r'<id_(\d{2})>', piece)
                    if m:
                        idx = int(m.group(1))
                        if idx < len(reserve_elements):
                            elem_html = reserve_elements[idx]
                            tag_match = re.match(r'<(\w+)', elem_html)
                            tag_n = tag_match.group(1).lower() if tag_match else ''
                            if tag_n != 'img':
                                clean_html = re.sub(r'\s+xmlns(?::\w+)?="[^"]*"', '', elem_html)
                                out.append(clean_html)
                    else:
                        out.append(html_module.escape(piece))
                return ''.join(out)

            parts = re.split(r'(</?p_\d{2}>)', text)
            result = []
            for part in parts:
                if part in tag_lookup:
                    result.append(tag_lookup[part])
                elif re.match(r'</?p_\d{2}>', part):
                    pass
                else:
                    result.append(_expand_reserve(part))
            return _apply_auto_wrap(''.join(result).strip())

        text = re.sub(r'</?p_\d{2}>', '', text)
        for idx in range(len(reserve_elements)):
            text = text.replace('<id_{:02d}>'.format(idx), '')
        text = re.sub(r'</?id_\d{2}>', '', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    if not reserve_elements:
        text = re.sub(r'</?id_\d{2}>', '', text)
        text = re.sub(r'</?p_\d{2}>', '', text)
        result = text.strip()
        return _apply_auto_wrap(result) if keep_html else result

    for idx, element_html in enumerate(reserve_elements):
        placeholder = '<id_{:02d}>'.format(idx)
        if placeholder not in text:
            continue
        if keep_html:
            tag_match = re.match(r'<(\w+)', element_html)
            tag_name = tag_match.group(1).lower() if tag_match else ''
            if tag_name == 'img':
                text = text.replace(placeholder, '')
            else:
                clean_html = re.sub(r'\s+xmlns(?::\w+)?="[^"]*"', '', element_html)
                text = text.replace(placeholder, clean_html)
        else:
            text = text.replace(placeholder, '')

    text = re.sub(r'</?id_\d{2}>', '', text)
    text = re.sub(r'</?p_\d{2}>', '', text)
    text = re.sub(r'  +', ' ', text)
    result = text.strip()
    return _apply_auto_wrap(result) if keep_html else result


def _int_color_to_css(color_int: int) -> str:
    r = (color_int >> 16) & 0xFF
    g = (color_int >> 8) & 0xFF
    b = color_int & 0xFF
    return f"#{r:02x}{g:02x}{b:02x}"

def _get_raw_text_for_inline(para: dict) -> str:
    text = (
        para.get('translated_text') if para.get('is_translated')
        else para.get('original_text', '')
    ) or ''

    reserve_elements = para.get('reserve_elements')
    if reserve_elements:
        placeholder_pattern = para.get('placeholder_pattern', '<id_{:02d}>')
        for idx in range(len(reserve_elements)):
            text = text.replace(placeholder_pattern.format(idx), '')

    text = re.sub(r'</?id_\d{2}>', '', text)
    text = re.sub(r'<nt_\d{2}/>', '', text)
    text = re.sub(r'  +', ' ', text)
    return text.strip()


def _build_inline_segments(text: str, inline_map: dict):
    segments = []
    stack = []
    pos = 0

    for m in re.finditer(r'<(/?p_(\d{2}))>', text):
        if m.start() > pos:
            seg = text[pos:m.start()]
            if seg:
                is_bold = any(inline_map.get(e, {}).get('tag') in ('b', 'strong') for e in stack)
                is_italic = any(inline_map.get(e, {}).get('tag') in ('i', 'em') for e in stack)
                is_underline = any(inline_map.get(e, {}).get('tag') == 'u' for e in stack)
                link_href = next(
                    (inline_map[e].get('attributes', {}).get('href', '')
                     for e in reversed(stack)
                     if inline_map.get(e, {}).get('tag') == 'a'),
                    ''
                )
                segments.append((seg, is_bold, is_italic, is_underline, link_href))

        eid = int(m.group(2))
        if m.group(0).startswith('</'):
            if stack and stack[-1] == eid:
                stack.pop()
        else:
            stack.append(eid)
        pos = m.end()

    if pos < len(text):
        seg = text[pos:]
        if seg:
            is_bold = any(inline_map.get(e, {}).get('tag') in ('b', 'strong') for e in stack)
            is_italic = any(inline_map.get(e, {}).get('tag') in ('i', 'em') for e in stack)
            is_underline = any(inline_map.get(e, {}).get('tag') == 'u' for e in stack)
            link_href = next(
                (inline_map[e].get('attributes', {}).get('href', '')
                 for e in reversed(stack)
                 if inline_map.get(e, {}).get('tag') == 'a'),
                ''
            )
            segments.append((seg, is_bold, is_italic, is_underline, link_href))

    return segments


def _fill_fb2_paragraph(p_elem, text: str, inline_map: dict, FB2_NS: str) -> None:
    if not inline_map or not re.search(r'</?p_\d{2}>', text):
        p_elem.text = text
        return
    _fill_fb2_from_segments(p_elem, _build_inline_segments(text, inline_map), FB2_NS)


def _fill_docx_paragraph(p, text: str, inline_map: dict) -> None:
    if not inline_map or not re.search(r'</?p_\d{2}>', text):
        p.add_run(text)
        return
    _fill_docx_from_segments(p, _build_inline_segments(text, inline_map))

def _parse_aligned_html_segments(aligned_html: str):
    BOLD_TAGS = {'b', 'strong'}
    ITALIC_TAGS = {'em', 'i'}
    UNDERLINE_TAGS = {'u'}

    try:
        elem = etree.fromstring(aligned_html.encode('utf-8'))
    except etree.XMLSyntaxError:
        try:
            wrapped = f'<_r xmlns="http://www.w3.org/1999/xhtml">{aligned_html}</_r>'
            elem = etree.fromstring(wrapped.encode('utf-8'))
        except Exception:
            clean = re.sub(r'<[^>]+>', '', aligned_html)
            return [(clean, False, False, False, '')] if clean.strip() else []

    segments = []

    def _walk(node, bold=False, italic=False, underline=False, link_href=''):
        if callable(node.tag):
            return
        t = etree.QName(node).localname.lower()
        b = bold or (t in BOLD_TAGS)
        i = italic or (t in ITALIC_TAGS)
        u = underline or (t in UNDERLINE_TAGS)
        href = link_href or (node.get('href', '') if t == 'a' else '')
        if node.text:
            segments.append((node.text, b, i, u, href))
        for child in node:
            _walk(child, b, i, u, href)
            if child.tail:
                segments.append((child.tail, bold, italic, underline, link_href))

    _walk(elem)
    return segments

def _fill_fb2_from_segments(p_elem, segments, FB2_NS: str) -> None:
    XLINK_NS = 'http://www.w3.org/1999/xlink'
    for item in segments:
        seg_text = item[0]
        bold = item[1]
        italic = item[2]
        link_href = item[4] if len(item) > 4 else ''
        if not seg_text:
            continue

        if link_href:
            a_elem = etree.SubElement(p_elem, f'{{{FB2_NS}}}a')
            a_elem.set(f'{{{XLINK_NS}}}href', link_href)
            if bold and italic:
                em = etree.SubElement(a_elem, f'{{{FB2_NS}}}emphasis')
                st = etree.SubElement(em, f'{{{FB2_NS}}}strong')
                st.text = seg_text
            elif italic:
                em = etree.SubElement(a_elem, f'{{{FB2_NS}}}emphasis')
                em.text = seg_text
            elif bold:
                st = etree.SubElement(a_elem, f'{{{FB2_NS}}}strong')
                st.text = seg_text
            else:
                a_elem.text = seg_text
        else:
            if bold and italic:
                em = etree.SubElement(p_elem, f'{{{FB2_NS}}}emphasis')
                st = etree.SubElement(em, f'{{{FB2_NS}}}strong')
                st.text = seg_text
            elif italic:
                em = etree.SubElement(p_elem, f'{{{FB2_NS}}}emphasis')
                em.text = seg_text
            elif bold:
                st = etree.SubElement(p_elem, f'{{{FB2_NS}}}strong')
                st.text = seg_text
            else:
                if len(p_elem) == 0:
                    p_elem.text = (p_elem.text or '') + seg_text
                else:
                    p_elem[-1].tail = (p_elem[-1].tail or '') + seg_text

def _fill_docx_from_segments(p, segments) -> None:
    for item in segments:
        seg_text = item[0]
        bold = item[1]
        italic = item[2]
        underline = item[3]
        if not seg_text:
            continue
        run = p.add_run(seg_text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if underline:
            run.underline = True

class PDFtoPDFCreator(QThread):
    finished = pyqtSignal(str, bool)

    def __init__(self, paragraphs, original_pdf_path, output_path):
        super().__init__()
        self.paragraphs = paragraphs
        self.original_pdf_path = original_pdf_path
        self.output_path = output_path

    def run(self):
        try:
            import fitz

            doc = fitz.open(self.original_pdf_path)
            pages_map: dict = {}
            for para in self.paragraphs:
                try:
                    page_num = int(para['item_href'])
                except (ValueError, KeyError):
                    continue
                pages_map.setdefault(page_num, []).append(para)

            for page_num, paras in pages_map.items():
                if page_num >= len(doc):
                    logger.warning(f"Page {page_num} out of range, skipping")
                    continue
                page = doc[page_num]
                translated = [p for p in paras if p.get('is_translated') and p.get('translated_text')]
                if not translated:
                    continue
                for para in translated:
                    bbox = para.get('bbox')
                    if not bbox or len(bbox) < 4:
                        continue
                    page.add_redact_annot(fitz.Rect(bbox))
                try:
                    page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)
                except (TypeError, AttributeError):
                    page.apply_redactions()
                for para in translated:
                    bbox = para.get('bbox')
                    if not bbox or len(bbox) < 4:
                        continue
                    meta = para.get('span_metadata', {})
                    font_size = meta.get('font_size', 12)
                    bold = meta.get('bold', False)
                    italic = meta.get('italic', False)
                    color_css = _int_color_to_css(meta.get('color', 0))
                    font_weight = "bold" if bold else "normal"
                    font_style = "italic" if italic else "normal"
                    escaped_text = html_module.escape(para['translated_text'])
                    html_text = (
                        f'<span style="font-size:{font_size}pt;'
                        f'font-weight:{font_weight};'
                        f'font-style:{font_style};'
                        f'color:{color_css};">'
                        f'{escaped_text}</span>'
                    )
                    page.insert_htmlbox(fitz.Rect(bbox), html_text, overlay=True)

            doc.save(self.output_path)
            doc.close()
            logger.info(f"PDFtoPDFCreator: saved to {self.output_path}")
            self.finished.emit(self.output_path, False)

        except Exception as e:
            logger.exception("PDFtoPDFCreator error")
            self.finished.emit(str(e), True)


class EPUBWriter(QThread):
 
    finished = pyqtSignal(str, bool)
 
    def __init__(self, paragraphs, output_path, source_title="Translated",
                 image_resources=None, source_format='epub', lang='en', css_class_styles=None):
        super().__init__()
        self.paragraphs = paragraphs
        self.output_path = output_path
        self.source_title = source_title
        self.image_resources = image_resources or []
        self.source_format = source_format
        self.lang = lang
        self.css_class_styles = css_class_styles or {}
 
    def run(self):
        dispatch = {
            'epub': self.run_from_epub,
            'fb2':  self.run_from_fb2,
            'docx': self.run_from_docx,
            'pdf':  self.run_from_pdf,
            'txt':  self.run_from_txt,
            'mobi': self.run_from_mobi,
            'azw':  self.run_from_azw,
            'azw3': self.run_from_azw3,
        }
        handler = dispatch.get(self.source_format, self.run_from_epub)
        handler()
 
    def run_from_epub(self):
        try:
            book = epub.EpubBook()

            title_frag = next(
                (p for p in self.paragraphs if p.get('element_type') == 'book_title'), None
            )
            if title_frag:
                resolved_title = _resolve_text(title_frag, keep_html=False)
                if resolved_title:
                    self.source_title = resolved_title

            book.set_title(self.source_title)
            book.set_language(self.lang)

            css_content = (
                'body { margin: 1em; }\n'
                'p { margin: 0.5em 0; text-indent: 1em; }\n'
                'h1, h2, h3, h4, h5, h6 { margin: 0.8em 0 0.4em; text-indent: 0; }\n'
                'blockquote { margin: 0.5em 2em; font-style: italic; }\n'
                '.footer { font-size: 0.8em; color: #666; }\n'
                'img { max-width: 100%; height: auto; display: block; margin: 0.5em auto; }\n'
                'a { text-decoration: underline; }\n'
            )
            nav_css = epub.EpubItem(
                uid='style_default',
                file_name='style/default.css',
                media_type='text/css',
                content=css_content.encode('utf-8'),
            )
            book.add_item(nav_css)

            img_lookup = {}
            for res in self.image_resources:
                raw_fn = res['file_name'] or ''
                file_name = raw_fn if '/' in raw_fn else 'images/' + raw_fn
                safe_id = re.sub(r'[^\w]', '_', res['id'])
                epub_img = epub.EpubItem(
                    uid='imgres_' + safe_id,
                    file_name=file_name,
                    media_type=res['media_type'],
                    content=res['content'],
                )
                book.add_item(epub_img)
                for key in [
                    res['id'],
                    '#' + res['id'],
                    raw_fn,
                    os.path.basename(raw_fn),
                    raw_fn.replace('\\', '/'),
                    os.path.normpath(raw_fn).replace('\\', '/'),
                ]:
                    if key and key not in img_lookup:
                        img_lookup[key] = file_name
                fn_parts = raw_fn.replace('\\', '/').split('/')
                if len(fn_parts) >= 2:
                    tail2 = '/'.join(fn_parts[-2:])
                    if tail2 not in img_lookup:
                        img_lookup[tail2] = file_name

            def _find_image(href):
                if not href:
                    return None
                candidates = [
                    href,
                    href.lstrip('#'),
                    os.path.basename(href),
                    href.replace('\\', '/'),
                    os.path.normpath(href).replace('\\', '/').lstrip('/'),
                ]
                parts = href.replace('\\', '/').split('/')
                if len(parts) >= 2:
                    candidates.append('/'.join(parts[-2:]))
                for c in candidates:
                    if c in img_lookup:
                        return img_lookup[c]
                return None

            cover_res = next((r for r in self.image_resources if r.get('is_cover')), None)
            if cover_res:
                safe_id = re.sub(r'[^\w]', '_', cover_res['id'])
                book.add_metadata('OPF', 'meta', '', {'name': 'cover', 'content': 'imgres_' + safe_id})

            pages_map = {}
            for para in self.paragraphs:
                if para.get('item_href') == 'metadata':
                    continue
                key = para.get('item_href', '0')
                pages_map.setdefault(key, []).append(para)

            chapters = []
            toc = []
            used_file_names = set()

            for idx, (chapter_key, paras) in enumerate(pages_map.items()):
                raw_name = (
                    os.path.splitext(os.path.basename(chapter_key))[0]
                    if chapter_key and chapter_key not in ('0', '')
                    else f'chapter_{idx}'
                )
                if not raw_name:
                    raw_name = f'chapter_{idx}'
                candidate = f'{raw_name}.xhtml'
                if candidate in used_file_names:
                    candidate = f'{raw_name}_{idx}.xhtml'
                used_file_names.add(candidate)
                chapter_file_name = candidate
                chapter_id = os.path.splitext(candidate)[0]

                chapter_title = None
                first_heading = next(
                    (p for p in paras if p.get('element_type') == 'heading'), None
                )
                if first_heading:
                    raw = _resolve_text(first_heading, keep_html=False)
                    if raw:
                        chapter_title = raw[:60]

                if not chapter_title:
                    first_text = next(
                        (p for p in paras
                         if p.get('element_type') not in ('image',)
                         and (p.get('translated_text') or p.get('original_text', ''))),
                        None
                    )
                    if first_text:
                        raw = _resolve_text(first_text, keep_html=False)
                        if raw:
                            words = raw.split()
                            chapter_title = ' '.join(words[:10])
                            if len(words) > 10:
                                chapter_title = chapter_title.rstrip(',.;:!?…') + '…'

                if not chapter_title:
                    chapter_title = f'Chapter {idx + 1}'

                body_parts = []
                for para in paras:
                    etype = para.get('element_type', 'paragraph')

                    if etype == 'image':
                        href = para.get('image_href', '')
                        file_name = _find_image(href)
                        if file_name:
                            body_parts.append(f'<img src="{file_name}" alt=""/>')
                        continue

                    text = _resolve_text(para, keep_html=True)
                    if not text:
                        continue
                    has_html = bool(re.search(r'<[^>]+>', text))
                    t = text if has_html else html_module.escape(text)
                    t = re.sub(r'(href="(?!https?://)(?!mailto:)[^"]*?)\.html((?:#[^"]*)?)"', r'\1.xhtml\2"', t)

                    style_attr = ''
                    ks = para.get('kindle_style', '')
                    if ks:
                        style_attr = f' style="{html_module.escape(ks)}"'

                    id_attr = ''
                    eid = para.get('element_id', '')
                    if eid:
                        id_attr = f' id="{html_module.escape(eid)}"'

                    if etype == 'heading':
                        body_parts.append(f'<h2{id_attr}{style_attr}>{t}</h2>')
                    elif etype == 'footer':
                        body_parts.append(f'<p class="footer"{id_attr}{style_attr}>{t}</p>')
                    elif etype == 'blockquote':
                        body_parts.append(f'<blockquote><p{id_attr}{style_attr}>{t}</p></blockquote>')
                    else:
                        body_parts.append(f'<p{id_attr}{style_attr}>{t}</p>')

                if not body_parts:
                    continue

                chapter_content = (
                    '<?xml version="1.0" encoding="utf-8"?>\n'
                    '<!DOCTYPE html>\n'
                    '<html xmlns="http://www.w3.org/1999/xhtml">\n'
                    '<head>\n'
                    f'  <title>{html_module.escape(chapter_title)}</title>\n'
                    '  <link rel="stylesheet" type="text/css" href="../style/default.css"/>\n'
                    '</head>\n'
                    '<body>\n'
                    + '\n'.join(body_parts) +
                    '\n</body>\n</html>'
                )

                chapter = epub.EpubHtml(
                    title=chapter_title,
                    file_name=chapter_file_name,
                    lang='en',
                )
                chapter.content = chapter_content.encode('utf-8')
                chapter.add_item(nav_css)
                book.add_item(chapter)
                chapters.append(chapter)
                toc.append(epub.Link(chapter_file_name, chapter_title, chapter_id))

            if not chapters:
                raise ValueError("No content to write")

            book.toc = toc
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())

            book.spine = chapters + ['nav']

            epub.write_epub(self.output_path, book)
            logger.info(f"EPUBWriter: saved to {self.output_path}")
            self.finished.emit(self.output_path, False)

        except Exception as e:
            logger.exception("EPUBWriter error")
            self.finished.emit(str(e), True)
 
    def run_from_fb2(self):
        try:
            book = epub.EpubBook()

            title_frag = next(
                (p for p in self.paragraphs if p.get('element_type') == 'book_title'), None
            )
            if title_frag:
                resolved_title = _resolve_text(title_frag, keep_html=False)
                if resolved_title:
                    self.source_title = resolved_title

            book.set_title(self.source_title)
            book.set_language(self.lang)

            logger.warning(f"EPUBWriter.run_from_fb2: source_title={self.source_title}")
            logger.warning(f"EPUBWriter.run_from_fb2: lang={self.lang}")
            logger.warning(f"EPUBWriter.run_from_fb2: paragraphs count={len(self.paragraphs)}")
            logger.warning(f"EPUBWriter.run_from_fb2: image_resources count={len(self.image_resources)}")
            logger.warning(f"EPUBWriter.run_from_fb2: image_resources ids={[r.get('id') for r in self.image_resources]}")
            logger.warning(f"EPUBWriter.run_from_fb2: image_resources file_names={[r.get('file_name') for r in self.image_resources]}")
            logger.warning(f"EPUBWriter.run_from_fb2: image_resources is_cover={[r.get('is_cover') for r in self.image_resources]}")

            css_content = (
                'body { margin: 1em; }\n'
                'p { margin: 0.5em 0; text-indent: 1em; }\n'
                'h1, h2, h3, h4, h5, h6 { margin: 0.8em 0 0.4em; text-indent: 0; }\n'
                'blockquote { margin: 0.5em 2em; font-style: italic; }\n'
                'p.empty-line { margin: 0.8em 0; text-indent: 0; }\n'
                'p.table-row { font-family: monospace; margin: 0.2em 0; text-indent: 0; }\n'
                'p.footnote { margin: 0.3em 0; text-indent: 0; font-size: 0.9em; }\n'
                '.footer { font-size: 0.8em; color: #666; }\n'
                'img { max-width: 100%; height: auto; display: block; margin: 0.5em auto; }\n'
                'a { text-decoration: underline; }\n'
                'sub { vertical-align: sub; font-size: 0.75em; }\n'
                'sup { vertical-align: super; font-size: 0.75em; }\n'
                's { text-decoration: line-through; }\n'
                'code { font-family: monospace; }\n'
            )
            nav_css = epub.EpubItem(
                uid='style_default',
                file_name='style/default.css',
                media_type='text/css',
                content=css_content.encode('utf-8'),
            )
            book.add_item(nav_css)

            cover_css_content = (
                'body { margin: 0; padding: 0; }\n'
                'div.cover-wrapper { width: 100%; height: 100%; text-align: center; }\n'
                'img.cover-image { max-width: 100%; height: auto; display: block; margin: 0 auto; }\n'
            )
            cover_css = epub.EpubItem(
                uid='style_cover',
                file_name='style/cover.css',
                media_type='text/css',
                content=cover_css_content.encode('utf-8'),
            )
            book.add_item(cover_css)

            img_lookup = {}
            for res in self.image_resources:
                raw_fn = res.get('file_name') or ''
                file_name = raw_fn if '/' in raw_fn else 'images/' + raw_fn
                safe_id = re.sub(r'[^\w]', '_', res['id'])
                epub_img = epub.EpubItem(
                    uid='imgres_' + safe_id,
                    file_name=file_name,
                    media_type=res['media_type'],
                    content=res['content'],
                )
                book.add_item(epub_img)
                for key in [
                    res['id'],
                    '#' + res['id'],
                    raw_fn,
                    os.path.basename(raw_fn),
                    raw_fn.replace('\\', '/'),
                    os.path.normpath(raw_fn).replace('\\', '/'),
                    raw_fn.lstrip('#'),
                ]:
                    if key and key not in img_lookup:
                        img_lookup[key] = file_name
                fn_parts = raw_fn.replace('\\', '/').split('/')
                if len(fn_parts) >= 2:
                    tail2 = '/'.join(fn_parts[-2:])
                    if tail2 not in img_lookup:
                        img_lookup[tail2] = file_name

            logger.warning(f"EPUBWriter.run_from_fb2: img_lookup keys={list(img_lookup.keys())}")

            def _find_image(href):
                if not href:
                    return None
                candidates = [
                    href,
                    href.lstrip('#'),
                    os.path.basename(href),
                    href.replace('\\', '/'),
                    os.path.normpath(href).replace('\\', '/').lstrip('/'),
                ]
                parts = href.replace('\\', '/').split('/')
                if len(parts) >= 2:
                    candidates.append('/'.join(parts[-2:]))
                for c in candidates:
                    if c in img_lookup:
                        logger.warning(f"EPUBWriter.run_from_fb2: _find_image href={href} -> matched candidate={c!r} -> {img_lookup[c]}")
                        return img_lookup[c]
                logger.warning(f"EPUBWriter.run_from_fb2: _find_image href={href} -> NOT FOUND candidates={candidates}")
                return None

            cover_res = next((r for r in self.image_resources if r.get('is_cover')), None)
            cover_chapter = None
            if cover_res:
                safe_id = re.sub(r'[^\w]', '_', cover_res['id'])
                book.add_metadata('OPF', 'meta', '', {'name': 'cover', 'content': 'imgres_' + safe_id})
                logger.warning(f"EPUBWriter.run_from_fb2: cover_res id={cover_res['id']} file_name={cover_res.get('file_name')}")

                raw_fn = cover_res.get('file_name') or ''
                cover_img_file = raw_fn if '/' in raw_fn else 'images/' + raw_fn
                cover_page_content = (
                    '<?xml version="1.0" encoding="utf-8"?>\n'
                    '<!DOCTYPE html>\n'
                    '<html xmlns="http://www.w3.org/1999/xhtml">\n'
                    '<head>\n'
                    '  <title>Cover</title>\n'
                    '  <link rel="stylesheet" type="text/css" href="style/cover.css"/>\n'
                    '</head>\n'
                    '<body>\n'
                    '<div class="cover-wrapper">\n'
                    f'  <img class="cover-image" src="{cover_img_file}" alt="Cover"/>\n'
                    '</div>\n'
                    '</body>\n</html>'
                )
                cover_chapter = epub.EpubHtml(
                    title='Cover',
                    file_name='cover.xhtml',
                    lang=self.lang,
                )
                cover_chapter.content = cover_page_content.encode('utf-8')
                cover_chapter.add_item(cover_css)
                book.add_item(cover_chapter)
                logger.warning(f"EPUBWriter.run_from_fb2: cover page created -> cover.xhtml with img={cover_img_file}")
            else:
                logger.warning(f"EPUBWriter.run_from_fb2: cover_res NOT FOUND in image_resources")

            pages_map = {}
            for para in self.paragraphs:
                if para.get('item_href') == 'metadata':
                    continue
                key = para.get('item_href', '0')
                pages_map.setdefault(key, []).append(para)

            logger.warning(f"EPUBWriter.run_from_fb2: pages_map keys={list(pages_map.keys())}")
            logger.warning(f"EPUBWriter.run_from_fb2: pages_map para counts={[(k, len(v)) for k, v in pages_map.items()]}")

            chapters = []
            toc = []
            used_file_names = set()
            used_file_names.add('cover.xhtml')

            for idx, (chapter_key, paras) in enumerate(pages_map.items()):
                raw_name = (
                    os.path.splitext(os.path.basename(chapter_key))[0]
                    if chapter_key and chapter_key not in ('0', '')
                    else f'chapter_{idx}'
                )
                if not raw_name:
                    raw_name = f'chapter_{idx}'
                candidate = f'{raw_name}.xhtml'
                if candidate in used_file_names:
                    candidate = f'{raw_name}_{idx}.xhtml'
                used_file_names.add(candidate)
                chapter_file_name = candidate
                chapter_id = os.path.splitext(candidate)[0]

                chapter_title = None
                first_heading = next(
                    (p for p in paras if p.get('element_type') == 'heading'), None
                )
                if first_heading:
                    raw = _resolve_text(first_heading, keep_html=False)
                    if raw:
                        chapter_title = raw[:60]

                if not chapter_title:
                    first_text = next(
                        (p for p in paras
                         if p.get('element_type') not in ('image', 'empty_line')
                         and (p.get('translated_text') or p.get('original_text', ''))),
                        None
                    )
                    if first_text:
                        raw = _resolve_text(first_text, keep_html=False)
                        if raw:
                            words = raw.split()
                            chapter_title = ' '.join(words[:10])
                            if len(words) > 10:
                                chapter_title = chapter_title.rstrip(',.;:!?…') + '…'

                if not chapter_title:
                    chapter_title = f'Chapter {idx + 1}'

                body_parts = []
                for para in paras:
                    etype = para.get('element_type', 'paragraph')

                    if etype == 'image':
                        href = para.get('image_href', '')
                        file_name = _find_image(href)
                        if file_name:
                            body_parts.append(f'<img src="{file_name}" alt=""/>')
                        continue

                    if etype == 'empty_line':
                        body_parts.append('<p class="empty-line">&#160;</p>')
                        continue

                    text = _resolve_text(para, keep_html=True)
                    if not text:
                        if etype == 'table_row':
                            continue
                        continue

                    has_html = bool(re.search(r'<[^>]+>', text))
                    t = text if has_html else html_module.escape(text)

                    style_attr = ''
                    ks = para.get('kindle_style', '')
                    if ks:
                        style_attr = f' style="{html_module.escape(ks)}"'

                    id_attr = ''
                    eid = para.get('element_id', '')
                    if eid:
                        id_attr = f' id="{html_module.escape(eid)}"'

                    if etype == 'heading':
                        body_parts.append(f'<h2{id_attr}{style_attr}>{t}</h2>')
                    elif etype == 'footer':
                        body_parts.append(f'<p class="footer"{id_attr}{style_attr}>{t}</p>')
                    elif etype == 'blockquote':
                        body_parts.append(f'<blockquote><p{id_attr}{style_attr}>{t}</p></blockquote>')
                    elif etype == 'table_row':
                        body_parts.append(f'<p class="table-row"{id_attr}{style_attr}>{t}</p>')
                    elif etype == 'footnote':
                        body_parts.append(f'<p class="footnote"{id_attr}{style_attr}>{t}</p>')
                    else:
                        body_parts.append(f'<p{id_attr}{style_attr}>{t}</p>')

                logger.warning(
                    f"EPUBWriter.run_from_fb2: chapter idx={idx} key={chapter_key!r} "
                    f"file={chapter_file_name!r} title={chapter_title!r} body_parts={len(body_parts)}"
                )

                if not body_parts:
                    logger.warning(f"EPUBWriter.run_from_fb2: chapter idx={idx} SKIPPED (no body_parts)")
                    continue

                chapter_content = (
                    '<?xml version="1.0" encoding="utf-8"?>\n'
                    '<!DOCTYPE html>\n'
                    '<html xmlns="http://www.w3.org/1999/xhtml">\n'
                    '<head>\n'
                    f'  <title>{html_module.escape(chapter_title)}</title>\n'
                    '  <link rel="stylesheet" type="text/css" href="style/default.css"/>\n'
                    '</head>\n'
                    '<body>\n'
                    + '\n'.join(body_parts) +
                    '\n</body>\n</html>'
                )

                chapter = epub.EpubHtml(
                    title=chapter_title,
                    file_name=chapter_file_name,
                    lang=self.lang,
                )
                chapter.content = chapter_content.encode('utf-8')
                chapter.add_item(nav_css)
                book.add_item(chapter)
                chapters.append(chapter)
                toc.append(epub.Link(chapter_file_name, chapter_title, chapter_id))

            logger.warning(f"EPUBWriter.run_from_fb2: total chapters created={len(chapters)}")

            if not chapters:
                raise ValueError("No content to write")

            book.toc = toc
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())

            spine_items = []
            if cover_chapter is not None:
                spine_items.append(cover_chapter)
            spine_items.extend(chapters)
            spine_items.append('nav')
            book.spine = spine_items

            epub.write_epub(self.output_path, book)
            logger.info(f"EPUBWriter: saved to {self.output_path}")
            self.finished.emit(self.output_path, False)

        except Exception as e:
            logger.exception("EPUBWriter error")
            self.finished.emit(str(e), True)
 
    def run_from_docx(self):
        try:
            book = epub.EpubBook()

            title_frag = next(
                (p for p in self.paragraphs if p.get('element_type') == 'book_title'), None
            )
            if title_frag:
                resolved_title = _resolve_text(title_frag, keep_html=False)
                if resolved_title:
                    self.source_title = resolved_title

            book.set_title(self.source_title)
            book.set_language(self.lang)

            css_content = (
                'body { margin: 1em; }\n'
                'p { margin: 0.5em 0; text-indent: 1em; }\n'
                'h1, h2, h3, h4, h5, h6 { margin: 0.8em 0 0.4em; text-indent: 0; }\n'
                'blockquote { margin: 0.5em 2em; font-style: italic; }\n'
                '.footer { font-size: 0.8em; color: #666; }\n'
                'img { max-width: 100%; height: auto; display: block; margin: 0.5em auto; }\n'
                'a { text-decoration: underline; }\n'
            )
            nav_css = epub.EpubItem(
                uid='style_default',
                file_name='style/default.css',
                media_type='text/css',
                content=css_content.encode('utf-8'),
            )
            book.add_item(nav_css)

            img_lookup = {}
            for res in self.image_resources:
                raw_fn = res['file_name'] or ''
                file_name = raw_fn if '/' in raw_fn else 'images/' + raw_fn
                safe_id = re.sub(r'[^\w]', '_', res['id'])
                epub_img = epub.EpubItem(
                    uid='imgres_' + safe_id,
                    file_name=file_name,
                    media_type=res['media_type'],
                    content=res['content'],
                )
                book.add_item(epub_img)
                for key in [
                    res['id'],
                    '#' + res['id'],
                    raw_fn,
                    os.path.basename(raw_fn),
                    raw_fn.replace('\\', '/'),
                    os.path.normpath(raw_fn).replace('\\', '/'),
                ]:
                    if key and key not in img_lookup:
                        img_lookup[key] = file_name
                fn_parts = raw_fn.replace('\\', '/').split('/')
                if len(fn_parts) >= 2:
                    tail2 = '/'.join(fn_parts[-2:])
                    if tail2 not in img_lookup:
                        img_lookup[tail2] = file_name

            def _find_image(href):
                if not href:
                    return None
                candidates = [
                    href,
                    href.lstrip('#'),
                    os.path.basename(href),
                    href.replace('\\', '/'),
                    os.path.normpath(href).replace('\\', '/').lstrip('/'),
                ]
                parts = href.replace('\\', '/').split('/')
                if len(parts) >= 2:
                    candidates.append('/'.join(parts[-2:]))
                for c in candidates:
                    if c in img_lookup:
                        return img_lookup[c]
                return None

            cover_res = next((r for r in self.image_resources if r.get('is_cover')), None)
            if cover_res:
                safe_id = re.sub(r'[^\w]', '_', cover_res['id'])
                book.add_metadata('OPF', 'meta', '', {'name': 'cover', 'content': 'imgres_' + safe_id})

            pages_map = {}
            for para in self.paragraphs:
                if para.get('item_href') == 'metadata':
                    continue
                key = para.get('item_href', '0')
                pages_map.setdefault(key, []).append(para)

            chapters = []
            toc = []
            used_file_names = set()

            for idx, (chapter_key, paras) in enumerate(pages_map.items()):
                raw_name = (
                    os.path.splitext(os.path.basename(chapter_key))[0]
                    if chapter_key and chapter_key not in ('0', '')
                    else f'chapter_{idx}'
                )
                if not raw_name:
                    raw_name = f'chapter_{idx}'
                candidate = f'{raw_name}.xhtml'
                if candidate in used_file_names:
                    candidate = f'{raw_name}_{idx}.xhtml'
                used_file_names.add(candidate)
                chapter_file_name = candidate
                chapter_id = os.path.splitext(candidate)[0]

                chapter_title = None
                first_heading = next(
                    (p for p in paras if p.get('element_type') == 'heading'), None
                )
                if first_heading:
                    raw = _resolve_text(first_heading, keep_html=False)
                    if raw:
                        chapter_title = raw[:60]

                if not chapter_title:
                    first_text = next(
                        (p for p in paras
                         if p.get('element_type') not in ('image',)
                         and (p.get('translated_text') or p.get('original_text', ''))),
                        None
                    )
                    if first_text:
                        raw = _resolve_text(first_text, keep_html=False)
                        if raw:
                            words = raw.split()
                            chapter_title = ' '.join(words[:10])
                            if len(words) > 10:
                                chapter_title = chapter_title.rstrip(',.;:!?…') + '…'

                if not chapter_title:
                    chapter_title = f'Chapter {idx + 1}'

                body_parts = []
                for para in paras:
                    etype = para.get('element_type', 'paragraph')

                    if etype == 'image':
                        href = para.get('image_href', '')
                        file_name = _find_image(href)
                        if file_name:
                            body_parts.append(f'<img src="{file_name}" alt=""/>')
                        continue

                    text = _resolve_text(para, keep_html=True)
                    if not text:
                        continue
                    has_html = bool(re.search(r'<[^>]+>', text))
                    t = text if has_html else html_module.escape(text)

                    style_attr = ''
                    ks = para.get('kindle_style', '')
                    if ks:
                        style_attr = f' style="{html_module.escape(ks)}"'

                    id_attr = ''
                    eid = para.get('element_id', '')
                    if eid:
                        id_attr = f' id="{html_module.escape(eid)}"'

                    if etype == 'heading':
                        body_parts.append(f'<h2{id_attr}{style_attr}>{t}</h2>')
                    elif etype == 'footer':
                        body_parts.append(f'<p class="footer"{id_attr}{style_attr}>{t}</p>')
                    elif etype == 'blockquote':
                        body_parts.append(f'<blockquote><p{id_attr}{style_attr}>{t}</p></blockquote>')
                    else:
                        body_parts.append(f'<p{id_attr}{style_attr}>{t}</p>')

                if not body_parts:
                    continue

                chapter_content = (
                    '<?xml version="1.0" encoding="utf-8"?>\n'
                    '<!DOCTYPE html>\n'
                    '<html xmlns="http://www.w3.org/1999/xhtml">\n'
                    '<head>\n'
                    f'  <title>{html_module.escape(chapter_title)}</title>\n'
                    '  <link rel="stylesheet" type="text/css" href="../style/default.css"/>\n'
                    '</head>\n'
                    '<body>\n'
                    + '\n'.join(body_parts) +
                    '\n</body>\n</html>'
                )

                chapter = epub.EpubHtml(
                    title=chapter_title,
                    file_name=chapter_file_name,
                    lang='en',
                )
                chapter.content = chapter_content.encode('utf-8')
                chapter.add_item(nav_css)
                book.add_item(chapter)
                chapters.append(chapter)
                toc.append(epub.Link(chapter_file_name, chapter_title, chapter_id))

            if not chapters:
                raise ValueError("No content to write")

            book.toc = toc
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())

            book.spine = chapters + ['nav']

            epub.write_epub(self.output_path, book)
            logger.info(f"EPUBWriter: saved to {self.output_path}")
            self.finished.emit(self.output_path, False)

        except Exception as e:
            logger.exception("EPUBWriter error")
            self.finished.emit(str(e), True)
 
    def run_from_pdf(self):
        try:
            book = epub.EpubBook()

            title_frag = next(
                (p for p in self.paragraphs if p.get('element_type') == 'book_title'), None
            )
            if title_frag:
                resolved_title = _resolve_text(title_frag, keep_html=False)
                if resolved_title:
                    self.source_title = resolved_title

            book.set_title(self.source_title)
            book.set_language(self.lang)

            css_content = (
                'body { margin: 1em; }\n'
                'p { margin: 0.5em 0; text-indent: 1em; }\n'
                'h1, h2, h3, h4, h5, h6 { margin: 0.8em 0 0.4em; text-indent: 0; }\n'
                'blockquote { margin: 0.5em 2em; font-style: italic; }\n'
                '.footer { font-size: 0.8em; color: #666; text-indent: 0; }\n'
                '.header { font-size: 0.8em; color: #666; text-indent: 0; }\n'
                '.list-item { margin-left: 2em; text-indent: -1em; }\n'
                'img { max-width: 100%; height: auto; display: block; margin: 0.5em auto; }\n'
                'a { text-decoration: underline; }\n'
            )
            nav_css = epub.EpubItem(
                uid='style_default',
                file_name='style/default.css',
                media_type='text/css',
                content=css_content.encode('utf-8'),
            )
            book.add_item(nav_css)

            img_lookup = {}
            for res in self.image_resources:
                raw_fn = res['file_name'] or ''
                file_name = raw_fn if '/' in raw_fn else 'images/' + raw_fn
                safe_id = re.sub(r'[^\w]', '_', res['id'])
                epub_img = epub.EpubItem(
                    uid='imgres_' + safe_id,
                    file_name=file_name,
                    media_type=res['media_type'],
                    content=res['content'],
                )
                book.add_item(epub_img)
                for key in [
                    res['id'],
                    '#' + res['id'],
                    raw_fn,
                    os.path.basename(raw_fn),
                    raw_fn.replace('\\', '/'),
                    os.path.normpath(raw_fn).replace('\\', '/'),
                ]:
                    if key and key not in img_lookup:
                        img_lookup[key] = file_name
                fn_parts = raw_fn.replace('\\', '/').split('/')
                if len(fn_parts) >= 2:
                    tail2 = '/'.join(fn_parts[-2:])
                    if tail2 not in img_lookup:
                        img_lookup[tail2] = file_name

            def _find_image(href):
                if not href:
                    return None
                candidates = [
                    href,
                    href.lstrip('#'),
                    os.path.basename(href),
                    href.replace('\\', '/'),
                    os.path.normpath(href).replace('\\', '/').lstrip('/'),
                ]
                parts = href.replace('\\', '/').split('/')
                if len(parts) >= 2:
                    candidates.append('/'.join(parts[-2:]))
                for c in candidates:
                    if c in img_lookup:
                        return img_lookup[c]
                return None

            def _apply_span_meta(text_html, meta):
                if not meta:
                    return text_html
                if meta.get('italic'):
                    text_html = f'<em>{text_html}</em>'
                if meta.get('bold'):
                    text_html = f'<b>{text_html}</b>'
                color = meta.get('color', 0)
                if color:
                    css_color = _int_color_to_css(color)
                    if css_color != '#000000':
                        text_html = f'<span style="color:{css_color};">{text_html}</span>'
                return text_html

            cover_res = next((r for r in self.image_resources if r.get('is_cover')), None)
            if cover_res:
                safe_id = re.sub(r'[^\w]', '_', cover_res['id'])
                book.add_metadata('OPF', 'meta', '', {'name': 'cover', 'content': 'imgres_' + safe_id})

            pages_map = {}
            for para in self.paragraphs:
                if para.get('item_href') == 'metadata':
                    continue
                key = para.get('item_href', '0')
                pages_map.setdefault(key, []).append(para)

            chapters = []
            toc = []
            used_file_names = set()

            for idx, (chapter_key, paras) in enumerate(pages_map.items()):
                raw_name = (
                    os.path.splitext(os.path.basename(chapter_key))[0]
                    if chapter_key and chapter_key not in ('0', '')
                    else f'chapter_{idx}'
                )
                if not raw_name:
                    raw_name = f'chapter_{idx}'
                candidate = f'{raw_name}.xhtml'
                if candidate in used_file_names:
                    candidate = f'{raw_name}_{idx}.xhtml'
                used_file_names.add(candidate)
                chapter_file_name = candidate
                chapter_id = os.path.splitext(candidate)[0]

                chapter_title = None
                first_heading = next(
                    (p for p in paras if p.get('element_type') == 'heading'), None
                )
                if first_heading:
                    raw = _resolve_text(first_heading, keep_html=False)
                    if raw:
                        chapter_title = raw[:60]

                if not chapter_title:
                    first_text = next(
                        (p for p in paras
                         if p.get('element_type') not in ('image',)
                         and (p.get('translated_text') or p.get('original_text', ''))),
                        None
                    )
                    if first_text:
                        raw = _resolve_text(first_text, keep_html=False)
                        if raw:
                            words = raw.split()
                            chapter_title = ' '.join(words[:10])
                            if len(words) > 10:
                                chapter_title = chapter_title.rstrip(',.;:!?…') + '…'

                if not chapter_title:
                    chapter_title = f'Chapter {idx + 1}'

                body_parts = []
                for para in paras:
                    etype = para.get('element_type', 'paragraph')

                    if etype == 'image':
                        href = para.get('image_href', '')
                        file_name = _find_image(href)
                        if file_name:
                            body_parts.append(f'<img src="{file_name}" alt=""/>')
                        continue

                    text = _resolve_text(para, keep_html=True)
                    if not text:
                        continue
                    has_html = bool(re.search(r'<[^>]+>', text))
                    t = text if has_html else html_module.escape(text)

                    t = _apply_span_meta(t, para.get('span_metadata'))

                    style_attr = ''
                    ks = para.get('kindle_style', '')
                    if ks:
                        style_attr = f' style="{html_module.escape(ks)}"'

                    id_attr = ''
                    eid = para.get('element_id', '')
                    if eid:
                        id_attr = f' id="{html_module.escape(eid)}"'

                    if etype == 'heading':
                        body_parts.append(f'<h2{id_attr}{style_attr}>{t}</h2>')
                    elif etype == 'footer':
                        body_parts.append(f'<p class="footer"{id_attr}{style_attr}>{t}</p>')
                    elif etype == 'header':
                        body_parts.append(f'<p class="header"{id_attr}{style_attr}>{t}</p>')
                    elif etype == 'list_item':
                        body_parts.append(f'<p class="list-item"{id_attr}{style_attr}>{t}</p>')
                    elif etype == 'blockquote':
                        body_parts.append(f'<blockquote><p{id_attr}{style_attr}>{t}</p></blockquote>')
                    else:
                        body_parts.append(f'<p{id_attr}{style_attr}>{t}</p>')

                if not body_parts:
                    continue

                chapter_content = (
                    '<?xml version="1.0" encoding="utf-8"?>\n'
                    '<!DOCTYPE html>\n'
                    '<html xmlns="http://www.w3.org/1999/xhtml">\n'
                    '<head>\n'
                    f'  <title>{html_module.escape(chapter_title)}</title>\n'
                    '  <link rel="stylesheet" type="text/css" href="../style/default.css"/>\n'
                    '</head>\n'
                    '<body>\n'
                    + '\n'.join(body_parts) +
                    '\n</body>\n</html>'
                )

                chapter = epub.EpubHtml(
                    title=chapter_title,
                    file_name=chapter_file_name,
                    lang=self.lang,
                )
                chapter.content = chapter_content.encode('utf-8')
                chapter.add_item(nav_css)
                book.add_item(chapter)
                chapters.append(chapter)
                toc.append(epub.Link(chapter_file_name, chapter_title, chapter_id))

            if not chapters:
                raise ValueError("No content to write")

            book.toc = toc
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())

            book.spine = chapters + ['nav']

            epub.write_epub(self.output_path, book)
            logger.info(f"EPUBWriter: saved to {self.output_path}")
            self.finished.emit(self.output_path, False)

        except Exception as e:
            logger.exception("EPUBWriter error")
            self.finished.emit(str(e), True)
 
    def run_from_txt(self):
        try:
            book = epub.EpubBook()

            title_frag = next(
                (p for p in self.paragraphs if p.get('element_type') == 'book_title'), None
            )
            if title_frag:
                resolved_title = _resolve_text(title_frag, keep_html=False)
                if resolved_title:
                    self.source_title = resolved_title

            book.set_title(self.source_title)
            book.set_language(self.lang)

            css_content = (
                'body { margin: 1em; }\n'
                'p { margin: 0.5em 0; text-indent: 1em; }\n'
                'h1, h2, h3, h4, h5, h6 { margin: 0.8em 0 0.4em; text-indent: 0; }\n'
                'blockquote { margin: 0.5em 2em; font-style: italic; }\n'
                '.footer { font-size: 0.8em; color: #666; }\n'
                'img { max-width: 100%; height: auto; display: block; margin: 0.5em auto; }\n'
                'a { text-decoration: underline; }\n'
            )
            nav_css = epub.EpubItem(
                uid='style_default',
                file_name='style/default.css',
                media_type='text/css',
                content=css_content.encode('utf-8'),
            )
            book.add_item(nav_css)

            img_lookup = {}
            for res in self.image_resources:
                raw_fn = res['file_name'] or ''
                file_name = raw_fn if '/' in raw_fn else 'images/' + raw_fn
                safe_id = re.sub(r'[^\w]', '_', res['id'])
                epub_img = epub.EpubItem(
                    uid='imgres_' + safe_id,
                    file_name=file_name,
                    media_type=res['media_type'],
                    content=res['content'],
                )
                book.add_item(epub_img)
                for key in [
                    res['id'],
                    '#' + res['id'],
                    raw_fn,
                    os.path.basename(raw_fn),
                    raw_fn.replace('\\', '/'),
                    os.path.normpath(raw_fn).replace('\\', '/'),
                ]:
                    if key and key not in img_lookup:
                        img_lookup[key] = file_name
                fn_parts = raw_fn.replace('\\', '/').split('/')
                if len(fn_parts) >= 2:
                    tail2 = '/'.join(fn_parts[-2:])
                    if tail2 not in img_lookup:
                        img_lookup[tail2] = file_name

            def _find_image(href):
                if not href:
                    return None
                candidates = [
                    href,
                    href.lstrip('#'),
                    os.path.basename(href),
                    href.replace('\\', '/'),
                    os.path.normpath(href).replace('\\', '/').lstrip('/'),
                ]
                parts = href.replace('\\', '/').split('/')
                if len(parts) >= 2:
                    candidates.append('/'.join(parts[-2:]))
                for c in candidates:
                    if c in img_lookup:
                        return img_lookup[c]
                return None

            cover_res = next((r for r in self.image_resources if r.get('is_cover')), None)
            if cover_res:
                safe_id = re.sub(r'[^\w]', '_', cover_res['id'])
                book.add_metadata('OPF', 'meta', '', {'name': 'cover', 'content': 'imgres_' + safe_id})

            pages_map = {}
            for para in self.paragraphs:
                if para.get('item_href') == 'metadata':
                    continue
                key = para.get('item_href', '0')
                pages_map.setdefault(key, []).append(para)

            chapters = []
            toc = []
            used_file_names = set()

            for idx, (chapter_key, paras) in enumerate(pages_map.items()):
                raw_name = (
                    os.path.splitext(os.path.basename(chapter_key))[0]
                    if chapter_key and chapter_key not in ('0', '')
                    else f'chapter_{idx}'
                )
                if not raw_name:
                    raw_name = f'chapter_{idx}'
                candidate = f'{raw_name}.xhtml'
                if candidate in used_file_names:
                    candidate = f'{raw_name}_{idx}.xhtml'
                used_file_names.add(candidate)
                chapter_file_name = candidate
                chapter_id = os.path.splitext(candidate)[0]

                chapter_title = None
                first_heading = next(
                    (p for p in paras if p.get('element_type') == 'heading'), None
                )
                if first_heading:
                    raw = _resolve_text(first_heading, keep_html=False)
                    if raw:
                        chapter_title = raw[:60]

                if not chapter_title:
                    first_text = next(
                        (p for p in paras
                         if p.get('element_type') not in ('image',)
                         and (p.get('translated_text') or p.get('original_text', ''))),
                        None
                    )
                    if first_text:
                        raw = _resolve_text(first_text, keep_html=False)
                        if raw:
                            words = raw.split()
                            chapter_title = ' '.join(words[:10])
                            if len(words) > 10:
                                chapter_title = chapter_title.rstrip(',.;:!?…') + '…'

                if not chapter_title:
                    chapter_title = f'Chapter {idx + 1}'

                body_parts = []
                for para in paras:
                    etype = para.get('element_type', 'paragraph')

                    if etype == 'image':
                        href = para.get('image_href', '')
                        file_name = _find_image(href)
                        if file_name:
                            body_parts.append(f'<img src="{file_name}" alt=""/>')
                        continue

                    text = _resolve_text(para, keep_html=True)
                    if not text:
                        continue
                    has_html = bool(re.search(r'<[^>]+>', text))
                    t = text if has_html else html_module.escape(text)

                    style_attr = ''
                    ks = para.get('kindle_style', '')
                    if ks:
                        style_attr = f' style="{html_module.escape(ks)}"'

                    id_attr = ''
                    eid = para.get('element_id', '')
                    if eid:
                        id_attr = f' id="{html_module.escape(eid)}"'

                    if etype == 'heading':
                        body_parts.append(f'<h2{id_attr}{style_attr}>{t}</h2>')
                    elif etype == 'footer':
                        body_parts.append(f'<p class="footer"{id_attr}{style_attr}>{t}</p>')
                    elif etype == 'blockquote':
                        body_parts.append(f'<blockquote><p{id_attr}{style_attr}>{t}</p></blockquote>')
                    else:
                        body_parts.append(f'<p{id_attr}{style_attr}>{t}</p>')

                if not body_parts:
                    continue

                chapter_content = (
                    '<?xml version="1.0" encoding="utf-8"?>\n'
                    '<!DOCTYPE html>\n'
                    '<html xmlns="http://www.w3.org/1999/xhtml">\n'
                    '<head>\n'
                    f'  <title>{html_module.escape(chapter_title)}</title>\n'
                    '  <link rel="stylesheet" type="text/css" href="../style/default.css"/>\n'
                    '</head>\n'
                    '<body>\n'
                    + '\n'.join(body_parts) +
                    '\n</body>\n</html>'
                )

                chapter = epub.EpubHtml(
                    title=chapter_title,
                    file_name=chapter_file_name,
                    lang='en',
                )
                chapter.content = chapter_content.encode('utf-8')
                chapter.add_item(nav_css)
                book.add_item(chapter)
                chapters.append(chapter)
                toc.append(epub.Link(chapter_file_name, chapter_title, chapter_id))

            if not chapters:
                raise ValueError("No content to write")

            book.toc = toc
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())

            book.spine = chapters + ['nav']

            epub.write_epub(self.output_path, book)
            logger.info(f"EPUBWriter: saved to {self.output_path}")
            self.finished.emit(self.output_path, False)

        except Exception as e:
            logger.exception("EPUBWriter error")
            self.finished.emit(str(e), True)
 
    def run_from_mobi(self):
        try:
            book = epub.EpubBook()
            title_frag = next(
                (p for p in self.paragraphs if p.get('element_type') == 'book_title'), None
            )
            if title_frag:
                resolved_title = _resolve_text(title_frag, keep_html=False)
                if resolved_title:
                    self.source_title = resolved_title
            book.set_title(self.source_title)
            book.set_language(self.lang)
            logger.warning(f"EPUBWriter.run_from_mobi: source_title={self.source_title}")
            logger.warning(f"EPUBWriter.run_from_mobi: paragraphs count={len(self.paragraphs)}")
            logger.warning(f"EPUBWriter.run_from_mobi: image_resources count={len(self.image_resources)}")
            logger.warning(f"EPUBWriter.run_from_mobi: image_resources ids={[r.get('id') for r in self.image_resources]}")
            logger.warning(f"EPUBWriter.run_from_mobi: image_resources file_names={[r.get('file_name') for r in self.image_resources]}")
            logger.warning(f"EPUBWriter.run_from_mobi: image_resources is_cover={[r.get('is_cover') for r in self.image_resources]}")
            css_content = (
                'body { margin: 1em; }\n'
                'p { margin: 0.3em 0; }\n'
                'h1, h2, h3, h4, h5, h6 { margin: 0.8em 0 0.4em; }\n'
                'blockquote { margin: 0.5em 2em; font-style: italic; }\n'
                '.footer { font-size: 0.8em; color: #666; }\n'
                'img { max-width: 100%; height: auto; display: block; margin: 0.5em auto; }\n'
                'a { text-decoration: underline; }\n'
            )
            nav_css = epub.EpubItem(
                uid='style_default',
                file_name='style/default.css',
                media_type='text/css',
                content=css_content.encode('utf-8'),
            )
            book.add_item(nav_css)
            img_lookup = {}
            for res in self.image_resources:
                raw_fn = res['file_name'] or ''
                file_name = raw_fn if '/' in raw_fn else 'images/' + raw_fn
                safe_id = re.sub(r'[^\w]', '_', res['id'])
                epub_img = epub.EpubItem(
                    uid='imgres_' + safe_id,
                    file_name=file_name,
                    media_type=res['media_type'],
                    content=res['content'],
                )
                book.add_item(epub_img)
                for key in [
                    res['id'],
                    '#' + res['id'],
                    raw_fn,
                    os.path.basename(raw_fn),
                    raw_fn.replace('\\', '/'),
                    os.path.normpath(raw_fn).replace('\\', '/'),
                ]:
                    if key and key not in img_lookup:
                        img_lookup[key] = file_name
                fn_parts = raw_fn.replace('\\', '/').split('/')
                if len(fn_parts) >= 2:
                    tail2 = '/'.join(fn_parts[-2:])
                    if tail2 not in img_lookup:
                        img_lookup[tail2] = file_name
            logger.warning(f"EPUBWriter.run_from_mobi: img_lookup keys={list(img_lookup.keys())}")
            def _find_image(href):
                if not href:
                    return None
                candidates = [
                    href,
                    href.lstrip('#'),
                    os.path.basename(href),
                    href.replace('\\', '/'),
                    os.path.normpath(href).replace('\\', '/').lstrip('/'),
                ]
                parts = href.replace('\\', '/').split('/')
                if len(parts) >= 2:
                    candidates.append('/'.join(parts[-2:]))
                for c in candidates:
                    if c in img_lookup:
                        logger.warning(f"EPUBWriter.run_from_mobi: _find_image href={href} -> matched candidate={c!r} -> {img_lookup[c]}")
                        return img_lookup[c]
                logger.warning(f"EPUBWriter.run_from_mobi: _find_image href={href} -> NOT FOUND candidates={candidates}")
                return None
            cover_chapter = None
            cover_res = next((r for r in self.image_resources if r.get('is_cover')), None)
            if cover_res:
                logger.warning(f"EPUBWriter.run_from_mobi: cover_res id={cover_res['id']} file_name={cover_res.get('file_name')}")
                safe_id = re.sub(r'[^\w]', '_', cover_res['id'])
                book.add_metadata('OPF', 'meta', '', {'name': 'cover', 'content': 'imgres_' + safe_id})
                raw_fn = cover_res['file_name'] or ''
                cover_file = raw_fn if '/' in raw_fn else 'images/' + raw_fn
                cover_html = (
                    '<?xml version="1.0" encoding="utf-8"?>\n'
                    '<!DOCTYPE html>\n'
                    '<html xmlns="http://www.w3.org/1999/xhtml">\n'
                    '<head><title>Cover</title>'
                    '<link rel="stylesheet" type="text/css" href="../style/default.css"/></head>\n'
                    '<body style="margin:0;padding:0;text-align:center">\n'
                    f'<div><img src="{cover_file}" alt="Cover" style="max-width:100%;height:auto"/></div>\n'
                    '</body>\n</html>'
                )
                cover_chapter = epub.EpubHtml(title='Cover', file_name='cover.xhtml', lang='en')
                cover_chapter.content = cover_html.encode('utf-8')
                cover_chapter.add_item(nav_css)
                book.add_item(cover_chapter)
            else:
                logger.warning(f"EPUBWriter.run_from_mobi: cover_res NOT FOUND in image_resources")
            pages_map = {}
            for para in self.paragraphs:
                if para.get('item_href') == 'metadata':
                    continue
                key = para.get('item_href', '0')
                pages_map.setdefault(key, []).append(para)
            logger.warning(f"EPUBWriter.run_from_mobi: pages_map keys={list(pages_map.keys())}")
            logger.warning(f"EPUBWriter.run_from_mobi: pages_map para counts={[(k, len(v)) for k, v in pages_map.items()]}")
            id_to_chapter = {}
            chapter_file_map = {}
            used_file_names_pre = set()
            for idx, (chapter_key, paras) in enumerate(pages_map.items()):
                raw_name = (
                    os.path.splitext(os.path.basename(chapter_key))[0]
                    if chapter_key and chapter_key not in ('0', '')
                    else f'chapter_{idx}'
                )
                if not raw_name:
                    raw_name = f'chapter_{idx}'
                candidate = f'{raw_name}.xhtml'
                if candidate in used_file_names_pre:
                    candidate = f'{raw_name}_{idx}.xhtml'
                used_file_names_pre.add(candidate)
                chapter_file_map[chapter_key] = candidate
                for para in paras:
                    eid = para.get('element_id', '')
                    if eid:
                        id_to_chapter[eid] = candidate
                    for extra_id in para.get('extra_anchor_ids', []):
                        id_to_chapter[extra_id] = candidate
                    for info in (para.get('inline_formatting_map') or {}).values():
                        if info.get('tag') == 'a':
                            for attr_name in ('id', 'name'):
                                val = info.get('attributes', {}).get(attr_name, '').strip()
                                if val:
                                    id_to_chapter[val] = candidate
            chapters = []
            toc = []
            for idx, (chapter_key, paras) in enumerate(pages_map.items()):
                chapter_file_name = chapter_file_map[chapter_key]
                chapter_id = os.path.splitext(chapter_file_name)[0]
                chapter_title = None
                first_heading = next(
                    (p for p in paras if p.get('element_type') == 'heading'), None
                )
                if first_heading:
                    raw = _resolve_text(first_heading, keep_html=False)
                    if raw:
                        chapter_title = raw[:60]
                if not chapter_title:
                    first_text = next(
                        (p for p in paras
                         if p.get('element_type') not in ('image',)
                         and (p.get('translated_text') or p.get('original_text', ''))),
                        None
                    )
                    if first_text:
                        raw = _resolve_text(first_text, keep_html=False)
                        if raw:
                            words = raw.split()
                            chapter_title = ' '.join(words[:10])
                            if len(words) > 10:
                                chapter_title = chapter_title.rstrip(',.;:!?…') + '…'
                if not chapter_title:
                    chapter_title = f'Chapter {idx + 1}'
                body_parts = []
                for para in paras:
                    etype = para.get('element_type', 'paragraph')
                    extra_anchors_html = ''.join(
                        f'<a id="{html_module.escape(aid)}"></a>'
                        for aid in para.get('extra_anchor_ids', [])
                    )
                    if etype == 'image':
                        href = para.get('image_href', '')
                        file_name = _find_image(href)
                        if file_name:
                            if extra_anchors_html:
                                body_parts.append(f'<p>{extra_anchors_html}</p>')
                            body_parts.append(f'<img src="{file_name}" alt=""/>')
                        continue
                    text = _resolve_text(para, keep_html=True)
                    if not text:
                        if extra_anchors_html:
                            body_parts.append(f'<p>{extra_anchors_html}</p>')
                        continue
                    has_html = bool(re.search(r'<[^>]+>', text))
                    t = text if has_html else html_module.escape(text)
                    style_attr = ''
                    ks = para.get('kindle_style', '')
                    if ks:
                        style_attr = f' style="{html_module.escape(ks)}"'
                    id_attr = ''
                    eid = para.get('element_id', '')
                    if eid:
                        id_attr = f' id="{html_module.escape(eid)}"'
                    if etype == 'heading':
                        hlevel = para.get('heading_level', 'h2')
                        body_parts.append(f'<{hlevel}{id_attr}{style_attr}>{extra_anchors_html}{t}</{hlevel}>')
                    elif etype == 'footer':
                        body_parts.append(f'<p class="footer"{id_attr}{style_attr}>{extra_anchors_html}{t}</p>')
                    elif etype == 'blockquote':
                        body_parts.append(f'<blockquote><p{id_attr}{style_attr}>{extra_anchors_html}{t}</p></blockquote>')
                    else:
                        body_parts.append(f'<p{id_attr}{style_attr}>{extra_anchors_html}{t}</p>')
                logger.warning(
                    f"EPUBWriter.run_from_mobi: chapter idx={idx} key={chapter_key!r} "
                    f"file={chapter_file_name!r} title={chapter_title!r} body_parts={len(body_parts)}"
                )
                if not body_parts:
                    logger.warning(f"EPUBWriter.run_from_mobi: chapter idx={idx} SKIPPED (no body_parts)")
                    continue
                def _fix_href(m, _current=chapter_file_name):
                    fragment = m.group(1)
                    target_file = id_to_chapter.get(fragment)
                    if target_file and target_file != _current:
                        return f'href="{target_file}#{fragment}"'
                    return m.group(0)
                body_html = re.sub(r'href="#([^"]+)"', _fix_href, '\n'.join(body_parts))
                chapter_content = (
                    '<?xml version="1.0" encoding="utf-8"?>\n'
                    '<!DOCTYPE html>\n'
                    '<html xmlns="http://www.w3.org/1999/xhtml">\n'
                    '<head>\n'
                    f' <title>{html_module.escape(chapter_title)}</title>\n'
                    ' <link rel="stylesheet" type="text/css" href="../style/default.css"/>\n'
                    '</head>\n'
                    '<body>\n'
                    + body_html +
                    '\n</body>\n</html>'
                )
                chapter = epub.EpubHtml(
                    title=chapter_title,
                    file_name=chapter_file_name,
                    lang='en',
                )
                chapter.content = chapter_content.encode('utf-8')
                chapter.add_item(nav_css)
                book.add_item(chapter)
                chapters.append(chapter)
                toc.append(epub.Link(chapter_file_name, chapter_title, chapter_id))
            logger.warning(f"EPUBWriter.run_from_mobi: total chapters created={len(chapters)}")
            if not chapters:
                raise ValueError("No content to write")
            book.toc = toc
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            spine_items = []
            if cover_chapter:
                spine_items.append(cover_chapter)
            spine_items.extend(chapters)
            spine_items.append('nav')
            book.spine = spine_items
            epub.write_epub(self.output_path, book)
            logger.info(f"EPUBWriter: saved to {self.output_path}")
            self.finished.emit(self.output_path, False)
        except Exception as e:
            logger.exception("EPUBWriter error")
            self.finished.emit(str(e), True)

    def run_from_azw(self):
        self.run_from_mobi()
 
    def _restore_inline_html(self, text: str, para: dict) -> str:
        inline_map = para.get('inline_formatting_map', {})
        non_trans = para.get('non_translatable_placeholders', {})
        reserve_elements = para.get('reserve_elements', [])
        prefix_tags = para.get('prefix_reserve_tags', [])
        suffix_tags = para.get('suffix_reserve_tags', [])
        processing_mode = para.get('processing_mode', 'inline')

        for tag_id, info in non_trans.items():
            nt_marker = f'<nt_{tag_id:02d}/>'
            text = text.replace(nt_marker, info.get('full_match', info.get('content', '')))

        def _sub_inline(m):
            tid = int(m.group(1))
            inner = m.group(2)
            if tid not in inline_map:
                return inner
            info = inline_map[tid]
            tag = info.get('tag', 'span')
            attrs = dict(info.get('attributes', {}))
            cls_str = attrs.get('class', '')
            if cls_str and self.css_class_styles:
                extra_style_parts = []
                for c in cls_str.split():
                    for prop, val in self.css_class_styles.get(c, {}).items():
                        extra_style_parts.append(f'{prop}:{val}')
                if extra_style_parts:
                    existing = attrs.get('style', '')
                    combined = (existing.rstrip('; ') + '; ' if existing else '') + '; '.join(extra_style_parts)
                    attrs['style'] = combined
            attr_str = ''.join(
                f' {k}="{html_module.escape(str(v))}"'
                for k, v in attrs.items()
                if not k.startswith('{')
            )
            return f'<{tag}{attr_str}>{inner}</{tag}>'

        for _ in range(10):
            new_text = re.sub(r'<p_(\d{2})>(.*?)</p_\1>', _sub_inline, text, flags=re.DOTALL)
            if new_text == text:
                break
            text = new_text

        text = re.sub(r'</?p_\d{2}>', '', text)

        if processing_mode == 'legacy':
            _res_split = re.compile(r'(<id_\d{2}>)')
            segments = _res_split.split(text)
            text = ''.join(
                s if re.match(r'^<id_\d{2}>$', s) else html_module.escape(s)
                for s in segments
            )

        def _sub_reserve(m):
            idx = int(m.group(1))
            if idx < len(reserve_elements):
                rv = reserve_elements[idx]
                rv = re.sub(r'\s+xmlns(?::\w+)?="[^"]*"', '', rv)
                return rv
            return ''

        text = re.sub(r'<id_(\d{2})>', _sub_reserve, text)

        def _boundary_html(tags):
            parts = []
            for tag in tags:
                bm = re.match(r'<id_(\d{2})>', tag)
                if bm:
                    idx = int(bm.group(1))
                    if idx < len(reserve_elements):
                        rv = reserve_elements[idx]
                        rv = re.sub(r'\s+xmlns(?::\w+)?="[^"]*"', '', rv)
                        parts.append(rv)
            return ''.join(parts)

        prefix_html = _boundary_html(prefix_tags)
        suffix_html = _boundary_html(suffix_tags)
        if prefix_html or suffix_html:
            text = prefix_html + text + suffix_html

        return text
 
    def run_from_azw3(self):
        _HEADING_TAGS = {'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}
        _HEADING_TYPES = _HEADING_TAGS | {'heading'}
        try:
            book = epub.EpubBook()
            title_frag = next(
                (p for p in self.paragraphs if p.get('element_type') == 'book_title'), None
            )
            if title_frag:
                resolved_title = _resolve_text(title_frag, keep_html=False)
                if resolved_title:
                    self.source_title = resolved_title
            book.set_title(self.source_title)
            book.set_language(self.lang)

            extra_css_rules = ''
            for cls_name, props in self.css_class_styles.items():
                decls = ' '.join(f'{k}: {v};' for k, v in props.items())
                if decls:
                    extra_css_rules += f'.{cls_name} {{ {decls} }}\n'
            css_content = (
                'body { margin: 1em; }\n'
                'p { margin: 0.3em 0; }\n'
                'h1 { font-size: 1.6em; margin: 1em 0 0.5em; }\n'
                'h2 { font-size: 1.4em; margin: 0.9em 0 0.4em; }\n'
                'h3 { font-size: 1.2em; margin: 0.8em 0 0.4em; }\n'
                'h4, h5, h6 { font-size: 1em; margin: 0.8em 0 0.4em; }\n'
                'blockquote { margin: 0.5em 2em; font-style: italic; }\n'
                '.footer { font-size: 0.8em; color: #666; }\n'
                'img { max-width: 100%; height: auto; display: block; margin: 0.5em auto; }\n'
                'a { text-decoration: underline; }\n'
            ) + (('\n' + extra_css_rules) if extra_css_rules else '')
            nav_css = epub.EpubItem(
                uid='style_default',
                file_name='style/default.css',
                media_type='text/css',
                content=css_content.encode('utf-8'),
            )
            book.add_item(nav_css)

            img_lookup = {}
            for res in self.image_resources:
                raw_fn = res['file_name'] or ''
                file_name = raw_fn if '/' in raw_fn else 'images/' + raw_fn
                safe_id = re.sub(r'[^\w]', '_', res['id'])
                epub_img = epub.EpubItem(
                    uid='imgres_' + safe_id,
                    file_name=file_name,
                    media_type=res['media_type'],
                    content=res['content'],
                )
                book.add_item(epub_img)
                for key in [
                    res['id'], '#' + res['id'], raw_fn, os.path.basename(raw_fn),
                    raw_fn.replace('\\', '/'), os.path.normpath(raw_fn).replace('\\', '/'),
                    '../' + raw_fn, '../' + raw_fn.replace('\\', '/'),
                ]:
                    if key and key not in img_lookup:
                        img_lookup[key] = file_name
                fn_parts = raw_fn.replace('\\', '/').split('/')
                if len(fn_parts) >= 2:
                    tail2 = '/'.join(fn_parts[-2:])
                    if tail2 not in img_lookup:
                        img_lookup[tail2] = file_name

            def _find_image(href):
                if not href:
                    return None
                candidates = [
                    href, href.lstrip('#'), os.path.basename(href),
                    href.replace('\\', '/'),
                    os.path.normpath(href).replace('\\', '/').lstrip('/'),
                    '../' + href,
                ]
                parts = href.replace('\\', '/').split('/')
                if len(parts) >= 2:
                    candidates.append('/'.join(parts[-2:]))
                for c in candidates:
                    if c in img_lookup:
                        return img_lookup[c]
                return None

            cover_res = next((r for r in self.image_resources if r.get('is_cover')), None)
            if cover_res:
                safe_id = re.sub(r'[^\w]', '_', cover_res['id'])
                book.add_metadata('OPF', 'meta', '', {'name': 'cover', 'content': 'imgres_' + safe_id})

            pages_map = {}
            for para in self.paragraphs:
                if para.get('item_href') == 'metadata':
                    continue
                key = para.get('item_href', '0')
                pages_map.setdefault(key, []).append(para)

            logger.warning(
                f"css_class_styles sample: class_s91c={self.css_class_styles.get('class_s91c')} "
                f"class1={self.css_class_styles.get('class1')} "
                f"class_s4v={self.css_class_styles.get('class_s4v')}"
            )

            chapters = []
            toc = []
            used_file_names = set()
            for idx, (chapter_key, paras) in enumerate(pages_map.items()):
                raw_name = (
                    os.path.splitext(os.path.basename(chapter_key))[0]
                    if chapter_key and chapter_key not in ('0', '')
                    else f'chapter_{idx}'
                )
                if not raw_name:
                    raw_name = f'chapter_{idx}'
                candidate = f'{raw_name}.xhtml'
                if candidate in used_file_names:
                    candidate = f'{raw_name}_{idx}.xhtml'
                used_file_names.add(candidate)
                chapter_file_name = candidate
                chapter_id = os.path.splitext(candidate)[0]

                chapter_title = None
                first_heading = next((p for p in paras if p.get('element_type') in _HEADING_TYPES), None)
                if first_heading:
                    raw = _resolve_text(first_heading, keep_html=False)
                    if raw:
                        chapter_title = raw[:60]
                if not chapter_title:
                    first_text = next((p for p in paras if p.get('element_type') not in ('image',) and (p.get('translated_text') or p.get('original_text', ''))), None)
                    if first_text:
                        raw = _resolve_text(first_text, keep_html=False)
                        if raw:
                            words = raw.split()
                            chapter_title = ' '.join(words[:10])
                            if len(words) > 10:
                                chapter_title = chapter_title.rstrip(',.;:!?…') + '…'
                if not chapter_title:
                    chapter_title = f'Chapter {idx + 1}'

                body_parts = []
                list_stack = []

                def _close_until_list(target_id, _stack=list_stack, _parts=body_parts):
                    while _stack and _stack[-1]['list_id'] != target_id:
                        top = _stack.pop()
                        _parts.append('</li>')
                        _parts.append(f'</{top["list_type"]}>')

                def _close_all_lists(_stack=list_stack, _parts=body_parts):
                    while _stack:
                        top = _stack.pop()
                        _parts.append('</li>')
                        _parts.append(f'</{top["list_type"]}>')

                for para in paras:
                    etype = para.get('element_type', 'paragraph')
                    lm = para.get('list_metadata')

                    if etype == 'image':
                        href = para.get('image_href', '')
                        file_name = _find_image(href)
                        if lm:
                            list_id = lm['list_id']
                            li_index = lm['li_index']
                            stack_ids = [s['list_id'] for s in list_stack]
                            if list_id not in stack_ids:
                                lt = lm['list_type']
                                lc = lm.get('list_class', '')
                                lc_attr = f' class="{html_module.escape(lc)}"' if lc else ''
                                body_parts.append(f'<{lt}{lc_attr}>')
                                lic = lm.get('li_class', '')
                                li_ca = f' class="{html_module.escape(lic)}"' if lic else ''
                                body_parts.append(f'<li{li_ca}>')
                                list_stack.append({'list_id': list_id, 'list_type': lt, 'current_li': li_index})
                            else:
                                _close_until_list(list_id)
                                top = list_stack[-1]
                                if top['current_li'] != li_index:
                                    body_parts.append('</li>')
                                    lic = lm.get('li_class', '')
                                    li_ca = f' class="{html_module.escape(lic)}"' if lic else ''
                                    body_parts.append(f'<li{li_ca}>')
                                    top['current_li'] = li_index
                            if file_name:
                                body_parts.append(f'<img src="{file_name}" alt=""/>')
                        else:
                            _close_all_lists()
                            if file_name:
                                body_parts.append(f'<img src="{file_name}" alt=""/>')
                        continue

                    text = _resolve_text(para, keep_html=True)
                    if not text:
                        continue
                    text = self._restore_inline_html(text, para)

                    if para.get('processing_mode') == 'inline':
                        t = text
                    else:
                        has_html = bool(re.search(r'<[^>]+>', text))
                        t = text if has_html else html_module.escape(text)

                    t = re.sub(r'(href="(?!https?://)(?!mailto:)[^"]*?)\.html((?:#[^"]*)?)"', r'\1.xhtml\2"', t)
                    t = re.sub(r'href="(?!https?://)(?!mailto:)[^"#]*/([^"/]+\.xhtml(?:#[^"]*)?)"', r'href="\1"', t)
                    t = re.sub(r'src="([^"]+)"', lambda m: 'src="' + (_find_image(m.group(1)) or m.group(1)) + '"', t)

                    if self.css_class_styles:
                        def _inject_css(m, _css=self.css_class_styles):
                            tn = m.group(1)
                            ap = m.group(2)
                            cm = re.search(r'\bclass="([^"]*)"', ap)
                            if not cm:
                                return m.group(0)
                            extra = []
                            for c in cm.group(1).split():
                                for prop, val in _css.get(c, {}).items():
                                    extra.append(f'{prop}:{val}')
                            if not extra:
                                return m.group(0)
                            sm = re.search(r'\bstyle="([^"]*)"', ap)
                            if sm:
                                ns = sm.group(1).rstrip('; ') + '; ' + '; '.join(extra)
                                ap = ap[:sm.start(1)] + ns + ap[sm.end(1):]
                            else:
                                ap = ap + f' style="{"; ".join(extra)}"'
                            return f'<{tn}{ap}>'
                        t = re.sub(r'<(\w+)((?:\s[^>]*)?)>', _inject_css, t)

                    style_attr = ''
                    ks = para.get('kindle_style', '')
                    if ks:
                        style_attr = f' style="{html_module.escape(ks)}"'
                    id_attr = ''
                    eid = para.get('element_id', '')
                    if eid:
                        id_attr = f' id="{html_module.escape(eid)}"'
                    class_attr = ''
                    ecls = para.get('element_class', '')
                    if ecls:
                        class_attr = f' class="{html_module.escape(ecls)}"'

                    if lm:
                        list_id = lm['list_id']
                        li_index = lm['li_index']
                        list_type = lm['list_type']
                        list_class = lm.get('list_class', '')
                        stack_ids = [s['list_id'] for s in list_stack]

                        if list_id not in stack_ids:
                            lc_attr = f' class="{html_module.escape(list_class)}"' if list_class else ''
                            body_parts.append(f'<{list_type}{lc_attr}>')
                            lic = lm.get('li_class', '')
                            li_ca = f' class="{html_module.escape(lic)}"' if lic else ''
                            body_parts.append(f'<li{li_ca}>')
                            list_stack.append({'list_id': list_id, 'list_type': list_type, 'current_li': li_index})
                        else:
                            _close_until_list(list_id)
                            top = list_stack[-1]
                            if top['current_li'] != li_index:
                                body_parts.append('</li>')
                                lic = lm.get('li_class', '')
                                li_ca = f' class="{html_module.escape(lic)}"' if lic else ''
                                body_parts.append(f'<li{li_ca}>')
                                top['current_li'] = li_index

                        if etype == 'li':
                            body_parts.append(t)
                        elif etype in _HEADING_TAGS:
                            body_parts.append(f'<{etype}{id_attr}{class_attr}{style_attr}>{t}</{etype}>')
                        elif etype == 'heading':
                            body_parts.append(f'<h2{id_attr}{class_attr}{style_attr}>{t}</h2>')
                        elif etype == 'footer':
                            body_parts.append(f'<p class="footer"{id_attr}{style_attr}>{t}</p>')
                        elif etype == 'blockquote':
                            body_parts.append(f'<blockquote><p{id_attr}{class_attr}{style_attr}>{t}</p></blockquote>')
                        else:
                            body_parts.append(f'<p{id_attr}{class_attr}{style_attr}>{t}</p>')
                    else:
                        _close_all_lists()

                        if etype in _HEADING_TAGS:
                            body_parts.append(f'<{etype}{id_attr}{class_attr}{style_attr}>{t}</{etype}>')
                        elif etype == 'heading':
                            body_parts.append(f'<h2{id_attr}{class_attr}{style_attr}>{t}</h2>')
                        elif etype == 'footer':
                            body_parts.append(f'<p class="footer"{id_attr}{style_attr}>{t}</p>')
                        elif etype == 'blockquote':
                            body_parts.append(f'<blockquote><p{id_attr}{class_attr}{style_attr}>{t}</p></blockquote>')
                        else:
                            body_parts.append(f'<p{id_attr}{class_attr}{style_attr}>{t}</p>')

                _close_all_lists()

                if not body_parts:
                    continue

                chapter_content = (
                    '<?xml version="1.0" encoding="utf-8"?>\n'
                    '<!DOCTYPE html>\n'
                    '<html xmlns="http://www.w3.org/1999/xhtml">\n'
                    '<head>\n'
                    f' <title>{html_module.escape(chapter_title)}</title>\n'
                    ' <link rel="stylesheet" type="text/css" href="style/default.css"/>\n'
                    '</head>\n'
                    '<body>\n'
                    + '\n'.join(body_parts) +
                    '\n</body>\n</html>'
                )

                chapter = epub.EpubHtml(title=chapter_title, file_name=chapter_file_name, lang=self.lang)
                chapter.content = chapter_content.encode('utf-8')
                chapter.add_item(nav_css)
                book.add_item(chapter)
                chapters.append(chapter)
                toc.append(epub.Link(chapter_file_name, chapter_title, chapter_id))

            if not chapters:
                raise ValueError("No content to write")
            book.toc = toc
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            book.spine = chapters + ['nav']
            epub.write_epub(self.output_path, book)
            logger.info(f"EPUBWriter: saved to {self.output_path}")
            self.finished.emit(self.output_path, False)
        except Exception as e:
            logger.exception("EPUBWriter error")
            self.finished.emit(str(e), True)

class FB2Writer(QThread):
    finished = pyqtSignal(str, bool)

    def __init__(self, paragraphs, output_path, source_title="Translated",
                 image_resources=None, source_format='epub'):
        super().__init__()
        self.paragraphs = paragraphs
        self.output_path = output_path
        self.source_title = source_title
        self.image_resources = image_resources or []
        self.source_format = source_format

    def run(self):
        self.run_from_fb2()

    def run_from_fb2(self):
        try:
            import base64 as _base64

            FB2_NS = 'http://www.gribuser.ru/xml/fictionbook/2.0'
            XLINK_NS = 'http://www.w3.org/1999/xlink'
            nsmap = {None: FB2_NS, 'l': XLINK_NS}

            root = etree.Element(f'{{{FB2_NS}}}FictionBook', nsmap=nsmap)
            desc = etree.SubElement(root, f'{{{FB2_NS}}}description')
            title_info = etree.SubElement(desc, f'{{{FB2_NS}}}title-info')
            book_title_elem = etree.SubElement(title_info, f'{{{FB2_NS}}}book-title')
            book_title_elem.text = self.source_title

            cover_res = next((r for r in self.image_resources if r.get('is_cover')), None)
            if cover_res:
                coverpage_elem = etree.SubElement(title_info, f'{{{FB2_NS}}}coverpage')
                cov_img = etree.SubElement(coverpage_elem, f'{{{FB2_NS}}}image')
                cov_img.set(f'{{{XLINK_NS}}}href', f'#{cover_res["id"]}')

            body = etree.SubElement(root, f'{{{FB2_NS}}}body')

            img_lookup = {}
            for res in self.image_resources:
                img_lookup[res['id']] = res['id']
                img_lookup['#' + res['id']] = res['id']
                img_lookup[res['file_name']] = res['id']
                img_lookup[os.path.basename(res['file_name'])] = res['id']

            pages_map = {}
            for para in self.paragraphs:
                key = para.get('item_href', '0')
                pages_map.setdefault(key, []).append(para)

            for chapter_key, paras in pages_map.items():
                section = etree.SubElement(body, f'{{{FB2_NS}}}section')
                for para in paras:
                    etype = para.get('element_type', 'paragraph')

                    if etype == 'image':
                        href = para.get('image_href', '')
                        binary_id = (
                            img_lookup.get(href)
                            or img_lookup.get(href.lstrip('#'))
                            or img_lookup.get(os.path.basename(href))
                        )
                        if binary_id:
                            img_elem = etree.SubElement(section, f'{{{FB2_NS}}}image')
                            img_elem.set(f'{{{XLINK_NS}}}href', f'#{binary_id}')
                        continue

                    if etype == 'empty_line':
                        etree.SubElement(section, f'{{{FB2_NS}}}empty-line')
                        continue

                    aligned_html = para.get('aligned_translated_html') if para.get('is_translated') else None
                    if aligned_html:
                        segs = _parse_aligned_html_segments(aligned_html)
                        if segs:
                            if etype == 'heading':
                                title_elem = etree.SubElement(section, f'{{{FB2_NS}}}title')
                                p_elem = etree.SubElement(title_elem, f'{{{FB2_NS}}}p')
                            elif etype == 'blockquote':
                                epigraph = etree.SubElement(section, f'{{{FB2_NS}}}epigraph')
                                p_elem = etree.SubElement(epigraph, f'{{{FB2_NS}}}p')
                            else:
                                p_elem = etree.SubElement(section, f'{{{FB2_NS}}}p')
                            _fill_fb2_from_segments(p_elem, segs, FB2_NS)
                            continue

                    text = _get_raw_text_for_inline(para)
                    if not text:
                        continue

                    inline_map = para.get('inline_formatting_map') or {}

                    if etype == 'heading':
                        title_elem = etree.SubElement(section, f'{{{FB2_NS}}}title')
                        p_elem = etree.SubElement(title_elem, f'{{{FB2_NS}}}p')
                    elif etype == 'blockquote':
                        epigraph = etree.SubElement(section, f'{{{FB2_NS}}}epigraph')
                        p_elem = etree.SubElement(epigraph, f'{{{FB2_NS}}}p')
                    else:
                        p_elem = etree.SubElement(section, f'{{{FB2_NS}}}p')

                    _fill_fb2_paragraph(p_elem, text, inline_map, FB2_NS)

            for res in self.image_resources:
                binary_elem = etree.SubElement(root, f'{{{FB2_NS}}}binary')
                binary_elem.set('id', res['id'])
                binary_elem.set('content-type', res['media_type'])
                binary_elem.text = _base64.b64encode(res['content']).decode('ascii')

            tree = etree.ElementTree(root)
            with open(self.output_path, 'wb') as f:
                tree.write(f, xml_declaration=True, encoding='utf-8', pretty_print=True)

            logger.info(f"FB2Writer: saved to {self.output_path}")
            self.finished.emit(self.output_path, False)

        except Exception as e:
            logger.exception("FB2Writer error")
            self.finished.emit(str(e), True)


class DOCXWriter(QThread):
    finished = pyqtSignal(str, bool)

    def __init__(self, paragraphs, output_path, source_title="Translated",
                 image_resources=None, source_format='epub'):
        super().__init__()
        self.paragraphs = paragraphs
        self.output_path = output_path
        self.source_title = source_title
        self.image_resources = image_resources or []
        self.source_format = source_format

    def run(self):
        self.run_from_docx()

    def run_from_docx(self):
        try:
            import io
            from docx import Document
            from docx.shared import Inches

            doc = Document()
            doc.add_heading(self.source_title, level=0)

            img_lookup = {}
            for res in self.image_resources:
                img_lookup[res['id']] = res['content']
                img_lookup['#' + res['id']] = res['content']
                img_lookup[res['file_name']] = res['content']
                img_lookup[os.path.basename(res['file_name'])] = res['content']

            for para in self.paragraphs:
                etype = para.get('element_type', 'paragraph')

                if etype == 'image':
                    href = para.get('image_href', '')
                    img_data = (
                        img_lookup.get(href)
                        or img_lookup.get(href.lstrip('#'))
                        or img_lookup.get(os.path.basename(href))
                    )
                    if img_data:
                        try:
                            doc.add_picture(io.BytesIO(img_data), width=Inches(4.5))
                        except Exception as img_err:
                            logger.warning(f"DOCXWriter: could not insert image {href}: {img_err}")
                    continue

                aligned_html = para.get('aligned_translated_html') if para.get('is_translated') else None
                if aligned_html:
                    segs = _parse_aligned_html_segments(aligned_html)
                    if segs:
                        if etype == 'heading':
                            p = doc.add_heading('', level=2)
                        elif etype == 'blockquote':
                            p = doc.add_paragraph(style='Quote')
                        else:
                            p = doc.add_paragraph()
                        _fill_docx_from_segments(p, segs)
                        continue

                text = _get_raw_text_for_inline(para)
                if not text:
                    continue

                inline_map = para.get('inline_formatting_map') or {}

                if etype == 'heading':
                    p = doc.add_heading('', level=2)
                elif etype == 'blockquote':
                    p = doc.add_paragraph(style='Quote')
                else:
                    p = doc.add_paragraph()

                _fill_docx_paragraph(p, text, inline_map)

            doc.save(self.output_path)
            logger.info(f"DOCXWriter: saved to {self.output_path}")
            self.finished.emit(self.output_path, False)

        except Exception as e:
            logger.exception("DOCXWriter error")
            self.finished.emit(str(e), True)
 

class FileProcessorFactory:
    @staticmethod
    def create_processor(file_type: str, app_settings: dict) -> FileProcessor:
        if file_type == "epub":
            return EPUBProcessor(app_settings)
        elif file_type == "srt":
            return SRTProcessor()
        elif file_type == "txt":
            return TXTProcessor()
        elif file_type == "pdf":
            return PDFProcessor(app_settings)
        elif file_type in ("mobi", "azw"):
            return MobiProcessor(app_settings)
        elif file_type == "azw3":
            return AZW3Processor(app_settings)
        elif file_type == "fb2":
            return FB2Processor(app_settings)
        elif file_type == "docx":
            return DOCXProcessor(app_settings)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def create_writer(output_format: str, paragraphs, output_path: str,
                      source_title: str = "Translated", image_resources=None,
                      source_format: str = 'epub', lang: str = 'en'):
        if output_format == 'fb2':
            return FB2Writer(paragraphs, output_path, source_title, image_resources, source_format)
        elif output_format == 'docx':
            return DOCXWriter(paragraphs, output_path, source_title, image_resources, source_format)
        else:
            return EPUBWriter(paragraphs, output_path, source_title, image_resources, source_format, lang)
 
